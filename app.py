import streamlit as st
import pandas as pd
import numpy as np
import matplotlib

matplotlib.use(
    "Agg"
)  # Set non-GUI backend for matplotlib to prevent issues in Streamlit
import matplotlib.pyplot as plt
import seaborn as sns
from langchain_openai import AzureChatOpenAI
from langchain_experimental.utilities import PythonREPL
import io
import os
import glob
import re
import traceback
import time
import asyncio
import aiohttp
import json
import xml.etree.ElementTree as ET
from PIL import Image
from contextlib import redirect_stdout
import tempfile
from typing import List, Tuple, Dict
from datetime import datetime
from dotenv import load_dotenv
from markdown_to_docx import generate_gpt_analysis_docx
from document_utils import process_document_file, create_document_chunks
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from io import BytesIO  # Already imported but good to note its use here
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64  # Import base64 for image encoding/decoding


# Import Pydantic models
from models import (
    AnalysisResults,
    AnalysisIteration,
    DocumentInfo,
    DataFrameMetadata,
    DocumentMetadata,
)

# --- Prompts (ensure prompts.py is in the same directory or accessible) ---
# Example: from prompts import SYSTEM_PROMPT_DATAFRAME_ANALYSIS, SYSTEM_PROMPT_GENERAL_QA
# Actual prompts used will be defined/loaded in the get_code_from_llm function or globally.
from prompts import (
    followup_system_prompt_template,
    system_prompt_expert_questions,
    expert1_system_prompt,
    expert2_system_prompt,
    expert3_system_prompt,
)


load_dotenv()

# --- DOCX Generation Function (Assumed to be in markdown_to_docx.py) ---
try:
    # Ensure markdown_to_docx.py is in the same directory or accessible via PYTHONPATH
    # This file should contain: def generate_gpt_analysis_docx(...)
    from markdown_to_docx import generate_gpt_analysis_docx

    DOCX_AVAILABLE = True
except ImportError:
    # st.warning( # Commented out to avoid premature Streamlit calls before UI rendering
    #     "`markdown_to_docx.py` not found or `generate_gpt_analysis_docx` is missing. "
    #     "DOCX generation will not be available."
    # )
    DOCX_AVAILABLE = False

    # Define a placeholder function if the import fails, so the app doesn't crash
    def generate_gpt_analysis_docx(*args, **kwargs):
        st.error(
            "DOCX generation is unavailable because the necessary module/function could not be imported."
        )
        return None


# --- App Configuration ---
st.set_page_config(
    page_title="Iterative AI Analyzer",
    layout="wide",  # Use wide layout for better display of code and plots
    page_icon="🤖",
    initial_sidebar_state="auto",
)

# --- Environment/Secret Keys ---
# Fetch from Streamlit secrets or environment variables
os.environ["AZURE_OPENAI_ENDPOINT"] = os.getenv("AZURE_OPENAI_ENDPOINT")
os.environ["AZURE_OPENAI_API_KEY"] = os.getenv("AZURE_OPENAI_API_KEY")

API_KEY = os.environ.get("AZURE_OPENAI_API_KEY", st.secrets.get("AZURE_OPENAI_API_KEY"))
AZURE_ENDPOINT = os.environ.get(
    "AZURE_OPENAI_ENDPOINT", st.secrets.get("AZURE_OPENAI_ENDPOINT")
)
PASSWORD = os.environ.get(
    "PASSWORD", st.secrets.get("password", "test_password")
)  # Default for local dev
AZURE_DEPLOYMENT = os.environ.get(
    "AZURE_DEPLOYMENT", st.secrets.get("AZURE_DEPLOYMENT")
)
API_VERSION = os.environ.get(
    "AZURE_API_VERSION", st.secrets.get("API_VERSION", "2023-05-15")
)  # Default to latest version
TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY", st.secrets.get("TAVILY_API_KEY"))


# --- Session State Initialization ---
if "password_correct" not in st.session_state:
    st.session_state.password_correct = False
if "login_attempts" not in st.session_state:
    st.session_state.login_attempts = 0

if "dataframes" not in st.session_state:  # Dictionary to store all uploaded dataframes
    st.session_state.dataframes = {}
if (
    "dataframe_metadata" not in st.session_state
):  # Dictionary to store metadata about each dataframe
    st.session_state.dataframe_metadata = {}
if (
    "active_df" not in st.session_state
):  # Name of the currently active dataframe for display
    st.session_state.active_df = None
if (
    "gpt_working_dfs" not in st.session_state
):  # Dictionary of dataframes being modified by GPT
    st.session_state.gpt_working_dfs = {}
if (
    "document_texts" not in st.session_state
):  # Dictionary to store extracted text from documents
    st.session_state.document_texts = {}
if (
    "document_vector_store" not in st.session_state
):  # Vector store for uploaded documents
    st.session_state.document_vector_store = None
if "use_web_search" not in st.session_state:  # Flag to enable web search
    st.session_state.use_web_search = True
if "analysis_results" not in st.session_state:  # Stores results of the last analysis
    st.session_state.analysis_results = AnalysisResults().model_dump()
if "outputs_path" not in st.session_state:  # Path for saving plots
    st.session_state.outputs_path = tempfile.mkdtemp(prefix="gpt_plots_")
if "current_analysis_timestamp" not in st.session_state:
    st.session_state.current_analysis_timestamp = ""
if "last_agent_question" not in st.session_state:
    st.session_state.last_agent_question = ""
if (
    "categorical_mappings" not in st.session_state
):  # To store how categoricals were mapped
    st.session_state.categorical_mappings = {}
if (
    "retrieved_chunks" not in st.session_state
):  # To store all retrieved chunks across iterations
    st.session_state.retrieved_chunks = []
if "search_history" not in st.session_state:  # To store search query history
    st.session_state.search_history = []
if "pubmed_search_terms" not in st.session_state:  # To store PubMed search terms
    st.session_state.pubmed_search_terms = ""
if "articles" not in st.session_state:  # To store PubMed articles
    st.session_state.articles = []
if (
    "older_pubmed_articles_alert" not in st.session_state
):  # Flag for older PubMed articles alert
    st.session_state.older_pubmed_articles_alert = False
if (
    "is_medical_question" not in st.session_state
):  # Flag to track if current question is medical
    st.session_state.is_medical_question = False
if "search_strategy" not in st.session_state:  # To store the search strategy used
    st.session_state.search_strategy = {}
if "ragas_model" not in st.session_state:  # To store the RAGAS evaluation model
    st.session_state.ragas_model = "gpt-4o-mini"
if (
    "validated_section1" not in st.session_state
):  # To store the validated section 1 content
    st.session_state.validated_section1 = ""
if "validation_results" not in st.session_state:  # To store validation results
    st.session_state.validation_results = None
if "current_section1" not in st.session_state:  # To store the current section 1 content
    st.session_state.current_section1 = ""
if "followup_chat_thread" not in st.session_state:
    st.session_state.followup_chat_thread = []
if "google_search_followup_results" not in st.session_state:
    st.session_state.google_search_followup_results = ""
if "experts" not in st.session_state:
    st.session_state.experts = []
if "expert_domains" not in st.session_state:
    st.session_state.expert_domains = []
if "expert_questions" not in st.session_state:
    st.session_state.expert_questions = []
if "expert_answers" not in st.session_state:
    st.session_state.expert_answers = []
if "experts" not in st.session_state:
    st.session_state.experts = []
if "expert_domains" not in st.session_state:
    st.session_state.expert_domains = []
if "expert_questions" not in st.session_state:
    st.session_state.expert_questions = []
if "expert_answers" not in st.session_state:
    st.session_state.expert_answers = []


# --- Emojis for Chat Display ---
role_emojis = {
    "user": "👤",
    "assistant": "🤖",
}


# --- Authentication ---
def check_password():
    """Checks password and manages login state."""
    if st.session_state.get("password_correct", False):  # Use .get for safety
        return True

    st.sidebar.title("🔐 Login Required")
    password_attempt = st.sidebar.text_input(
        "Password", type="password", key="password_input_main"
    )

    if st.sidebar.button("Login", key="login_button_main"):
        if password_attempt == PASSWORD:
            st.session_state.password_correct = True
            st.session_state.login_attempts = 0
            st.rerun()  # Rerun to hide login and show app
        else:
            st.session_state.password_correct = False
            st.session_state.login_attempts = (
                st.session_state.get("login_attempts", 0) + 1
            )
            st.sidebar.error(
                f"Incorrect password. Attempts: {st.session_state.login_attempts}"
            )
    return False


# --- Helper function to add hyperlink to Word documents ---
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --- Helper function to format a single citation (AMA-like) ---
def format_citation(citation, idx=None):
    title = citation.get("title") or citation.get("context") or ""
    url = citation.get("url") or citation.get("link") or ""
    year = citation.get("year", "")
    pmid = citation.get("id", "")
    citation_str = f"{title}"
    if year:
        citation_str += f". {year}."
    if url:
        citation_str += f" {url}"
    if pmid:
        citation_str += f" PMID: {pmid}."
    if idx is not None:
        citation_str = f"{idx}. {citation_str}"
    return citation_str.strip()

def extract_section1(summary_text: str) -> Tuple[str, str]:
    """
    Extracts Section 1 from the given summary_text.
    Section 1 is defined as the text before the first occurrence of 
    "2. Additional Insights from the Model's Knowledge".
    Returns a tuple (section1, remainder). Caches section1 in st.session_state.
    """
    import re
    if st.session_state.get("current_section1"):
        section1 = st.session_state.current_section1
        return section1, ""
    pattern = r"(2\. Additional Insights from the Model's Knowledge.*)"
    split_content = re.split(pattern, summary_text, maxsplit=1, flags=re.DOTALL)
    if len(split_content) > 1:
        section1 = split_content[0].rstrip()
        remainder = split_content[1].lstrip()
    else:
        section1 = summary_text
        remainder = ""
    st.session_state.current_section1 = section1
    return section1, remainder


# --- Helper function to convert Markdown text to a Word document ---
def markdown_to_word(markdown_text, original_question_text=None, citations_list=None):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    if original_question_text:
        h_user_q = doc.add_heading("User Question", level=2)
        h_user_q.runs[0].font.name = "Calibri"
        h_user_q.runs[0].font.size = Pt(16)
        h_user_q.runs[0].bold = True
        h_user_q.paragraph_format.space_after = Pt(6)
        p_user_q = doc.add_paragraph(original_question_text)
        p_user_q.paragraph_format.space_after = Pt(6)
        for run_item in p_user_q.runs:
            run_item.font.name = "Calibri"
            run_item.font.size = Pt(11)

    lines = markdown_text.split("\n")
    for line in lines:
        if line.startswith("#"):
            heading_level = len(line) - len(line.lstrip("#"))
            text = line[heading_level:].strip()
            h = doc.add_heading(text, level=min(heading_level, 6))
            if h.runs:
                h_run = h.runs[0]
                h_run.font.name = "Calibri"
                if heading_level == 1:
                    h_run.font.size = Pt(18)
                elif heading_level == 2:
                    h_run.font.size = Pt(16)
                elif heading_level == 3:
                    h_run.font.size = Pt(14)
                else:
                    h_run.font.size = Pt(12)
                h_run.bold = True
            h.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            segments = re.split(r"(\[.*?\]\(.*?\)|\*\*.*?\*\*|\*.*?\*)", line)
            for seg in segments:
                if seg.startswith("[") and seg.endswith(")"):
                    match = re.match(r"\[([^]]+)\]\(([^)]+)\)", seg)
                    if match:
                        add_hyperlink(p, match.group(2), match.group(1))
                    else:
                        p.add_run(seg).font.name = "Calibri"
                elif seg.startswith("**") and seg.endswith("**"):
                    p.add_run(seg[2:-2]).bold = True
                    p.runs[-1].font.name = "Calibri"
                elif seg.startswith("*") and seg.endswith("*"):
                    p.add_run(seg[1:-1]).italic = True
                    p.runs[-1].font.name = "Calibri"
                else:
                    p.add_run(seg).font.name = "Calibri"

    if citations_list and isinstance(citations_list, list) and len(citations_list) > 0:
        unique_urls = sorted(
            list(
                set(
                    c.get("url") or c.get("link")
                    for c in citations_list
                    if c.get("url") or c.get("link")
                )
            )
        )
        if unique_urls:
            doc.add_page_break()
            h_urls = doc.add_heading("Source URLs", level=2)
            if h_urls.runs:
                h_urls_run = h_urls.runs[0]
                h_urls_run.font.name = "Calibri"
                h_urls_run.font.size = Pt(16)
                h_urls_run.bold = True
            h_urls.paragraph_format.space_after = Pt(6)
            for url_item in unique_urls:
                p_url = doc.add_paragraph()
                p_url.paragraph_format.space_after = Pt(3)
                add_hyperlink(p_url, url_item, url_item)

        h_refs = doc.add_heading("References", level=2)  # Add "References" heading
        if h_refs.runs:  # Check if the newly added heading has runs
            h_refs_run = h_refs.runs[0]
            h_refs_run.font.name = "Calibri"
            h_refs_run.font.size = Pt(16)
            h_refs_run.bold = True
        h_refs.paragraph_format.space_after = Pt(
            6
        )  # Style the paragraph object directly

        for idx, citation_item in enumerate(citations_list, 1):
            citation_str = format_citation(citation_item, idx)
            p_ref = doc.add_paragraph(citation_str, style="List Number")
            p_ref.paragraph_format.space_after = Pt(3)
            for run_item in p_ref.runs:
                run_item.font.name = "Calibri"
                run_item.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Helper function to extract expert information from JSON ---
def extract_expert_info(json_input: str) -> Tuple[List[str], List[str], List[str]]:
    experts, domains, expert_questions = [], [], []
    try:
        data = json.loads(json_input)
        if not isinstance(data, dict) or "rephrased_questions" not in data:
            print(
                "JSON input for extract_expert_info is missing 'rephrased_questions' key or is not a dict."
            )
            return experts, domains, expert_questions

        rephrased_list = data["rephrased_questions"]
        if not isinstance(rephrased_list, list):
            print("'rephrased_questions' is not a list in extract_expert_info.")
            return experts, domains, expert_questions

        for item in rephrased_list:
            if isinstance(item, dict) and all(
                k in item for k in ["expert", "domain", "question"]
            ):
                experts.append(str(item["expert"]))
                domains.append(str(item["domain"]))
                expert_questions.append(str(item["question"]))
            else:
                print(f"Skipping malformed item in 'rephrased_questions': {item}")
    except json.JSONDecodeError as e:
        print(f"Failed to decode JSON in extract_expert_info: {e}")
    except Exception as e:
        print(f"Unexpected error in extract_expert_info: {e}")
    return experts, domains, expert_questions


# --- Helper function to generate DOCX from chat thread ---
def chat_thread_to_docx(
    chat_thread: List[Dict[str, str]],
    original_question: str,
    initial_summary: str,
    role_emojis_param: Dict[str, str],  # Renamed to avoid conflict with global
) -> BytesIO:
    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add original question
    h_q = doc.add_heading("Original Question", level=2)
    h_q.runs[0].font.name = "Calibri"
    p_q = doc.add_paragraph(original_question)
    p_q.paragraph_format.space_after = Pt(6)
    for run_item in p_q.runs:  # Ensure font for question
        run_item.font.name = "Calibri"
        run_item.font.size = Pt(11)

    # Add initial summary
    h_s = doc.add_heading("Initial Summary Provided by AI", level=2)
    h_s.runs[0].font.name = "Calibri"
    # For the summary, which might be markdown, we add it paragraph by paragraph.
    # This is a simplified way to handle potential markdown newlines.
    # More complex markdown (like lists, bold, italic) will appear as raw markdown text.
    summary_lines = initial_summary.split("\n")
    for line in summary_lines:
        p_summary_line = doc.add_paragraph(line)
        p_summary_line.paragraph_format.space_after = Pt(
            3
        )  # Smaller space for lines within summary
        for run_item in p_summary_line.runs:  # Ensure font for summary
            run_item.font.name = "Calibri"
            run_item.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)  # Spacer after summary

    h_chat = doc.add_heading("Follow-up Conversation", level=2)
    h_chat.runs[0].font.name = "Calibri"

    for msg in chat_thread:
        if msg["role"] != "system":  # Don't include system prompts in the document
            role_display = (
                f"{role_emojis_param.get(msg['role'], '')} {msg['role'].capitalize()}"
            )

            p_role = doc.add_paragraph()
            run_role = p_role.add_run(f"{role_display}:")
            run_role.bold = True
            run_role.font.name = "Calibri"
            run_role.font.size = Pt(11)

            # Add message content with basic markdown handling (bold/italic)
            content = msg["content"]
            # Using a paragraph per message line for content to respect newlines from chat
            message_lines = content.split("\n")
            for line_idx, line_text in enumerate(message_lines):
                # If it's not the first line of the message, create a new paragraph for it
                # to preserve multiline formatting from the chat.
                # The first line is added to the paragraph with the role.
                if line_idx == 0:
                    p_content = p_role  # Continue the paragraph with the role
                else:
                    p_content = (
                        doc.add_paragraph()
                    )  # New paragraph for subsequent lines
                    p_content.paragraph_format.left_indent = Pt(
                        36
                    )  # Indent subsequent lines

                segments = re.split(
                    r"(\*\*.*?\*\*|\*.*?\*|```[\s\S]*?```|`.*?`)", line_text
                )
                for seg in segments:
                    if not seg:
                        continue
                    run_seg = p_content.add_run()
                    run_seg.font.name = "Calibri"
                    run_seg.font.size = Pt(11)
                    if seg.startswith("**") and seg.endswith("**"):
                        run_seg.text = seg[2:-2]
                        run_seg.bold = True
                    elif seg.startswith("*") and seg.endswith("*"):
                        run_seg.text = seg[1:-1]
                        run_seg.italic = True
                    elif seg.startswith("```") and seg.endswith("```"):
                        # Add code block with a distinct style (e.g., Courier New)
                        # Remove triple backticks and potential language specifier
                        code_text = seg[3:-3]
                        if (
                            "\n" in code_text and code_text.split("\n")[0].isalnum()
                        ):  # Likely language specifier
                            code_text = "\n".join(code_text.split("\n")[1:])

                        # Add each line of the code block
                        for code_line_text in code_text.strip("\n").split("\n"):
                            p_code_line = doc.add_paragraph(code_line_text)
                            p_code_line.style = "NoSpacing"  # Use a style with minimal spacing for code blocks
                            for run_code in p_code_line.runs:
                                run_code.font.name = (
                                    "Courier New"  # Monospace font for code
                                )
                                run_code.font.size = Pt(10)
                            # Indent code blocks further
                            p_code_line.paragraph_format.left_indent = Pt(72)

                    elif seg.startswith("`") and seg.endswith("`"):
                        run_seg.text = seg[1:-1]
                        run_seg.font.name = "Courier New"  # Monospace for inline code
                        run_seg.font.size = Pt(10)
                    else:
                        run_seg.text = seg

                if (
                    line_idx == 0 and len(message_lines) > 1
                ):  # If first line and more lines to come
                    p_content.paragraph_format.space_after = Pt(
                        0
                    )  # No space after first line of multi-line message
                elif line_idx < len(message_lines) - 1:  # Intermediate lines
                    p_content.paragraph_format.space_after = Pt(0)
                else:  # Last line of message or single line message
                    p_content.paragraph_format.space_after = Pt(6)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- LLM and REPL Setup ---
def get_rag_search_available():
    """Checks if RAG search is available by verifying required environment variables."""
    from rag_utils import (
        GOOGLE_API_KEY,
        GOOGLE_CSE_ID,
        AZURE_OPENAI_API_KEY,
        AZURE_OPENAI_ENDPOINT,
    )

    if not GOOGLE_API_KEY or not GOOGLE_CSE_ID:
        return False, "Google API key or CSE ID not found"

    if not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT:
        return False, "Azure OpenAI credentials not found"

    return True, "RAG search is available"


@st.cache_data(ttl=3600)
def optimize_pubmed_search_terms(query: str, cutting_edge: bool = False) -> str:
    """
    Optimize a query for PubMed search using Azure OpenAI.
    Uses Streamlit caching to avoid repeated optimization of the same query.

    Args:
        query: The original query to optimize
        cutting_edge: Whether to optimize for cutting-edge research (vs. consensus)

    Returns:
        Optimized PubMed search terms
    """
    # Define the system prompt based on whether cutting-edge research is requested
    if cutting_edge:
        system_prompt = """**Role**: You are a highly specialized AI designed to create precise PubMed search queries for medical professionals. Your task is to transform any user question into a reasonably broad search query that retrieves high-quality, evidence-based, and cutting-edge literature, intended to answer the physician's question. Your outputs must follow the provided guidelines and examples precisely.

### **Guidelines for Query Optimization**:

1. **Define the Core Concepts**:
   - Identify the main topic, condition, intervention, or outcome mentioned in the question.
   - Translate these into appropriate MeSH terms, relevant text words, and emerging concepts.

2. **Prioritize High-Quality and Recent Evidence**:
   - Include terms to retrieve both foundational evidence (e.g., "systematic review," "meta-analysis," "guideline") and cutting-edge research (e.g., "novel," "emerging therapies," "recent advances").
   - Emphasize publication types such as "clinical trial," "randomized controlled trial," "cohort study," or "case series" alongside traditional review articles.

3. **Streamline and Expand Terms**:
   - Avoid extraneous words or phrases that do not contribute to the search focus.
   - Include terms that broaden the search to capture emerging trends and innovative approaches.

4. **Leverage Boolean Operators**:
   - Combine MeSH terms, text words, and emerging keywords using **AND**, **OR**, and parentheses for logical grouping.
   - Use these operators to ensure inclusivity (OR) and specificity (AND).

5. **Focus on Both Practical and Innovative Applications**:
   - Align the query with both the practical intent and the potential for novel findings, such as new diagnostic tools, treatments, or clinical decision-making insights.

### **Output Requirement**:
- Provide only the optimized PubMed query for each input question. Do not include additional commentary or extraneous information."""
    else:
        system_prompt = """**Role**: You are a highly specialized AI designed to create precise PubMed search queries for medical professionals. Your task is to transform any user question into a reasonably broad search query that retrieves high-quality, evidence-based literature
intended to retrieve citations that will contain the answer the phyician's question. Your outputs must follow the provided guidelines and examples precisely.

### **Guidelines for Query Optimization**:

1. **Define the Core Concepts**:
   - Identify the main topic, condition, intervention, or outcome mentioned in the question.
   - Translate these into appropriate MeSH terms and relevant text words.
   - Often the question will be very specific, so to ensure relevant articles, the parent topic may be needed.

2. **Prioritize High-Quality Evidence**:
   - When appropriate, include terms that emphasize evidence quality, such as "systematic review," "meta-analysis," "guideline," or "consensus."
   - Include publication types like "practice guideline" or "review" to ensure relevance.
   - The goal is to return the best evidence available, not to have a high barrier.

3. **Streamline Terms**:
   - Avoid extraneous words or phrases that do not contribute to the search focus.
   - Include only essential and related terms.

4. **Leverage Boolean Operators**:
   - Combine MeSH terms and text words using **AND**, **OR**, and parentheses for logical grouping.
   - Use these operators to ensure inclusivity (OR) and specificity (AND).

5. **Incorporate Specific Examples**:
   - When relevant, include both broad categories and specific entities. For example:
     - Medications: (Anti-Bacterial Agents[MeSH Terms] OR antibiotics[Text Word] OR doxycycline[Text Word] OR amoxicillin[Text Word])
     - Conditions: (Hypertension[MeSH Terms] OR high blood pressure[Text Word])

6. **Exclude Overly Restrictive Features**:
   - Do not use quotation marks to narrow the results excessively.
   - Avoid unnecessary limits unless explicitly mentioned (e.g., age group, gender, etc.).

7. **Focus on Practical Application**:
   - Align the query with the practical intent of the user's question, such as treatment options, diagnostic approaches, or clinical decision-making.

### **Output Requirement**:
- Provide only the optimized PubMed query for each input question. Do not include additional commentary or extraneous information."""

    # Create the messages for the API call
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": query},
    ]

    # Get the LLM instance
    llm = get_llm_instance()
    if llm is None:
        return query  # Return original query if LLM initialization failed

    try:
        # Invoke the LLM
        response = llm.invoke(messages)
        optimized_query = (
            response.content if hasattr(response, "content") else str(response)
        )
        return optimized_query.strip()
    except Exception as e:
        print(f"Error optimizing PubMed search terms: {e}")
        return query  # Return original query if optimization failed


@st.cache_resource(ttl=3600)
def get_llm_instance():
    """
    Initializes and returns the AzureChatOpenAI instance.
    Uses Streamlit caching to avoid repeated initialization.
    """
    if not all([API_KEY, AZURE_ENDPOINT, AZURE_DEPLOYMENT]):
        st.error(
            "Azure OpenAI service credentials (API Key, Base URL, Deployment Name) are not configured. Please check secrets/environment variables."
        )
        return None
    try:
        return AzureChatOpenAI(
            azure_deployment=AZURE_DEPLOYMENT,
            api_version=API_VERSION,  # Or your preferred API version
            azure_endpoint=AZURE_ENDPOINT,
            api_key=API_KEY,
            temperature=0.0,  # Low temperature for more deterministic code generation
            max_tokens=4000,
            timeout=120,  # Increased timeout for potentially long code generations
            max_retries=3,
            seed=42,  # For reproducibility if supported
        )

    except Exception as e:
        st.error(f"Failed to initialize LLM: {e}")
        return None


def get_python_repl():
    """Initializes and returns a PythonREPL instance."""
    return PythonREPL()


# --- Helper: Clean plot directory ---
def clean_plot_directory(plot_path):
    """Removes old plot files from the specified directory."""
    if not os.path.exists(plot_path):
        os.makedirs(plot_path)  # Ensure directory exists
    image_exts = ["png", "jpg", "jpeg", "svg", "pdf"]
    for ext in image_exts:
        for img_file in glob.glob(os.path.join(plot_path, f"*.{ext}")):
            try:
                os.remove(img_file)
            except Exception as e:
                st.warning(f"Could not remove old plot file {img_file}: {e}")


# --- PubMed Search Functions ---

# Standard timeouts for network requests
AIOHTTP_TIMEOUT = aiohttp.ClientTimeout(
    total=60
)  # 60 seconds total for aiohttp requests


# Extract abstract text from PubMed XML data for a given PMID
async def extract_abstract_from_xml(xml_data: str, pmid: str) -> str:
    """
    Extract the abstract text from PubMed XML data for a specific PMID.

    Args:
        xml_data: The XML data from PubMed
        pmid: The PubMed ID to extract the abstract for

    Returns:
        The extracted abstract text or a message if no abstract is available
    """
    try:
        root = ET.fromstring(xml_data)
        for article in root.findall(".//PubmedArticle"):
            medline_citation = article.find("MedlineCitation")
            if medline_citation is not None:
                pmid_element = medline_citation.find("PMID")
                if pmid_element is not None and pmid_element.text == pmid:
                    abstract_element = medline_citation.find(".//Abstract")
                    if abstract_element is not None:
                        abstract_texts = []
                        for elem in abstract_element.findall("AbstractText"):
                            label = elem.get("Label")
                            # Get text content and normalize whitespace
                            text = ET.tostring(
                                elem, encoding="unicode", method="text"
                            ).strip()
                            # Replace any newlines with spaces to prevent one-character line issues
                            text = " ".join(text.splitlines())

                            if label is not None and label != "":
                                abstract_texts.append(f"{label}: {text}")
                            else:
                                abstract_texts.append(text)

                        # Join all abstract sections with proper spacing
                        return " ".join(abstract_texts).strip()
        return "No abstract available"
    except ET.ParseError:
        print(f"Error parsing XML for PMID {pmid}")
        return "Error extracting abstract"


# Fetch additional PubMed result IDs if needed
async def fetch_additional_results(
    session: aiohttp.ClientSession,
    search_query: str,
    max_results: int,
    current_count: int,
) -> List[str]:
    """
    Fetch additional PubMed result IDs if needed.

    Args:
        session: The aiohttp ClientSession to use for requests
        search_query: The search query to use
        max_results: The maximum number of results to fetch
        current_count: The current number of results already fetched

    Returns:
        A list of additional PubMed IDs
    """
    additional_needed = max_results - current_count
    url = (
        f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
        f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={additional_needed}&"
        f"api_key={st.secrets.get('pubmed_api_key', '')}"
    )
    try:
        async with session.get(url, timeout=AIOHTTP_TIMEOUT) as response:
            response.raise_for_status()
            data = await response.json()
            if "esearchresult" in data and isinstance(data["esearchresult"], dict):
                return data["esearchresult"].get("idlist", [])
            print(f"Unexpected data structure in fetch_additional_results: {data}")
            return []
    except (aiohttp.ClientError, asyncio.TimeoutError) as e:
        print(f"Error fetching additional results: {e}")
        return []


# Fetch PubMed article details and abstract XML data
async def fetch_article_details(
    session: aiohttp.ClientSession,
    id: str,
    details_url: str,
    abstracts_url: str,
    semaphore: asyncio.Semaphore,
) -> Tuple[str, Dict, str]:
    """
    Fetch PubMed article details and abstract XML data.

    Args:
        session: The aiohttp ClientSession to use for requests
        id: The PubMed ID
        details_url: The URL to fetch article details from
        abstracts_url: The URL to fetch abstract XML data from
        semaphore: A semaphore to limit concurrent requests

    Returns:
        A tuple of (id, details_data, abstracts_data)
    """
    async with semaphore:
        try:
            async with session.get(
                details_url, timeout=AIOHTTP_TIMEOUT
            ) as details_response:
                details_response.raise_for_status()
                details_data = await details_response.json()
            async with session.get(
                abstracts_url, timeout=AIOHTTP_TIMEOUT
            ) as abstracts_response:
                abstracts_response.raise_for_status()
                abstracts_data = await abstracts_response.text()
            return id, details_data, abstracts_data
        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            print(f"Error fetching article details for ID {id}: {e}")
            return id, {}, ""


# Retrieve and filter PubMed abstracts based on search terms and relevance
async def pubmed_abstracts(
    search_terms: str,
    search_type: str,
    max_results_param: int,
    initial_years_back_param: int,
    filter_relevance_param: bool,
    relevance_threshold_param: float,
    original_question: str,
) -> Tuple[List[Dict[str, str]], bool]:
    """
    Retrieve and filter PubMed abstracts based on search terms and relevance.

    Args:
        search_terms: The search terms to use
        search_type: The type of search to perform (not used in query)
        max_results_param: The maximum number of results to return
        initial_years_back_param: The number of years back to search initially
        filter_relevance_param: Whether to filter results by relevance
        relevance_threshold_param: The relevance threshold to use
        original_question: The original question for relevance filtering

    Returns:
        A tuple of (list of article dictionaries, flag indicating if older fallback articles were used)
    """

    async def _internal_search_logic(
        session_param: aiohttp.ClientSession,
        search_terms_for_helper: str,
        year_start_for_helper: int,
        year_end_for_helper: int,
        max_results_to_fetch: int,
        filter_relevance_for_helper: bool,
        relevance_threshold_for_helper: float,
        original_question_for_helper: str,
    ) -> List[Dict[str, str]]:
        search_query = f"{search_terms_for_helper}+AND+{year_start_for_helper}[PDAT]:{year_end_for_helper}[PDAT]"
        url = (
            f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?"
            f"db=pubmed&term={search_query}&sort=relevance&retmode=json&retmax={max_results_to_fetch}&"
            f"api_key={st.secrets.get('pubmed_api_key', '')}"
        )
        try:
            async with session_param.get(url, timeout=AIOHTTP_TIMEOUT) as response:
                response.raise_for_status()
                data = await response.json()
                if (
                    not isinstance(data.get("esearchresult"), dict)
                    or "count" not in data["esearchresult"]
                ):
                    st.error("Unexpected response format from PubMed API (esearch)")
                    return []
                ids = data["esearchresult"].get("idlist", [])
                print(
                    f"PubMed esearch for query '{search_query}' found {len(ids)} article IDs."
                )
                if not ids:
                    return []

            articles_data = []  # Renamed from 'articles' to avoid confusion
            semaphore = asyncio.Semaphore(10)
            tasks = []
            for id_str in ids:
                details_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?"
                    f"db=pubmed&id={id_str}&retmode=json&api_key={st.secrets.get('pubmed_api_key', '')}"
                )
                abstracts_url = (
                    f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?"
                    f"db=pubmed&id={id_str}&retmode=xml&rettype=abstract&api_key={st.secrets.get('pubmed_api_key', '')}"
                )
                tasks.append(
                    fetch_article_details(
                        session_param, id_str, details_url, abstracts_url, semaphore
                    )
                )
            results = await asyncio.gather(*tasks)

            processed_articles = []
            for id_str, details_data, abstracts_data in results:
                if "result" in details_data and str(id_str) in details_data["result"]:
                    article_detail = details_data["result"][str(id_str)]
                    year = article_detail["pubdate"].split(" ")[0]
                    if year.isdigit():
                        abstract = await extract_abstract_from_xml(
                            abstracts_data, id_str
                        )
                        article_url = f"https://pubmed.ncbi.nlm.nih.gov/{id_str}"
                        if abstract.strip() and abstract != "No abstract available":
                            processed_articles.append(
                                {
                                    "id": id_str,
                                    "title": article_detail["title"],
                                    "year": year,
                                    "abstract": abstract.strip(),
                                    "link": article_url,
                                }
                            )
                        else:
                            print(f"No valid abstract found for article ID {id_str}")

            if not processed_articles:
                return []

            if filter_relevance_for_helper:
                articles_prompt = "\n".join(
                    [
                        f"ID: {article['id']} - Title: {article['title']}"
                        for article in processed_articles
                    ]
                )
                messages = [
                    {
                        "role": "system",
                        "content": "You are an assistant evaluating the relevance of articles to a query. For each article provided, return a relevance score between 0.0 and 1.0 as a JSON object mapping the article's ID to its score, For example return only: {'12345': 0.9, '67890': 0.7}",
                    },
                    {
                        "role": "user",
                        "content": f"Query: {original_question_for_helper}\nArticles:\n{articles_prompt}\n\nReturn a JSON object without additional characters.",
                    },
                ]
                print("Filtering PubMed articles for question relevance...")
                try:
                    # Use Azure OpenAI for relevance filtering
                    llm = get_llm_instance()
                    response = llm.invoke(messages)
                    response_content = (
                        response.content
                        if hasattr(response, "content")
                        else str(response)
                    )
                    print(f"Relevance filtering response content: {response_content}")

                    if not response_content:
                        print(
                            "Empty response content received from relevance filtering."
                        )
                        relevance_scores = {}
                    else:
                        try:
                            # Clean up the response content to ensure it's valid JSON
                            # Remove any non-JSON text before or after the JSON object
                            json_start = response_content.find("{")
                            json_end = response_content.rfind("}") + 1

                            if json_start >= 0 and json_end > json_start:
                                json_content = response_content[json_start:json_end]
                                # Replace single quotes with double quotes if needed
                                json_content = json_content.replace("'", '"')
                                relevance_scores = json.loads(json_content)
                            else:
                                # If no JSON object is found, try to parse it as a dictionary-like string
                                # Convert the string representation of a dict to an actual dict
                                if ":" in response_content and "," in response_content:
                                    relevance_scores = {}
                                    # Extract key-value pairs using regex
                                    import re

                                    pairs = re.findall(
                                        r'[\'"]?(\w+)[\'"]?\s*:\s*(\d+\.?\d*)',
                                        response_content,
                                    )
                                    for key, value in pairs:
                                        relevance_scores[key] = float(value)
                                else:
                                    print("No valid JSON structure found in response")
                                    relevance_scores = {}
                        except json.JSONDecodeError as e:
                            print(f"Error parsing JSON from relevance filtering: {e}")
                            print("Attempting alternative parsing methods...")

                            try:
                                # Try eval as a last resort (with safety checks)
                                if "{" in response_content and "}" in response_content:
                                    # Only evaluate if it looks like a dictionary
                                    import ast

                                    relevance_scores = ast.literal_eval(
                                        response_content
                                    )
                                else:
                                    relevance_scores = {}
                            except (SyntaxError, ValueError) as eval_error:
                                print(f"Alternative parsing also failed: {eval_error}")
                                relevance_scores = {}

                    final_filtered_articles = [
                        article
                        for article in processed_articles
                        if float(relevance_scores.get(str(article["id"]), 0))
                        >= relevance_threshold_for_helper
                    ]
                    print(
                        f"Found {len(final_filtered_articles)} relevant articles after filtering."
                    )
                    return final_filtered_articles
                except Exception as e:
                    print(f"Error during relevance filtering: {e}")
                    return processed_articles  # Fallback to unfiltered if error
            else:  # No relevance filtering
                return processed_articles

        except (aiohttp.ClientError, asyncio.TimeoutError) as e:
            print(f"Error fetching PubMed articles for query '{search_query}': {e}")
            return []

    # --- Main logic for pubmed_abstracts ---
    final_articles_to_return = []
    used_older_fallback_articles = False
    current_year = datetime.now().year

    async with aiohttp.ClientSession() as session:
        # Initial search
        start_year_initial = current_year - initial_years_back_param
        print(
            f"Performing initial PubMed search for '{search_terms}' from {start_year_initial} to {current_year} (last {initial_years_back_param} years)."
        )
        initial_articles_found = await _internal_search_logic(
            session,
            search_terms,
            start_year_initial,
            current_year,
            max_results_param,
            filter_relevance_param,
            relevance_threshold_param,
            original_question,
        )

        if initial_articles_found:
            final_articles_to_return = initial_articles_found
            print(
                f"Initial PubMed search found {len(final_articles_to_return)} relevant articles."
            )
        else:  # Initial search found nothing
            if initial_years_back_param < 15:  # Max fallback depth is 15 years
                print(
                    f"Initial PubMed search (last {initial_years_back_param} years) yielded no relevant articles. Extending search to 15 years."
                )
                start_year_extended = current_year - 15  # Fixed 15 years

                extended_articles_found = await _internal_search_logic(
                    session,
                    search_terms,
                    start_year_extended,
                    current_year,
                    max_results_param,
                    filter_relevance_param,
                    relevance_threshold_param,
                    original_question,
                )

                if extended_articles_found:
                    final_articles_to_return = extended_articles_found
                    used_older_fallback_articles = True
                    print(
                        f"Extended PubMed search (last 15 years) found {len(final_articles_to_return)} relevant articles."
                    )
                else:
                    print(
                        "Extended PubMed search (last 15 years) also found no relevant articles."
                    )
            else:  # Initial search found nothing, and initial_years_back_param was already >= 15
                print(
                    f"Initial PubMed search (last {initial_years_back_param} years) yielded no relevant articles. Not extending further as search already covered {initial_years_back_param} years."
                )

    # Ensure we return only up to max_results_param
    print(
        f"Total articles to be added to the database: {len(final_articles_to_return[:max_results_param])}"
    )
    return final_articles_to_return[:max_results_param], used_older_fallback_articles


# --- Core Analysis Functions (Continued from Part 1) ---


def get_code_from_llm(
    llm,
    question,
    dfs_snapshot,
    iteration,
    max_iters,
    prev_code,
    prev_output,
    prev_error,
    compare_web=False,
    doc_available=False,
    research_mode=False,
):
    """
    Generates Python code using the LLM based on the question and previous iteration results.
    Handles both general questions and dataframe-specific questions.
    """
    df_info_str = "No dataframes are currently loaded. Focus on answering the general text-based question."
    dataframes_info = {}

    # Check if RAG search is available
    search_available, search_error_msg = (
        get_rag_search_available()
    )  # Renamed search_error to avoid conflict
    medical_search_context_for_llm = (
        ""  # To hold info about what searches were done for medical Qs
    )
    medical_search_results = None  # Initialize to fix F821 error

    # Determine if this is a medical question
    # Import here to avoid circular imports if not already at top level
    from rag_utils import (
        is_medical_topic,
        get_embeddings_model,
        google_search,
        scrape_webpage,
    )
    from langchain_core.documents import Document
    from langchain_community.vectorstores import FAISS
    from urllib.parse import urlparse  # For document naming from URLs

    is_medical_question = is_medical_topic(question)
    st.session_state["is_medical_question"] = is_medical_question
    print(f"Question identified as medical: {is_medical_question}")

    # Initialize medical_context_for_prompt here to ensure it's always defined
    medical_context_for_prompt = ""

    if is_medical_question and search_available:
        st.sidebar.info(
            "Processing as a medical question: gathering information from PubMed and Web to build a comprehensive knowledge base for this query."
        )
        all_documents_for_vs = []  # Accumulator for documents from all sources
        embeddings = get_embeddings_model()

        if not embeddings:
            st.error(
                "Embeddings model not available. Cannot build knowledge base for medical question."
            )
            medical_search_context_for_llm = (
                "Error: Embeddings model unavailable. Cannot perform RAG."
            )
        else:
            # --- PubMed Search and Processing ---
            with st.spinner("Performing PubMed search and processing..."):
                cutting_edge_value = st.session_state.get("cutting_edge", False)
                pubmed_search_terms = optimize_pubmed_search_terms(
                    question, cutting_edge_value
                )
                st.session_state.pubmed_search_terms = pubmed_search_terms

                articles_list, older_fetched_flag = asyncio.run(
                    pubmed_abstracts(
                        search_terms=pubmed_search_terms,
                        search_type="all",
                        max_results_param=st.session_state.get("max_results", 6),
                        initial_years_back_param=st.session_state.get("years_back", 4),
                        filter_relevance_param=st.session_state.get(
                            "filter_relevance", True
                        ),
                        relevance_threshold_param=st.session_state.get(
                            "relevance_threshold", 0.7
                        ),  # Using sidebar value
                        original_question=question,
                    )
                )
                st.session_state.articles = articles_list  # For display later
                st.session_state.older_pubmed_articles_alert = older_fetched_flag

                pubmed_docs_for_vs_temp = []
                if articles_list:
                    # Use web_chunk_size and web_chunk_overlap for PubMed content
                    pubmed_chunk_size = st.session_state.get("web_chunk_size", 1500)
                    pubmed_chunk_overlap = st.session_state.get("web_chunk_overlap", 200)
                    for article in articles_list:
                        # Using abstract as primary content. Full-text scraping from PubMed is complex.
                        content = f"Source: PubMed Article\nTitle: {article['title']}\nYear: {article['year']}\nLink: {article['link']}\nAbstract: {article['abstract']}"
                        # Chunk the content using the web/PubMed chunk size/overlap
                        chunks = create_document_chunks(
                            content,
                            chunk_size=pubmed_chunk_size,
                            chunk_overlap=pubmed_chunk_overlap,
                        )
                        for i, chunk in enumerate(chunks):
                            metadata = DocumentMetadata(
                                source=article["link"],
                                title=article["title"],
                                document_name=f"PubMed_{article['id']}",
                                document_type="pubmed_article",
                                is_scraped=False,
                                chunk=i,
                            ).model_dump()
                            pubmed_docs_for_vs_temp.append(
                                Document(page_content=chunk, metadata=metadata)
                            )
                            # --- Add to retrieved_chunks for UI/RAGAS ---
                            if "retrieved_chunks" not in st.session_state:
                                st.session_state.retrieved_chunks = []
                            if chunk not in st.session_state.retrieved_chunks:
                                st.session_state.retrieved_chunks.append(chunk)
                    all_documents_for_vs.extend(pubmed_docs_for_vs_temp)
                    print(
                        f"Prepared {len(pubmed_docs_for_vs_temp)} documents from PubMed."
                    )
                    medical_search_context_for_llm += f"\n- Searched PubMed with terms: '{pubmed_search_terms}'. Found {len(pubmed_docs_for_vs_temp)} relevant abstracts."

            # --- Web Search and Processing ---
            with st.spinner("Performing web search and processing relevant content..."):
                current_date_str = datetime.now().strftime("%B %Y")
                web_search_query = (
                    f"{question} medical context {current_date_str}"  # General query
                )

                num_web_results_to_process = st.session_state.get("rag_num_results", 5)
                raw_web_results = google_search(
                    web_search_query,
                    num_results=num_web_results_to_process,
                    scrape_content=st.session_state.get("rag_scrape_content", True),
                )  # Use the scrape_content setting from session state

                web_docs_for_vs_temp = []
                if raw_web_results:
                    # Use web_chunk_size and web_chunk_overlap for web content
                    web_chunk_size = st.session_state.get("web_chunk_size", 1500)
                    web_chunk_overlap = st.session_state.get("web_chunk_overlap", 200)
                    processed_urls_count = 0
                    for result in raw_web_results:
                        if result.link:
                            print(f"Scraping web content from: {result.link}")
                            scraped_content = scrape_webpage(
                                result.link
                            )  # from rag_utils
                            if (
                                scraped_content
                                and not scraped_content.startswith("Error:")
                                and not scraped_content.startswith("Skipped:")
                                and len(scraped_content) > 150
                            ):
                                content = f"Source: Web Page\nTitle: {result.title}\nURL: {result.link}\nContent: {scraped_content}"
                                # Chunk the content using the web/PubMed chunk size/overlap
                                chunks = create_document_chunks(
                                    content,
                                    chunk_size=web_chunk_size,
                                    chunk_overlap=web_chunk_overlap,
                                )
                                for i, chunk in enumerate(chunks):
                                    metadata = DocumentMetadata(
                                        source=result.link,
                                        title=result.title,
                                        document_name=urlparse(
                                            result.link
                                        ).netloc,  # Use domain as name
                                        document_type="web_page",
                                        is_scraped=True,
                                        chunk=i,
                                    ).model_dump()
                                    web_docs_for_vs_temp.append(
                                        Document(page_content=chunk, metadata=metadata)
                                    )
                                    # --- Add to retrieved_chunks for UI/RAGAS ---
                                    if "retrieved_chunks" not in st.session_state:
                                        st.session_state.retrieved_chunks = []
                                    if chunk not in st.session_state.retrieved_chunks:
                                        st.session_state.retrieved_chunks.append(chunk)
                                processed_urls_count += 1
                            else:
                                print(
                                    f"Skipping or failed to scrape useful content from: {result.link}"
                                )
                    all_documents_for_vs.extend(web_docs_for_vs_temp)
                    print(
                        f"Prepared {len(web_docs_for_vs_temp)} documents from {processed_urls_count} scraped web pages."
                    )
                    medical_search_context_for_llm += f"\n- Searched Web for '{web_search_query}'. Added content from {len(web_docs_for_vs_temp)} pages."

            # --- Update/Create Vector Store ---
            if all_documents_for_vs:  # If we have new documents from PubMed or Web
                with st.spinner("Updating knowledge base with PubMed/Web findings..."):
                    if st.session_state.document_vector_store is not None:
                        # Add new documents to the existing store (which might contain uploaded files)
                        st.session_state.document_vector_store.add_documents(
                            all_documents_for_vs
                        )
                        st.success(
                            f"Added {len(all_documents_for_vs)} PubMed/Web documents to the existing knowledge base."
                        )
                        medical_search_context_for_llm += "\n- These new findings have been added to the existing knowledge base (which may include user-uploaded documents)."
                    else:
                        # Create a new store if one doesn't exist
                        st.session_state.document_vector_store = FAISS.from_documents(
                            all_documents_for_vs, embeddings
                        )
                        st.success(
                            f"Knowledge base created with {len(all_documents_for_vs)} PubMed/Web documents."
                        )
                        medical_search_context_for_llm += "\n- A new knowledge base has been created with these PubMed/Web findings."
            elif st.session_state.document_vector_store is not None:
                st.info(
                    "No new PubMed or Web documents were added to the knowledge base for this query, but using existing uploaded documents."
                )
                medical_search_context_for_llm += "\n- No new PubMed or Web documents found. Will rely on previously uploaded documents if any."
            else:
                st.info(
                    "No new PubMed or Web documents were found, and no documents were previously uploaded. The knowledge base is empty."
                )
                medical_search_context_for_llm += "\n- No PubMed, Web, or uploaded documents available in the knowledge base."

    # Fallback for non-medical questions or if RAG search is not available for medical
    search_likely_needed = False  # Reset for non-medical or if medical RAG failed
    if (
        not is_medical_question and search_available
    ):  # Standard search check for non-medical
        # Import here to avoid circular imports
        import search_utils  # Ensure it's imported for this path too

        if compare_web:  # from button
            search_likely_needed = True
        elif st.session_state.document_vector_store is None:  # No uploaded docs
            search_likely_needed = search_utils.is_search_needed(question)
        # else: (docs available, not comparing) search_likely_needed remains False

    search_tool_info = ""
    doc_vector_store_available = (
        st.session_state.get("document_vector_store") is not None
    )

    if search_available:
        files_uploaded = bool(st.session_state.dataframes) or bool(
            st.session_state.document_texts
        )
        augment_with_pubmed_web = st.session_state.get("augment_with_pubmed_web", False)
        if files_uploaded and not augment_with_pubmed_web:
            # User uploaded files and did NOT request PubMed/Web augmentation: restrict to uploaded files only
            if doc_vector_store_available:
                search_tool_info = """
You have uploaded data files and/or documents. The analysis will be confined ONLY to your uploaded files. PubMed and web search will NOT be used unless you enable it in the sidebar.

**Available Tools:**
- Python REPL for code execution and data analysis
- Uploaded CSV/Excel data files (as pandas DataFrames)
- Uploaded PDF/Word documents (as searchable chunks)
- Document search via the `search_utils.web_search` function (searches only your uploaded documents)

**Instructions:**
- You MUST base your answer solely on the information found in the uploaded files.
- If you cannot answer the question with the available files, clearly state this.
- Do NOT use PubMed or web search unless the user enables it in the sidebar.
- Use the `web_search` function to search your uploaded documents for relevant information.

Example usage:
```python
from search_utils import web_search
results = web_search("your specific search query here")
print(results)
```
"""
            else:
                search_tool_info = """
No data files or documents are currently uploaded. Please upload files to enable document-based analysis, or enable PubMed/Web search in the sidebar.
"""
        else:
            # Either no files uploaded (so use all tools), or user explicitly enabled PubMed/Web augmentation
            if is_medical_question:
                # Specific prompt for medical questions after data gathering
                search_tool_info = f"""
This is a medical/health question. The knowledge base has been prepared for you.
Summary of information gathering: {medical_search_context_for_llm}

Your primary task is to synthesize an answer based on the information retrieved from this comprehensive knowledge base.
To query the knowledge base, use the `search_utils.web_search` function.
You MUST try multiple variations of the question to retrieve the most relevant chunks. For example:
- The original question: "{question}"
- A more general phrasing of the question.
- Queries focusing on specific sub-aspects or keywords from the question.

Example of querying the knowledge base:
```python
# Import the search tool
from search_utils import web_search

# Query with the original question
results1 = search_utils.web_search("{question}")
print("Results for original question:\\n" + results1)

# Query with a variation (e.g., broader or more specific)
results2 = search_utils.web_search("alternative or focused phrasing of '{question}'")
print("Results for alternative phrasing:\\n" + results2)

# After gathering information from multiple queries, synthesize your answer.
# Your final answer should be based on the content of these search results.
# If you have relevant medical knowledge that is NOT in the search results,
# you may include it ONLY in a separate section at the end of your response
# titled "From the AI's Own Knowledge:".
```

IMPORTANT:
- The `web_search` function returns a pre-formatted STRING. Print it to see the retrieved chunks.
- Your final response should integrate information from these chunks.
- Prioritize recent, high-quality evidence (systematic reviews, meta-analyses, clinical guidelines) if discernible from the chunks.
- At the end of your answer, include a "References for further reading" section listing the source URLs/PubMed links found in the retrieved chunks.
"""
            elif (
                doc_vector_store_available and not compare_web
            ):  # Non-medical, docs uploaded, not comparing with web
                search_tool_info = """
You have access to a RAG (Retrieval Augmented Generation) system that can search uploaded documents.

IMPORTANT: The user has uploaded documents that MUST be used to answer the question. You MUST constrain your response to ONLY information found in these documents.

To use the search tool, include code like this:
```python
# Import the search tool
from search_utils import web_search

# Search uploaded documents - ALWAYS include this code
# Try multiple search queries with different phrasings to maximize chances of finding relevant information
search_results1 = web_search("your specific search query here")
print(search_results1)

# Try a more general version of the query to catch broader matches
search_results2 = web_search("broader terms related to the topic")
print(search_results2)
```

IMPORTANT: 
1. You MUST include the web_search code in your response to search the documents.
2. Try multiple search queries with different phrasings to maximize the chance of finding relevant information.
3. The web_search function returns a pre-formatted STRING with the search results, not a list or dictionary.
4. Do not try to parse or process the results - just print them directly.
5. If you see "No relevant documents found" for one query, try different search terms.
6. Documents HAVE been uploaded to the system. There are document chunks available for searching.

Your response MUST be based ONLY on information found in the uploaded documents. If after multiple search attempts you still can't find relevant information in the documents, clearly state this.

If you have relevant knowledge that is NOT in the documents, you may include it ONLY in a separate section at the end of your response titled "From the AI's Own Knowledge:".
"""
            elif search_likely_needed or compare_web:
                search_tool_info = """
You have access to a RAG (Retrieval Augmented Generation) system that can be used to look up current information or facts that might not be in your training data.
Based on the question, it appears that web search would be HIGHLY BENEFICIAL for providing an accurate and up-to-date answer.

IMPORTANT: You MUST use the search tool for this question as it likely requires current information.

To use the search tool, include code like this:
```python
# Import the search tool
from search_utils import web_search

# Use the search tool to get information - it returns a FORMATTED STRING
search_results = web_search("your specific search query here")
# Simply print the results directly - DO NOT try to parse them as a list or dictionary
print(search_results)
```

For example, if asked "What are the latest developments in AI?", your code should include:
```python
search_results = web_search("latest developments in artificial intelligence")
print(search_results)  # Already formatted, just print directly
```

IMPORTANT: The web_search function returns a pre-formatted STRING with the search results, not a list or dictionary.
Do not try to parse or process the results - just print them directly.

For questions about news or current events, ALWAYS use the search tool first before attempting to answer.
"""
            else:
                search_tool_info = """
You have access to a RAG (Retrieval Augmented Generation) system that can be used to look up current information or facts that might not be in your training data.
Use this tool when the question requires up-to-date information, specific facts, or data that you're uncertain about.

To use the search tool, include code like this:
```python
# Import the search tool
from search_utils import web_search

# Use the search tool to get information - it returns a FORMATTED STRING
search_results = web_search("your specific search query here")
# Simply print the results directly - DO NOT try to parse them as a list or dictionary
print(search_results)
```

IMPORTANT: The web_search function returns a pre-formatted STRING with the search results, not a list or dictionary.
Do not try to parse or process the results - just print them directly.

Only use the search tool if you believe the question requires current information or facts that might not be in your training data.
"""

    # Prepare dataframe context if dataframes are loaded
    if dfs_snapshot and any(not df.empty for df in dfs_snapshot.values()):
        # Build information about each dataframe
        for df_name, df in dfs_snapshot.items():
            if df is not None and not df.empty:
                df_cols_str = ", ".join(df.columns)
                df_head_str = df.head().to_string()
                dataframes_info[df_name] = {
                    "columns": list(df.columns),
                    "head": df_head_str,
                    "shape": df.shape,
                    "dtypes": {col: str(df[col].dtype) for col in df.columns},
                }

        # Create a comprehensive dataframe context string
        df_info_parts = ["The user has loaded the following pandas DataFrames:"]

        for df_name, info in dataframes_info.items():
            df_info_parts.append(f"\n## DataFrame: `{df_name}`")
            df_info_parts.append(f"Shape: {info['shape']}")
            df_info_parts.append(f"Columns: {', '.join(info['columns'])}")
            df_info_parts.append(f"First 5 rows of `{df_name}`:")
            df_info_parts.append(f"{info['head']}\n")

        df_info_parts.append("""
For each dataframe, you have access to two versions:
1. `{df_name}` - This is your working copy of the dataframe. It reflects any modifications made in previous steps of this analysis. You can modify this dataframe as needed.
2. `original_{df_name}` - This is the initial, unmodified dataframe (read-only reference). Use this if you need to start a calculation from the original data state.

Based on the user's question, determine which dataframe(s) to use for your analysis. If the question doesn't specify a dataframe but requires one, use the most appropriate dataframe based on the column names and data types.

If the user refers to a column name using a synonym or in a different case, always match it to the correct intended column name in the appropriate dataframe, ignoring case.

Before performing any analysis that requires numeric data (such as correlation heatmaps, PCA, or regression), always check for categorical columns (object dtype or string values):
-   If a categorical column has exactly 2 unique values, convert it to numeric by mapping the most common value to 0 and the least common value to 1. Use the `safe_map_categorical(series, mapping)` function for this conversion. Print a message indicating which columns were converted and how.
-   If a categorical column has more than 2 unique values, use one-hot encoding (e.g., `df = pd.get_dummies(df, columns=['your_column_name'], prefix='your_column_name')`) to create additional columns as needed. Print a message indicating which columns were one-hot encoded.
-   Always check for and handle NaN values in categorical columns before mapping or encoding.

**Important:** The unique values for categorical columns in the current dataframes might be printed in the 'Previous Output/Error' section for your reference. Use this information to correctly identify and handle categorical values.""")

        df_info_str = "\n".join(df_info_parts)
    else:  # No dataframe context
        df_info_str = """The user has not uploaded any dataframes. Your task is to answer their general text-based question.
If the question implies data analysis but no data is provided, state that you need data.
Otherwise, answer the question directly. You can use Python for calculations or general knowledge tasks if helpful, but do not assume any dataframes exist."""

    medical_context_for_prompt = ""  # Initialize here to ensure it's always defined

    # --- Base Prompt Construction (Iteration 1) ---
    if iteration == 1:
        # Standard mode - original prompt construction
        # medical_context and medical_references handling needs to align with the new medical flow.
        # search_tool_info is now set correctly for medical questions above.
        # medical_context_for_prompt is also set.
        medical_references = ""  # References are now expected to be extracted by LLM from retrieved chunks.

        # medical_context_for_prompt is already prepared earlier based on is_medical_question
        # and whether medical_search_results (old flow) or medical_search_context_for_llm (new flow) exists.
        # For the new medical flow, search_tool_info is comprehensive.

        prompt = f"""You are an expert Python data analyst and helpful assistant.
The user has asked the following question:
"{question}"

{medical_context_for_prompt}
{df_info_str}
{search_tool_info}

Your task is to write Python code to answer the question or perform the requested analysis.

**Instructions for Code Generation:**
1.  **Imports:** ALWAYS start your Python code block with:
    ```python
    import matplotlib.pyplot as plt
    import seaborn as sns
    import numpy as np
    import pandas as pd
    import streamlit as st  # ALWAYS import streamlit
    import time # Import time for sleep
    ```
2.  **Dataframe Handling (if `df` exists):**
    *   If modifying `df` (e.g., filtering, transforming), the changes will persist for subsequent iterations.
    *   Use `original_df` for calculations requiring the initial state of the uploaded data.
3.  **Plotting & Illustrative Images:**
    *   If a plot is needed for data analysis, or if a key concept/point can be effectively illustrated with an image generated by Python (e.g., a simple diagram, a conceptual chart, a mathematical function plot), create it.
    *   **Use Streamlit's native display capabilities - no need to save files to disk.**
    *   For matplotlib plots, simply create the figure and then display it directly with:
        ```python
        # After creating your plot
        plt.tight_layout()
        st.pyplot(fig)  # Pass the figure object
        ```
    *   Alternatively, you can just use `plt.tight_layout()` followed by `st.pyplot()` to display the current figure.
    *   **DO NOT use `plt.savefig()` or `plt.show()` - use Streamlit's display functions instead.**
    *   If the question is ambiguous about plotting, make a reasonable choice for a helpful visualization. If no plot is relevant, state so in a comment within the code.
4.  **Output:**
    *   For text output, you can use BOTH `print()` AND Streamlit's display functions:
        - For simple text: `st.write("Your text here")`
        - For important text: `st.info("Important information")`
        - For warnings: `st.warning("Warning message")`
        - For errors: `st.error("Error message")`
        - For success messages: `st.success("Success message")`
    *   For tables and dataframes, use `st.dataframe(your_dataframe)` to display them interactively.
    *   For code snippets or raw text, use `st.code(your_code_or_text)`.
5.  **Error Handling:** If the question is unanswerable with the provided information (or lack of dataframe), or if an operation is impossible, raise an `Exception` with a clear, helpful message within the code.
6.  **Focus:** If no dataframe is loaded, focus on answering the general question. Python can be used for calculations or demonstrations.
7.  **Document Constraints:** If documents are uploaded, ONLY use information from those documents. If you have relevant knowledge not in the documents, include it ONLY in a section at the end titled "From the AI's Own Knowledge:".
8.  **IMPORTANT:** Documents HAVE been uploaded to the system. There are document chunks available for searching.

{("For medical/health questions, always include a 'References for further reading' section at the end of your answer, using ONLY the following URLs from the search results:" + chr(10) + medical_references) if is_medical_question else ""}

Respond ONLY with a single block of valid Python code. Do not include any surrounding text like "Here is the code:".
"""
    # --- Refinement Prompt Construction (Iterations > 1) ---
    else:
        error_info = (
            f"\nThe previous code (iteration {iteration - 1}) resulted in this error:\n```\n{prev_error}\n```"
            if prev_error
            else ""
        )

        # medical_context_for_prompt and search_tool_info are determined before this 'else' block.
        # medical_references are expected to be handled by the LLM based on retrieved chunks.
        medical_references = (
            ""  # Reset for safety, though not directly used in prompt string here.
        )

        if research_mode:
            # Research Assistant Mode refinement prompt
            prompt = f"""You are an expert Research Assistant with access to powerful tools, in a multi-turn refinement process.
The researcher's original question was:
"{question}"

{medical_context_for_prompt}
{df_info_str}

This is **iteration {iteration} of {max_iters}**.

**Previous attempt (iteration {iteration - 1}):**
Previous Code:
```python
{prev_code}
```
Previous Output/Error:
```
{prev_output if prev_output else "No output was generated."}
```
{error_info}

{search_tool_info}

**Your Task for Iteration {iteration}:**
Continue your research process by reviewing the previous attempt and improving the Python code to better answer the researcher's question, or to fix any errors.

**Research Process - Current Stage:**
1. **ASSESS PROGRESS**: Evaluate what you've learned so far and what gaps remain in answering the research question.
2. **REFINE APPROACH**: Based on your assessment, refine your research strategy for this iteration.
3. **EXECUTE NEXT STEPS**: Implement the next steps in your research plan, building on previous findings.
4. **SYNTHESIZE**: Continue to integrate new information with previous findings.

**Key Considerations for Improvement:**
*   **Research Depth:** Are you exploring the question thoroughly? Consider multiple angles and sources.
*   **Error Resolution:** If there was an error, diagnose and fix it.
*   **Tool Selection:** Are you using the most appropriate tools for each part of your research?
*   **Information Integration:** How can you better combine information from different sources?
*   **Visualization:** If using visualizations, ensure they effectively communicate key findings.
*   **Dataframe Analysis:** If analyzing data, ensure your methods are appropriate for the research question.
*   **Document Utilization:** If documents are uploaded, ensure you're extracting relevant information effectively.
*   **Interactive Presentation:** Use Streamlit's display functions to present your findings in an interactive and visually appealing way.

**Instructions for Code Generation (Reminder):**
1.  **Imports:** ALWAYS start your Python code block with the standard imports (plt, sns, np, pd, st, time).
2.  **Plotting & Illustrative Images:** If a plot is needed for data analysis, or if a key concept/point can be effectively illustrated with an image generated by Python (e.g., a simple diagram, a conceptual chart, a mathematical function plot), create it. **Use Streamlit's native display capabilities - no need to save files to disk.** For matplotlib plots, simply create the figure and then display it directly with `st.pyplot(fig)` or just `st.pyplot()` for the current figure. Always call `plt.tight_layout()` before displaying.
4.  **Output:** Use both `print()` and Streamlit display functions (st.write, st.dataframe, etc.) for better presentation of results.
5.  **Research Progress:** Include comments that explain your research progress and next steps.

{("For medical/health questions, include a 'References for further reading' section at the end of your answer, using ONLY the following URLs from the search results:" + chr(10) + medical_references) if is_medical_question else ""}

Respond ONLY with a single block of improved Python code. Do not include any surrounding text.
"""
        else:
            # Standard mode refinement prompt
            prompt = f"""You are an expert Python data analyst and helpful assistant, in a multi-turn refinement process.
The user's original question was:
"{question}"

{medical_context_for_prompt}
{df_info_str}

This is **iteration {iteration} of {max_iters}**.

**Previous attempt (iteration {iteration - 1}):**
Previous Code:
```python
{prev_code}
```
Previous Output/Error:
```
{prev_output if prev_output else "No output was generated."}
```
{error_info}

{search_tool_info}

**Your Task for Iteration {iteration}:**
Review the previous attempt and improve the Python code to better and more completely answer the user's question, or to fix any errors.

**Key Considerations for Improvement:**
*   **Correctness & Completeness:** Does the previous output fully and accurately answer the question? If not, how can the code be modified?
*   **Error Resolution:** If there was an error, diagnose and fix it.
*   **Plotting & Illustrative Images:** If a plot is needed for data analysis, or if a key concept/point can be effectively illustrated with an image generated by Python (e.g., a simple diagram, a conceptual chart, a mathematical function plot), create it. **Use Streamlit's native display capabilities - no need to save files to disk.** For matplotlib plots, simply create the figure and then display it directly with `st.pyplot(fig)` or just `st.pyplot()` for the current figure. Always call `plt.tight_layout()` before displaying.
*   **Clarity of Output:** Is the printed output clear and easy to understand? Consider using Streamlit's display functions (st.write, st.info, st.warning, st.error, st.success) for better presentation.
*   **Dataframe Modifications (if `df` exists):** Remember that changes to `df` persist. If you need to revert to the original data for a step, use `original_df`.
*   **Categorical Encoding (if `df` exists):** Ensure `safe_map_categorical` or `pd.get_dummies` is used appropriately if new categorical columns are being processed.
*   **Document Constraints:** If documents are uploaded, ONLY use information from those documents. If you have relevant knowledge not in the documents, include it ONLY in a section at the end titled "From the AI's Own Knowledge:".

**Instructions for Code Generation (Reminder):**
1.  **Imports:** ALWAYS start your Python code block with the standard imports (plt, sns, np, pd, st, time).
2.  **Plotting:** Save plots to the directory specified by the `plot_output_dir` variable. You MUST use an f-string like this: `plt.savefig(f'{{plot_output_dir}}/plot_iter{{iteration}}.png')`. Call `time.sleep(0.1)` immediately after `plt.savefig()`, and then call `plt.close()`. No `plt.show()`.
3.  **Output:** Use both `print()` and Streamlit display functions (st.write, st.dataframe, etc.) for better presentation of results.

{("For medical/health questions, always include a 'References for further reading' section at the end of your answer, using ONLY the following URLs from the search results:" + chr(10) + medical_references) if is_medical_question else ""}

Respond ONLY with a single block of improved Python code. Do not include any surrounding text.
"""

    if llm is None:  # Safety check if LLM failed to initialize
        st.error("LLM is not available. Cannot generate code.")
        return ""  # Return empty string or specific error code

    try:
        with st.spinner(f"🤖 LLM is thinking... (Iteration {iteration}/{max_iters})"):
            response = llm.invoke(prompt)
        code_content = (
            response.content if hasattr(response, "content") else str(response)
        )
        # Clean the code: remove markdown backticks and "python" language specifier
        cleaned_code = re.sub(
            r"^```python\s*|^```\s*|\s*```$",
            "",
            code_content,
            flags=re.MULTILINE | re.DOTALL,
        ).strip()
        return cleaned_code
    except Exception as e:
        st.error(f"Error communicating with LLM: {e}")
        return f"# LLM communication error: {e}\nraise Exception('LLM communication error')"


def run_code_in_repl(
    repl, code_to_run, dfs_current_state, dfs_original_state, iteration=1
):
    """
    Executes the provided Python code in the REPL environment.
    Captures output, errors, and new plot images.
    Manages all dataframes in the REPL's global scope.
    Returns the textual output, error message (if any), list of new image paths,
    and the state of all dataframes after execution.
    """
    output_capture = io.StringIO()
    error_occurred_msg = None
    new_image_data_b64 = []  # Changed to store base64 encoded image data
    dfs_after_execution = (
        {name: df.copy() for name, df in dfs_current_state.items()}
        if dfs_current_state
        else {}
    )
    retrieved_chunks = []  # To store any retrieved chunks from this execution

    # Check if RAG search is available and add it to globals if it is
    search_module = None
    search_available, _ = get_rag_search_available()
    if search_available:
        import search_utils

        search_module = search_utils

    # Prepare REPL globals with all dataframes
    current_globals = {
        "plt": plt,
        "sns": sns,
        "np": np,
        "pd": pd,
        "st": st,  # Make streamlit available for direct display
        "Image": Image,  # Make Pillow Image available
        "__file__": "streamlit_app.py",  # Mock __file__ for some libraries
        "__name__": "__main__",  # Mock __name__
        "categorical_mappings": st.session_state.get(
            "categorical_mappings", {}
        ).copy(),  # Pass current mappings
        "search_utils": search_module,  # Add search module if available
        "time": time,  # Make time module available for sleep
        "iteration": iteration,  # Add current iteration number for plot naming
        "io": io,  # Make io module available
        "BytesIO": BytesIO,  # Make BytesIO available
    }

    # Add all current dataframes to globals
    for df_name, df in dfs_current_state.items():
        if df is not None:
            current_globals[df_name] = df.copy()

    # Add all original dataframes to globals with 'original_' prefix
    for df_name, df in dfs_original_state.items():
        if df is not None:
            current_globals[f"original_{df_name}"] = df.copy()

    # Helper function for safe categorical mapping to be injected into the REPL
    # Helper function for safe categorical mapping to be injected into the REPL
    safe_mapping_helper_code = """
# Patch the built-in print function to handle search results and retrieved documents. Always include:

original_print = print
import os
import base64

def patched_print(*args, **kwargs):
    # Check if this is a search result or document retrieval output
    text = " ".join(str(arg) for arg in args)
    if text.startswith("Search results for:") or text.startswith("Retrieved documents for query:") or text.startswith("DOCUMENTS FOUND:"):
        # Still print to the output capture for processing, but don't display in Streamlit
        original_print(*args, **kwargs)
    else:
        # Normal print behavior
        original_print(*args, **kwargs)
        
# Replace the built-in print function with our patched version
print = patched_print

# === Streamlit-native plotting helper functions ===
import streamlit as st
import matplotlib.pyplot as plt
from io import BytesIO
import os
import base64

def display_matplotlib_figure(fig=None, caption=None):
    '''
    Displays a matplotlib figure directly in Streamlit without saving to disk
    and closes the figure.
    
    Args:
        fig: The matplotlib figure to display (defaults to current figure).
        caption: Optional caption for the image.
    
    Returns:
        None
    '''
    if fig is None:
        fig = plt.gcf()
    
    # Prevent creating an empty figure if no plot exists
    if not fig.axes and not fig.get_children(): # Check if figure has axes or other artists
        if plt.get_fignums(): # Check if there are any figures at all
            original_print(f"Info: Attempted to display Figure {fig.number}, but it appears to be empty. Skipping display.")
        # else: # No figures exist at all, gcf() might have created a new empty one.
            # original_print(f"Info: No active Matplotlib figure to display.")
        plt.close(fig) # Close the potentially empty figure created by gcf()
        return

    buf = BytesIO()
    try:
        # Apply tight_layout if possible, otherwise continue
        fig.tight_layout()
    except Exception:
        # original_print(f"Warning: tight_layout() failed for Figure {fig.number}. Displaying as is.")
        pass # Continue saving without tight_layout if it fails
        
    fig.savefig(buf, format='png')
    buf.seek(0)
    
    # Display the image directly in Streamlit
    st.image(buf, caption=caption or f"Figure {fig.number}", use_container_width=True)
    
    # Close the figure to free memory after displaying
    plt.close(fig)
    
    return None

# Monkey patch plt.figure to ensure it's tracked (no changes to this specific patch's logic)
if not hasattr(plt.figure, "__is_patched_by_app__"):
    original_figure = plt.figure
    def patched_figure(*args, **kwargs):
        fig = original_figure(*args, **kwargs)
        return fig
    patched_figure.__is_patched_by_app__ = True
    plt.figure = patched_figure

# Monkey patch plt.savefig to use Streamlit's native display instead of saving to disk
# if not hasattr(plt.savefig, "__is_patched_by_app__"):
# original_savefig_func = plt.savefig # Keep reference if needed for complex pass-through
def patched_savefig(fname, *args, **kwargs): # This is the function definition
    '''
    Intercepts plt.savefig(), displays the current figure in Streamlit,
    and bypasses actual file saving.
    '''
    import os
    current_fig = plt.gcf()

    # Prevent displaying an empty figure
    if not current_fig.axes and not current_fig.get_children():
        if plt.get_fignums():
            original_print(f"Info: Attempted to save/display Figure {current_fig.number} (intended as '{os.path.basename(fname)}'), but it appears to be empty. Skipping display.")
        plt.close(current_fig) # Close the potentially empty figure
        return None

    # If line 86 of the *entire helper string* is here, this is where the issue is.
    # For example, if this line below is considered line 86:
    buf = BytesIO() # This line does not reference 'base64'
    
    valid_savefig_kwargs = {k: v for k, v in kwargs.items() if k not in ['format']}
    
    try:
        current_fig.tight_layout()
    except Exception:
        pass # Continue if tight_layout fails
        
    current_fig.savefig(buf, format='png', **valid_savefig_kwargs)
    buf.seek(0)
    
    fig_num = current_fig.number
    caption_text = f"Figure {fig_num}"
    if isinstance(fname, str):
        caption_text = f"Figure {fig_num}: {os.path.basename(fname)} (Displayed, not saved to disk)"
    
    st.image(buf, caption=caption_text, use_container_width=True)
    
    return None
patched_savefig.__is_patched_by_app__ = True
plt.savefig = patched_savefig

# Monkey patch plt.show to use Streamlit's native display
if not hasattr(plt.show, "__is_patched_by_app__"):
    # original_show_func = plt.show # Keep reference if needed
    def patched_show(*args, **kwargs): # plt.show() args/kwargs are usually for blocking behavior, not relevant here.
        '''
        Intercepts plt.show(), displays the current figure in Streamlit,
        and closes the figure.
        '''
        current_fig = plt.gcf()
        # display_matplotlib_figure will handle checks for empty figure, display, and close it.
        display_matplotlib_figure(current_fig) 
        return None # plt.show() typically doesn't return a value.
    patched_show.__is_patched_by_app__ = True
    plt.show = patched_show
# === End of plotting helper functions ===

def safe_map_categorical(series_to_map, mapping_dict, default_value_for_unmapped=None):
    '''
    Safely maps categorical values in a pandas Series.
    Handles NaN values by preserving them.
    Unmapped non-NaN values can be set to a default or preserved.
    Updates a global 'categorical_mappings' dictionary.
    '''
    # Ensure global categorical_mappings exists in the REPL's scope
    if 'categorical_mappings' not in globals():
        globals()['categorical_mappings'] = {}

    # Make a copy to avoid modifying the original Series passed if it's part of a DataFrame slice
    mapped_series = series_to_map.copy()
    
    # Ensure pandas and numpy are available (usually imported by user code or main app)
    import pandas as pd 
    import numpy as np

    for index, val in series_to_map.items():
        if pd.isna(val):
            mapped_series.loc[index] = np.nan # Preserve NaNs
        elif val in mapping_dict:
            mapped_series.loc[index] = mapping_dict[val]
        elif default_value_for_unmapped is not None:
            mapped_series.loc[index] = default_value_for_unmapped
        # Else (no default_value_for_unmapped and val not in mapping_dict), original value is kept by Series.copy()

    # Store the mapping for this series if a name exists (i.e., it's a DataFrame column)
    if hasattr(series_to_map, 'name') and series_to_map.name:
        globals()['categorical_mappings'][series_to_map.name] = mapping_dict
        print(f"Info: Column '{series_to_map.name}' was mapped using: {mapping_dict}")
    return mapped_series

# Patch pd.DataFrame.__getitem__ to provide helpful error messages for missing columns,
# and to return an empty DataFrame when a list of keys is requested but none are present.
import pandas as _pd # Use a distinct alias for patching
if not hasattr(_pd.DataFrame, "_original_getitem"):
    _pd.DataFrame._original_getitem = _pd.DataFrame.__getitem__
    def _safe_getitem(self, key):
        try:
            return self._original_getitem(key)
        except KeyError as e:
            if isinstance(key, list): # Check if key is a list (for selecting multiple columns)
                existing_cols = [col for col in key if col in self.columns]
                if existing_cols: # If some requested columns exist
                    # original_print(f"Warning: Requested columns {key}, but only found {existing_cols}. Proceeding with found columns.")
                    return self._original_getitem(existing_cols)
                else: # If no requested columns exist
                    original_print(f"Warning: None of the requested columns {key} are in the DataFrame. Available columns: {list(self.columns)}. Returning an empty DataFrame.")
                    return self.iloc[0:0] # Return an empty DataFrame with original columns
            else: # Single key not found
                original_print(f"KeyError: Column '{key}' not found in DataFrame. Available columns: {list(self.columns)}")
                original_print("This error often occurs if a column name was misspelled, or if one-hot encoding or LLM-generated code referenced a non-existent column.")
                raise e # Re-raise the original error after printing info
    _pd.DataFrame.__getitem__ = _safe_getitem

# Print unique values for categorical columns for LLM reference (for all dataframes)
# Ensure pandas is available for this part too.
try:
    import pandas as pd
    global_vars = dict(globals()) # Create a copy to avoid "dictionary changed size during iteration"
    for df_name, df_obj in global_vars.items():
        if isinstance(df_obj, pd.DataFrame) and not df_obj.empty and not df_name.startswith('original_'):
            # Check if it's not one of the specific dataframes to exclude from this printout
            excluded_df_names = [] # Add any df names here if they cause issues or are too verbose
            if df_name in excluded_df_names:
                continue

            # original_print(f"\\n--- Unique Values in Categorical Columns for '{df_name}' (before current execution) ---")
            # categorical_cols = df_obj.select_dtypes(include=['object', 'category']).columns
            # if not categorical_cols.empty:
            #     for col_name in categorical_cols:
            #         try:
            #             unique_vals = df_obj[col_name].unique()
            #             # Truncate long lists of unique values for brevity if necessary
            #             display_limit = 10
            #             num_unique = len(unique_vals)
            #             display_vals = list(unique_vals[:display_limit])
            #             if pd.NA in unique_vals or any(pd.isna(v) for v in unique_vals):
            #                 has_nan_val_str = " (contains NA/NaN)"
            #             else:
            #                 has_nan_val_str = ""
            #             
            #             if num_unique > display_limit:
            #                 original_print(f"Column '{col_name}' ({num_unique} unique): {display_vals}... {has_nan_val_str}")
            #             else:
            #                 original_print(f"Column '{col_name}' ({num_unique} unique): {display_vals}{has_nan_val_str}")
            #         except Exception as e_unique:
            #             original_print(f"Could not get unique values for column '{col_name}': {e_unique}")
            # else:
            #     original_print(f"No object or category columns found in '{df_name}'.")
            # original_print("---------------------------------------------------------------------\\n")
except ImportError:
    original_print("Pandas not available, skipping unique value printing for DataFrames.")
except Exception as e_inspect:
    original_print(f"Error during DataFrame inspection for unique values: {e_inspect}")
    """
    # Prepend helper function definition to the code LLM generates
    final_code_to_execute = safe_mapping_helper_code + "\n" + code_to_run.strip()

    try:
        with redirect_stdout(output_capture):
            # Execute the code within the REPL's global scope
            exec(final_code_to_execute, current_globals)

        # Retrieve all modified dataframes from the REPL's globals
        # These will become the input dfs_current_state for the *next* iteration
        for df_name in dfs_current_state.keys():
            if df_name in current_globals and isinstance(
                current_globals[df_name], pd.DataFrame
            ):
                dfs_after_execution[df_name] = current_globals[df_name].copy()
            elif (
                df_name in dfs_current_state and dfs_current_state[df_name] is not None
            ):
                # If a dataframe wasn't modified or became None, revert to previous state
                dfs_after_execution[df_name] = dfs_current_state[df_name].copy()

        # Update global session state for categorical_mappings
        if "categorical_mappings" in current_globals:
            st.session_state.categorical_mappings.update(
                current_globals["categorical_mappings"]
            )

    except Exception:
        error_occurred_msg = traceback.format_exc()
        # If an error occurs, dfs_after_execution should be the state *before* this failed execution
        # which is dfs_current_state (already copied at the start of this function)
        dfs_after_execution = (
            {name: df.copy() for name, df in dfs_current_state.items()}
            if dfs_current_state
            else {}
        )
        # Do not update categorical_mappings from a failed execution
    finally:
        # Always capture the output, whether success or failure
        text_output = output_capture.getvalue()
        output_capture.close()

        # Check if matplotlib has created any figures that haven't been displayed yet
        if plt.get_fignums():
            print(
                f"Found {len(plt.get_fignums())} unsaved matplotlib figures. Displaying them now."
            )
            for fig_num in plt.get_fignums():
                try:
                    fig = plt.figure(fig_num)
                    # Convert to BytesIO
                    buf = BytesIO()
                    fig.tight_layout()
                    fig.savefig(buf, format="png")
                    buf.seek(0)

                    # Convert BytesIO to base64 string for storage
                    buf_copy = BytesIO(buf.getvalue())  # Make a copy for storage
                    buf_copy.seek(0)
                    img_base64 = base64.b64encode(buf_copy.read()).decode("utf-8")
                    new_image_data_b64.append(img_base64)

                    # Display directly in Streamlit (using the original buffer)
                    buf.seek(0)
                    st.image(buf, caption=f"Figure {fig_num}", use_container_width=True)
                    plt.close(fig)
                except Exception as e_display:
                    print(f"Error displaying figure {fig_num}: {e_display}")

    # Extract retrieved chunks if present in the output
    if (
        "Retrieved documents for query:" in text_output
        or "Retrieved documents from UPLOADED FILES for query:" in text_output
        or "DOCUMENTS FOUND:" in text_output
    ):
        # Parse the retrieval sections
        retrieval_sections = []
        lines = text_output.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith("Retrieved documents") or line.startswith(
                "DOCUMENTS FOUND:"
            ):
                section_lines = [lines[i]]
                j = i + 1
                # Continue until we hit the next section or end of text
                while j < len(lines) and not (
                    lines[j].strip().startswith("Search results for:")
                    or lines[j].strip().startswith("Retrieved documents")
                ):
                    section_lines.append(lines[j])
                    j += 1
                retrieval_sections.append("\n".join(section_lines))
                i = j - 1
            i += 1

        # Add to retrieved_chunks
        retrieved_chunks = retrieval_sections

        # Store in session state
        st.session_state.retrieved_chunks.extend(retrieved_chunks)

        # Print confirmation for debugging
        print(f"Captured {len(retrieval_sections)} document retrieval sections")
        for i, section in enumerate(retrieval_sections):
            first_line = section.split("\n")[0] if "\n" in section else section[:50]
            print(f"Section {i + 1} starts with: {first_line}")

        # Also check for content sections that might not be properly formatted
        content_sections = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if "Content:" in line and (
                "Source:" in "\n".join(lines[max(0, i - 5) : i + 1])
                or "Title:" in "\n".join(lines[max(0, i - 5) : i + 1])
            ):
                # This looks like a content line that might be part of a document chunk
                # Try to find the start of this chunk
                start_idx = i
                while start_idx > 0 and not lines[start_idx].strip().startswith(
                    tuple(str(n) + "." for n in range(1, 11))
                ):
                    start_idx -= 1

                if start_idx < i:  # We found a potential chunk start
                    # Find the end of this chunk
                    end_idx = i + 1
                    while end_idx < len(lines) and not (
                        lines[end_idx]
                        .strip()
                        .startswith(tuple(str(n) + "." for n in range(1, 11)))
                        or lines[end_idx].strip() == ""
                    ):
                        end_idx += 1

                    chunk_text = "\n".join(lines[start_idx:end_idx])
                    if chunk_text not in content_sections:
                        content_sections.append(chunk_text)
            i += 1

        # If we found content sections that weren't part of a proper retrieval section,
        # create a synthetic retrieval section for them
        if content_sections and not retrieval_sections:
            synthetic_section = (
                "Retrieved documents from UPLOADED FILES (reconstructed):\n\n"
                + "\n\n".join(content_sections)
            )
            retrieval_sections.append(synthetic_section)
            retrieved_chunks = [synthetic_section]
            st.session_state.retrieved_chunks.append(synthetic_section)
            print(
                f"Created synthetic retrieval section with {len(content_sections)} content chunks"
            )

    return (
        text_output,
        error_occurred_msg,
        new_image_data_b64,
        dfs_after_execution,
        retrieved_chunks,
    )


# --- Main Application UI and Logic ---
def get_all_retrieved_chunks(results_data):
    """
    Collects and deduplicates all retrieved chunks from session state and history.
    Returns a list of unique chunks.
    """
    all_chunks = []
    if st.session_state.get("retrieved_chunks"):
        for chunk in st.session_state.retrieved_chunks:
            if chunk and chunk not in all_chunks:
                all_chunks.append(chunk)
    if results_data.get("history"):
        for item in results_data["history"]:
            if "chunks" in item and item["chunks"]:
                for chunk in item["chunks"]:
                    if chunk and chunk not in all_chunks:
                        all_chunks.append(chunk)
            if "output" in item and "Retrieved documents for query:" in item["output"]:
                output_text = item["output"]
                lines = output_text.splitlines()
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    if line.startswith("Retrieved documents for query:"):
                        section_start = i
                        j = i + 1
                        while j < len(lines) and not lines[j].strip().startswith("Search results for:"):
                            j += 1
                        retrieval_section = "\n".join(lines[section_start:j])
                        if retrieval_section and retrieval_section not in all_chunks:
                            all_chunks.append(retrieval_section)
                        i = j - 1
                    i += 1
    if not all_chunks and results_data.get("citations"):
        for citation in results_data.get("citations"):
            if isinstance(citation, dict):
                context = citation.get("context", "")
                url = citation.get("url", "")
                chunk = f"Source: {url}\nContent: {context}"
                if chunk and chunk not in all_chunks:
                    all_chunks.append(chunk)
    return all_chunks

def main_application():
    st.title("🤖 AI Analyzer")
    st.markdown(
        """
        Welcome to the AI Analyzer! Upload your data or documents and ask questions.
        The AI will use a multi-step process, including code execution and document retrieval,
        to provide a comprehensive answer. This is for educational use only. David Liebovitz, MD
        """
    )

    # --- Sidebar Setup ---
    st.sidebar.header("⚙️ Controls & Settings")

    # Data Upload - Multiple files
    st.sidebar.subheader("File Upload")

    # Create two separate uploaders - one for data files, one for documents
    data_files = st.sidebar.file_uploader(
        "Upload Data Files (CSV/Excel)",
        type=["csv", "xlsx"],
        accept_multiple_files=True,
        key="data_uploader",
    )

    document_files = st.sidebar.file_uploader(
        "Upload Documents (PDF/Word/Image)",
        type=["pdf", "docx", "doc", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
        key="document_uploader",
    )

    # Process data files (CSV/Excel)
    if data_files:
        import pandas as pd  # Ensure pandas is available in this scope

        # Process each uploaded file
        for uploaded_file in data_files:
            try:
                # Create a clean dataframe name from the filename
                df_name = os.path.splitext(uploaded_file.name)[0]
                # Replace invalid characters with underscore and ensure it starts with a letter
                df_name = re.sub(r"[^a-zA-Z0-9_]", "_", df_name)
                if not df_name[0].isalpha():
                    df_name = "df_" + df_name

                # Read the file into a dataframe
                if uploaded_file.name.endswith(".csv"):
                    df = pd.read_csv(uploaded_file)
                else:
                    df = pd.read_excel(uploaded_file)

                # Store the original dataframe
                st.session_state.dataframes[df_name] = df

                # Create a working copy
                st.session_state.gpt_working_dfs[df_name] = df.copy()

                # Store metadata about the dataframe using Pydantic model
                st.session_state.dataframe_metadata[df_name] = DataFrameMetadata(
                    columns=list(df.columns),
                    original_columns=list(df.columns),
                    dtypes={col: str(df[col].dtype) for col in df.columns},
                    filename=uploaded_file.name,
                    shape=df.shape,
                ).model_dump()

                # Set as active dataframe if it's the first one
                if st.session_state.active_df is None:
                    st.session_state.active_df = df_name

                st.sidebar.success(
                    f"Uploaded data file '{uploaded_file.name}' as '{df_name}'"
                )
            except Exception as e:
                st.sidebar.error(f"Error loading data file '{uploaded_file.name}': {e}")

        # Reset categorical mappings on new file upload
        st.session_state.categorical_mappings = {}
        st.session_state.experts = []
        st.session_state.expert_domains = []
        st.session_state.expert_questions = []
        st.session_state.expert_answers = []

    def save_modified_dataframe(new_name: str, modified_df, original_df):
        """
        Save a modified dataframe under a new name safely.
        Preserves the original list of columns from the original dataframe.
        Updates the session state for both gpt_working_dfs and dataframe_metadata.
        """
        # Save the modified dataframe as a new working copy and original copy
        st.session_state.gpt_working_dfs[new_name] = modified_df.copy()
        st.session_state.dataframes[new_name] = modified_df.copy()
        # Preserve the original columns from the original dataframe
        original_columns = list(original_df.columns)
        st.session_state.dataframe_metadata[new_name] = DataFrameMetadata(
            columns=list(modified_df.columns),
            original_columns=original_columns,
            dtypes={col: str(modified_df[col].dtype) for col in modified_df.columns},
            filename=f"modified_{new_name}",
            shape=modified_df.shape,
        ).model_dump()

    if st.session_state.dataframes:
        st.sidebar.info(
            f"Using {len(st.session_state.dataframes)} previously uploaded dataframe(s)."
        )
    else:
        st.sidebar.info("Upload CSV or Excel files to analyze your data.")

    # Process document files (PDF/DOCX)
    if document_files:
        from rag_utils import get_embeddings_model
        from langchain_community.vectorstores import FAISS
        from langchain_core.documents import Document

        # Check if we need to create a new vector store
        create_new_vector_store = False

        # Process each uploaded document
        for uploaded_file in document_files:
            try:
                file_ext = uploaded_file.name.lower().split(".")[-1]
                # If image, process directly with vision model
                if file_ext in ["png", "jpg", "jpeg"]:
                    st.sidebar.info(
                        f"Processing image '{uploaded_file.name}' with GPT-4o vision..."
                    )
                    try:
                        import openai

                        openai_api_key = st.secrets.get("OPENAI_API_KEY")
                        if not openai_api_key:
                            st.sidebar.error(
                                "OPENAI_API_KEY not found in secrets. Cannot use GPT-4o vision."
                            )
                            continue

                        image_bytes = uploaded_file.read()
                        vision_prompt = (
                            "Describe the content of this image. "
                            "Extract any text present, summarize the main visual elements, and provide a concise description."
                        )
                        client = openai.OpenAI(api_key=openai_api_key)
                        mime_type = "image/png" if file_ext == "png" else "image/jpeg"
                        import base64

                        image_b64 = base64.b64encode(image_bytes).decode("utf-8")
                        data_url = f"data:{mime_type};base64,{image_b64}"
                        completion = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "system",
                                    "content": "You are a document analysis assistant.",
                                },
                                {
                                    "role": "user",
                                    "content": [
                                        {"type": "text", "text": vision_prompt},
                                        {
                                            "type": "image_url",
                                            "image_url": {"url": data_url},
                                        },
                                    ],
                                },
                            ],
                            max_tokens=2048,
                            temperature=0.0,
                        )
                        vision_text = completion.choices[0].message.content.strip()
                        if vision_text:
                            doc_name = os.path.splitext(uploaded_file.name)[0]
                            doc_info = DocumentInfo(
                                text=vision_text,
                                filename=uploaded_file.name,
                                type=file_ext,
                            )
                            st.session_state.document_texts[doc_name] = (
                                doc_info.model_dump()
                            )
                            st.sidebar.success(
                                f"Processed image '{uploaded_file.name}' using GPT-4o vision."
                            )
                            create_new_vector_store = True
                        else:
                            st.sidebar.error(
                                f"GPT-4o vision did not return any text for image '{uploaded_file.name}'."
                            )
                    except Exception as e_vision:
                        st.sidebar.error(
                            f"Failed to process image '{uploaded_file.name}' with GPT-4o vision: {e_vision}"
                        )
                    continue  # Skip to next file

                # Process the document to extract text (for PDF/DOCX)
                extracted_text, success = process_document_file(uploaded_file)

                if success:
                    # Store the extracted text using Pydantic model
                    doc_name = os.path.splitext(uploaded_file.name)[0]
                    doc_info = DocumentInfo(
                        text=extracted_text,
                        filename=uploaded_file.name,
                        type=uploaded_file.name.split(".")[-1].lower(),
                    )
                    st.session_state.document_texts[doc_name] = doc_info.model_dump()
                    st.sidebar.success(
                        f"Processed document '{uploaded_file.name}' successfully"
                    )
                    create_new_vector_store = True
                else:
                    # Fallback: If PDF and failed, try OpenAI GPT-4o vision
                    if uploaded_file.name.lower().endswith(".pdf"):
                        st.sidebar.warning(
                            f"Standard PDF text extraction failed for '{uploaded_file.name}'. Trying GPT-4o vision analysis..."
                        )
                        try:
                            import openai

                            openai_api_key = st.secrets.get("OPENAI_API_KEY")
                            if not openai_api_key:
                                st.sidebar.error(
                                    "OPENAI_API_KEY not found in secrets. Cannot use GPT-4o vision."
                                )
                                continue

                            # Read PDF bytes
                            pdf_bytes = uploaded_file.read()
                            # Prepare the prompt for vision model
                            vision_prompt = (
                                "Extract all readable text and summarize the main content of this PDF. "
                                "If the PDF contains tables or images with text, describe them as well. "
                                "Return the extracted text and a summary."
                            )
                            client = openai.OpenAI(api_key=openai_api_key)
                            # Pass the file as a tuple: (filename, bytes, mime_type)
                            completion = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": "You are a document analysis assistant.",
                                    },
                                    {
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": vision_prompt},
                                            {
                                                "type": "file",
                                                "file": (
                                                    uploaded_file.name,
                                                    pdf_bytes,
                                                    "application/pdf",
                                                ),
                                            },
                                        ],
                                    },
                                ],
                                max_tokens=4096,
                                temperature=0.0,
                            )
                            vision_text = completion.choices[0].message.content.strip()
                            if vision_text:
                                doc_name = os.path.splitext(uploaded_file.name)[0]
                                doc_info = DocumentInfo(
                                    text=vision_text,
                                    filename=uploaded_file.name,
                                    type="pdf",
                                )
                                st.session_state.document_texts[doc_name] = (
                                    doc_info.model_dump()
                                )
                                st.sidebar.success(
                                    f"Processed '{uploaded_file.name}' using GPT-4o vision."
                                )
                                create_new_vector_store = True
                            else:
                                st.sidebar.error(
                                    f"GPT-4o vision did not return any text for '{uploaded_file.name}'."
                                )
                        except Exception as e_vision:
                            st.sidebar.error(
                                f"Failed to process '{uploaded_file.name}' with GPT-4o vision: {e_vision}"
                            )
                    else:
                        st.sidebar.error(
                            f"Failed to process document '{uploaded_file.name}': {extracted_text}"
                        )
            except Exception as e:
                st.sidebar.error(
                    f"Error processing document '{uploaded_file.name}': {e}"
                )

        # Create vector store from all documents if needed
        if create_new_vector_store and st.session_state.document_texts:
            try:
                with st.spinner("Creating document vector store..."):
                    # Get embeddings model
                    embeddings = get_embeddings_model()

                    # --- Ensure sidebar settings are used for chunking ---
                    # These are set in the sidebar expander below, so must be read here
                    chunk_size = st.session_state.get("rag_chunk_size", 1500)
                    chunk_overlap = st.session_state.get("rag_chunk_overlap", 200)

                    if embeddings:
                        # Create documents for vector store
                        documents = []

                        for (
                            doc_name,
                            doc_info,
                        ) in st.session_state.document_texts.items():
                            # Split text into chunks using sidebar settings
                            chunks = create_document_chunks(
                                doc_info["text"],
                                chunk_size=chunk_size,
                                chunk_overlap=chunk_overlap,
                            )

                            # Create a document for each chunk using Pydantic model
                            for i, chunk_text in enumerate(chunks):
                                metadata = DocumentMetadata(
                                    source=doc_info["filename"],
                                    chunk=i,
                                    document_name=doc_name,
                                    document_type=doc_info["type"],
                                )
                                documents.append(
                                    Document(
                                        page_content=chunk_text,
                                        metadata=metadata.model_dump(),
                                    )
                                )

                        if documents:
                            # Create vector store
                            vector_store = FAISS.from_documents(documents, embeddings)
                            st.session_state.document_vector_store = vector_store
                            # Use the k_docs setting from session state when creating the vector store
                            st.session_state.document_vector_store.k = st.session_state.get("rag_k_docs", 5)
                            st.sidebar.success(
                                f"Created vector store with {len(documents)} chunks from {len(st.session_state.document_texts)} documents"
                            )
                        else:
                            st.sidebar.warning("No document chunks were created")
                    else:
                        st.sidebar.error(
                            "Could not initialize embeddings model for document processing"
                        )
            except Exception as e:
                st.sidebar.error(f"Error creating document vector store: {e}")
    else:
        st.sidebar.info(
            "Upload PDF, Word, or Image files to include them in the analysis."
        )

    # --- New: Unified "Augment with PubMed and Web Search" Option ---
    # If any data files or document files are uploaded, default to confining analysis to those sources.
    # Add a sidebar checkbox to allow user to augment with PubMed and Web Search (off by default).
    files_uploaded = bool(st.session_state.dataframes) or bool(
        st.session_state.document_texts
    )
    # Default value is True if no files are uploaded, False if files are uploaded
    default_augment_value = not files_uploaded
    st.session_state["augment_with_pubmed_web"] = st.sidebar.checkbox(
        "Augment with PubMed and Web Search",
        value=st.session_state.get("augment_with_pubmed_web", default_augment_value),
        help="If checked, the analysis will also use PubMed and web search in addition to your uploaded files. By default, only your uploaded files are used.",
    )

    # For backward compatibility, set use_web_search and include_medical_web_search accordingly
    if files_uploaded:
        st.session_state.use_web_search = st.session_state["augment_with_pubmed_web"]
        st.session_state.include_medical_web_search = st.session_state[
            "augment_with_pubmed_web"
        ]
    else:
        # If no files uploaded, allow web search by default
        st.session_state.use_web_search = True
        st.session_state.include_medical_web_search = True
    # Dataframe selector if multiple dataframes are loaded
    if len(st.session_state.dataframes) > 0:
        df_options = list(st.session_state.dataframes.keys())
        selected_df = st.sidebar.selectbox(
            "Select active dataframe for preview:",
            df_options,
            index=df_options.index(st.session_state.active_df)
            if st.session_state.active_df in df_options
            else 0,
        )
        st.session_state.active_df = selected_df

        # Display the selected dataframe
        with st.sidebar.expander("View Selected Data Head"):
            st.dataframe(st.session_state.dataframes[selected_df].head())
            st.text(
                f"Shape: {st.session_state.dataframe_metadata[selected_df]['shape']}"
            )

    # All other settings under a single Advanced Settings expander
    with st.sidebar.expander("Advanced Settings", expanded=False):
        # Iteration Control
        st.subheader("Analysis Settings")
        max_iterations = st.slider(
            "Max Refinement Iterations:",
            min_value=1,
            max_value=10,
            value=5,
            help="Number of times the AI will try to refine its answer or fix errors.",
        )

        # RAG Advanced Settings
        st.subheader("🔍 RAG Settings")
        st.write("Configure Retrieval Augmented Generation settings:")
        st.subheader("RAGAS Model Selection")
        selected_ragas_model = st.selectbox("Select RAGAS Model to use", options=["gpt-4o-mini", "gpt-4.1-mini", "gpt-4.1"], index=0)
        st.session_state.ragas_model = selected_ragas_model

        # Number of search results
        if "rag_num_results" not in st.session_state:
            st.session_state.rag_num_results = 5
        st.session_state.rag_num_results = st.slider(
            "Number of search results:",
            min_value=3,
            max_value=10,
            value=st.session_state.rag_num_results,
            help="Number of web pages to retrieve from search",
        )

        # Chunk size for uploaded documents
        if "rag_chunk_size" not in st.session_state:
            st.session_state.rag_chunk_size = 1500
        st.session_state.rag_chunk_size = st.number_input(
            "Chunk size (uploaded documents):",
            min_value=500,
            max_value=2000,
            value=st.session_state.rag_chunk_size,
            step=100,
            help="Size of text chunks for processing uploaded document files",
        )

        # Chunk overlap for uploaded documents
        if "rag_chunk_overlap" not in st.session_state:
            st.session_state.rag_chunk_overlap = 200
        st.session_state.rag_chunk_overlap = st.number_input(
            "Chunk overlap (uploaded documents):",
            min_value=50,
            max_value=500,
            value=st.session_state.rag_chunk_overlap,
            step=50,
            help="Overlap between text chunks for uploaded document files",
        )

        # Chunk size for web/PubMed content
        if "web_chunk_size" not in st.session_state:
            st.session_state.web_chunk_size = 1500
        st.session_state.web_chunk_size = st.number_input(
            "Chunk size (web/PubMed):",
            min_value=500,
            max_value=2000,
            value=st.session_state.web_chunk_size,
            step=100,
            help="Size of text chunks for PubMed and web search content",
        )

        # Chunk overlap for web/PubMed content
        if "web_chunk_overlap" not in st.session_state:
            st.session_state.web_chunk_overlap = 200
        st.session_state.web_chunk_overlap = st.number_input(
            "Chunk overlap (web/PubMed):",
            min_value=50,
            max_value=500,
            value=st.session_state.web_chunk_overlap,
            step=50,
            help="Overlap between text chunks for PubMed and web search content",
        )

        # Scrape content toggle
        if "rag_scrape_content" not in st.session_state:
            st.session_state.rag_scrape_content = True
        st.session_state.rag_scrape_content = st.checkbox(
            "Scrape full web content",
            value=st.session_state.rag_scrape_content,
            help="If enabled, will attempt to scrape full content from web pages",
        )

        # Number of retrieved documents
        if "rag_k_docs" not in st.session_state:
            st.session_state.rag_k_docs = 5
        st.session_state.rag_k_docs = st.slider(
            "Number of documents to retrieve:",
            min_value=1,
            max_value=10,
            value=st.session_state.rag_k_docs,
            help="Number of documents to retrieve from vector store",
        )

        # PubMed search settings
        st.subheader("📚 Medical & PubMed Search Settings")

        # Auto-enable PubMed for medical questions
        auto_pubmed = st.checkbox(
            "Automatically use PubMed for medical questions",
            value=True,
            help="When enabled, medical questions will automatically trigger PubMed searches",
        )
        st.session_state["auto_pubmed"] = auto_pubmed

        # Years back for PubMed search
        years_back = st.slider(
            "Years Back for PubMed Search",
            min_value=1,
            max_value=15,
            value=6,
            step=1,
            help="Set the number of years back to search PubMed. A larger range may find more relevant articles.",
        )
        st.session_state["years_back"] = years_back

        # Number of abstracts to review
        max_results = st.slider(
            "Number of Abstracts to Review",
            min_value=3,
            max_value=20,
            value=6,
            step=1,
            help="Set the number of abstracts to review.",
        )
        st.session_state["max_results"] = max_results

        # Filter relevance toggle
        filter_relevance = st.toggle(
            "Filter Relevance of PubMed searching",
            value=True,
            help="Toggle to deselect.",
        )
        st.session_state["filter_relevance"] = filter_relevance

        if filter_relevance:
            relevance_threshold = st.slider(
                "Relevance Threshold",
                min_value=0.3,
                max_value=1.0,
                value=0.7,
                step=0.05,
                help="Set the minimum relevance score to consider an item relevant.",
            )
            st.session_state["relevance_threshold"] = relevance_threshold
        else:
            relevance_threshold = 0.65
            st.session_state["relevance_threshold"] = relevance_threshold
            st.write("Top sources will be added to the database regardless.")

        # Cutting-edge research toggle
        cutting_edge = st.checkbox(
            "Include Cutting-Edge Research in PubMed (default is consensus review articles)",
            help="Check to include latest, not yet consensus, articles in the search for medical content.",
            value=False,
        )
        st.session_state["cutting_edge"] = cutting_edge

        # Medical search strategy
        search_strategy_options = st.radio(
            "Medical Search Strategy",
            ["Balanced (Web + PubMed)", "PubMed Focused", "Web Focused"],
            index=0,
            help="Choose how to balance web search and PubMed for medical questions",
        )
        st.session_state["medical_search_strategy"] = search_strategy_options

    # --- Main Page Layout ---
    # Create two main columns for the interface
    main_col1, main_col2 = st.columns([1, 1])  # Adjust ratio as needed

    with main_col1:
        st.subheader("💬 Ask Your Question")
        user_question = st.text_area(
            "Enter your question about the data/documents (if uploaded) or a general query:",
            key="user_question_input",
            height=100,
            placeholder="E.g., 'Analyze the sales data in 'sales_df'', 'Summarize the key points in the uploaded PDF', or 'Explain quantum computing.'",
        )

        # The Analyze button will be in the second column, but the logic is triggered here
        # We need a way to trigger the analysis based on the button click in the other column
        # Streamlit handles this automatically when the button's state changes

    # The analyze_clicked logic needs to be outside the column context to affect the whole page state
    # The button itself is placed in main_col2 below
    analyze_clicked = False  # Initialize here

    # Web search is now controlled via the sidebar checkbox
    compare_web = False

    # Medical web search option is now handled in the sidebar

    # Place the Analyze button in the second column, outside the results display logic
    with main_col2:
        # Add some vertical space to align the button better with the text area in col1
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        analyze_clicked = st.button(
            "🚀 Analyze / Ask", key="analyze_button_col2", use_container_width=True
        )

    # The analysis execution logic should run regardless of which column the button is in
    # It should only run if analyze_clicked is True and a question is entered
    if analyze_clicked and user_question.strip():
        # --- Reset all relevant session state for a truly clean analysis ---
        keys_to_clear = [
            "validated_section1",
            "validation_results",
            "current_section1",
            "followup_chat_thread",
            "google_search_followup_results",
            "retrieved_chunks",
            "experts",
            "expert_domains",
            "expert_questions",
            "expert_answers",
            "analysis_results",
            "all_analysis_images",
            "pubmed_search_terms",
            "articles",
            "older_pubmed_articles_alert",
            "is_medical_question",
            "search_strategy",
            "ragas_model",
            "outputs_path",
            "current_analysis_timestamp",
            "last_agent_question",
            "categorical_mappings",
            "document_vector_store",
        ]
        for key in keys_to_clear:
            if key in st.session_state:
                # For outputs_path, optionally clean up the directory if needed
                if key == "outputs_path" and st.session_state[key]:
                    try:
                        import shutil
                        shutil.rmtree(st.session_state[key], ignore_errors=True)
                    except Exception:
                        pass
                del st.session_state[key]
        # Re-initialize outputs_path and other required keys
        st.session_state.outputs_path = tempfile.mkdtemp(prefix="gpt_plots_")
        st.session_state.retrieved_chunks = []
        st.session_state.followup_chat_thread = []
        st.session_state.google_search_followup_results = ""
        st.session_state.all_analysis_images = []
        st.session_state.categorical_mappings = {}
        st.session_state.experts = []
        # -----------------------------------------------------

        # Check if this is a medical question and apply the medical web search preference
        from rag_utils import is_medical_topic

        is_medical = is_medical_topic(user_question)
        files_uploaded = bool(st.session_state.dataframes) or bool(
            st.session_state.document_texts
        )
        # New logic: If files are uploaded, only use PubMed/Web if user checked the box
        if files_uploaded:
            compare_web = st.session_state.get("augment_with_pubmed_web", False)
            # print(f"Files uploaded. Augment with PubMed/Web: {compare_web}") # Removed to prevent I/O error
        else:
            # If no files uploaded, use web search by default
            compare_web = True
            # print("No files uploaded. Defaulting to allow PubMed/Web search.") # Removed to prevent I/O error

        llm = get_llm_instance()
        repl = get_python_repl()

        if llm is None:  # LLM initialization failed
            return  # Stop further processing

        # Clean plot directory before new analysis
        clean_plot_directory(st.session_state.outputs_path)

        # Reset retrieved chunks for new analysis
        st.session_state.retrieved_chunks = []

        # Clear vector store to prevent old content from cluttering the assessment
        st.session_state.document_vector_store = (
            None  # Corrected from vector_store to document_vector_store
        )

        st.session_state.last_agent_question = user_question  # Store current question
        st.session_state.current_analysis_timestamp = str(int(time.time()))
        current_history = []

        # Initialize working dataframes for this analysis run
        # Create copies of all working dataframes
        current_dfs_for_analysis = (
            {name: df.copy() for name, df in st.session_state.gpt_working_dfs.items()}
            if st.session_state.gpt_working_dfs
            else {}
        )
        original_dfs_for_analysis = (
            {name: df.copy() for name, df in st.session_state.dataframes.items()}
            if st.session_state.dataframes
            else {}
        )

        # Reset categorical mappings for a new analysis run if dataframes exist
        if current_dfs_for_analysis:
            st.session_state.categorical_mappings = {}

        # Reset all analysis images for this run
        st.session_state.all_analysis_images = []

        # --- Iterative Analysis Loop (Hidden in an expander in the right column) ---
        with main_col2:
            progress_bar = st.progress(0)
            status_text = st.empty()
            with st.expander("Intermediate Iteration Details", expanded=False):
                final_code, final_output, final_error, final_images = "", "", None, []
                
                for i in range(1, max_iterations + 1):
                    status_text.info(f"🧠 Processing Iteration {i}/{max_iterations}...")
                    progress_bar.progress(i / max_iterations)
                
                    # Pass the current state of the working DataFrames to the LLM
                    # Also pass document availability flag and research mode
                    doc_available = st.session_state.document_vector_store is not None
                    research_mode = st.session_state.get("research_assistant_mode", False)
                
                    # If there was an error in the previous iteration, prepend a message to the user question
                    # to encourage the LLM to interpret and fix the error in the next code generation
                    user_question_for_llm = user_question
                    if final_error:
                        user_question_for_llm = (
                            f"{user_question}\n\n"
                            f"NOTE: The previous code (iteration {i - 1}) resulted in this error:\n"
                            f"{final_error}\n"
                            f"Please interpret the error, diagnose the likely cause, and generate corrected code to fix it in this iteration."
                        )
                
                    code = get_code_from_llm(
                        llm,
                        user_question_for_llm,
                        current_dfs_for_analysis,  # Pass snapshot of current dfs state
                        i,
                        max_iterations,
                        final_code,
                        final_output,
                        final_error,  # Results from previous iteration
                        compare_web=compare_web,  # Pass web search preference
                        doc_available=doc_available,  # Pass document availability flag
                        research_mode=research_mode,  # Pass research mode flag
                    )
                
                    if not code.strip() or code.startswith("# LLM communication error"):
                        st.error(
                            f"LLM failed to generate valid code in iteration {i}. Stopping analysis."
                        )
                        final_error = (
                            code if code.startswith("# LLM") else "LLM returned empty code."
                        )
                        break  # Exit loop if no code or LLM error
                
                    text_output, error_msg, new_images_b64, dfs_after_run, new_chunks = (
                        run_code_in_repl(
                            repl,
                            code,
                            current_dfs_for_analysis,  # Pass current dfs state to REPL
                            original_dfs_for_analysis,  # Pass original dfs state to REPL
                            iteration=i,  # Pass the current iteration number
                        )
                    )
                
                    # Update current_dfs_for_analysis for the *next* iteration
                    current_dfs_for_analysis = (
                        dfs_after_run  # These dfs_after_run are from REPL
                    )
                
                    # Accumulate images from this iteration
                    if new_images_b64:
                        st.session_state.all_analysis_images.extend(new_images_b64)
                
                    # Convert BytesIO objects to base64 strings for storage in the Pydantic model
                    processed_images = []
                    for img in new_images_b64:
                        if isinstance(img, BytesIO):
                            # Convert BytesIO to base64 string
                            img.seek(0)
                            img_base64 = base64.b64encode(img.read()).decode("utf-8")
                            processed_images.append(img_base64)
                        elif isinstance(img, str):
                            # Already a string (path or base64)
                            processed_images.append(img)
                        else:
                            # Skip invalid types
                            print(f"Skipping invalid image type: {type(img)}")
                
                    # Store iteration details using Pydantic model
                    iteration = AnalysisIteration(
                        iteration=i,
                        code=code,
                        output=text_output,
                        error=error_msg,
                        images=processed_images,  # Store processed images (base64 strings)
                        chunks=new_chunks,
                    )
                    current_history.append(iteration.model_dump())
                
                    # Update overall final results with this iteration's outcome
                    final_code, final_output, final_error, final_images = (
                        code,
                        text_output,
                        error_msg,
                        new_images_b64,
                    )
                
                    if (
                        not final_error and i < max_iterations
                    ):  # If successful and not last iteration, check for completeness
                        completion_check_prompt = f"""Review the following:
        User Question: "{user_question}"
        Code Executed:
        ```python
        {final_code}
        ```
        Output/Result:
        ```
        {final_output}
        ```
        Does this output completely and correctly answer the user's question?
        Respond with only "YES" or "NO".
        """
                        try:
                            with st.spinner(
                                f"🤖 AI checking answer completeness (Iteration {i})..."
                            ):
                                comp_response = llm.invoke(completion_check_prompt)
                            comp_answer = comp_response.content.strip().upper()
                            if comp_answer == "YES":
                                status_text.success(
                                    f"✅ Analysis complete after {i} iterations."
                                )
                                break  # Exit loop if LLM deems answer complete
                        except Exception as e_comp:
                            st.warning(f"Could not check answer completeness: {e_comp}")
                    elif final_error and i == max_iterations:
                        status_text.warning(
                            f"⚠️ Analysis completed {max_iterations} iterations with an error in the final step."
                        )
                    elif not final_error and i == max_iterations:
                        status_text.success(
                            f"✅ Analysis completed after {max_iterations} iterations."
                        )
                    elif final_error and i < max_iterations:
                        status_text.warning(
                            f"⚠️ Error in iteration {i}. Attempting refinement..."
                        )
                
                progress_bar.empty()  # Clear progress bar
                if not final_error:
                    status_text.success("✅ Analysis Finished!")
                else:
                    status_text.error(
                        f"⚠️ Analysis finished with an error: {final_error.splitlines()[-1]}"
                    )


        # Use current_history to build a results_data-like dict for get_all_retrieved_chunks
        results_data_for_summary = {
            "history": current_history,
            "citations": [],
        }
        collected_chunks_for_summary = get_all_retrieved_chunks(results_data_for_summary)

        # Check if we have any images to include in the summary
        image_info_for_summary = ""
        if final_images:
            image_info_for_summary = f"\n\n{len(final_images)} plot(s) were generated and are included in the analysis."
        elif st.session_state.all_analysis_images:
            image_info_for_summary = f"\n\n{len(st.session_state.all_analysis_images)} plot(s) were generated across all iterations and are included in the analysis."

        if collected_chunks_for_summary:
            all_retrieved_chunks_text = (
                "\n\n---\n**Retrieved Context (from documents/web search):**\n---\n"
                + "\n\n---\n".join(collected_chunks_for_summary)
                + "\n---\n"
            )
        else:
            all_retrieved_chunks_text = "\n\n---\nNo specific context was retrieved from documents or web search during the analysis.\n-\n"

        summary_prompt_text = f"""The user asked: "{user_question}"

The final Python code executed was:
```python
{final_code}
```
The output from this code was:
```
{final_output}
```
{len(final_images)} plot(s) were generated: {", ".join([f"plot_iter{i + 1}.png" for i, _ in enumerate(final_images)]) if final_images else "None"}.
{"An error occurred: " + final_error if final_error else "No errors occurred."}
{image_info_for_summary}
{all_retrieved_chunks_text}
Based on all the information above (user's question, code, output, errors, and retrieved context), provide a comprehensive answer.

The full response structure MUST be:

### **1. Best Answer from Retrieved Context**  
- Include findings **exclusively from the retrieved context** provided above. If no context was retrieved or it's not relevant, state that clearly.
- Organize content to facilitate mastery of the topic. Use technical terms and avoid disclaimers.
- If applicable (e.g., medical topics), clearly indicate the **strength of evidence** (e.g., systematic reviews, RCTs, meta-analyses) found in the context.  
- Highlight **clinical guidelines, expert consensus**, and authoritative sources if present in the context.  
- When citing sources from the context, integrate them **in-line** or at the end of this section.  
- This section should only include information based on the retrieved context. Do not introduce new assertions or facts.
- IMPORTANT: Any URLs included in this section MUST ONLY be those explicitly found in the retrieved context. DO NOT create, guess, or generate any URLs that aren't directly present in the context.

   **Example Format (if applicable, using information from context):**  
   - "A 2023 meta-analysis of 12 RCTs published in *Journal X* (from context) found that [intervention] reduced [outcome] by X% (high-quality evidence)."

### **2. Additional Insights from the Model's Knowledge**  
- Provide a **comprehensive answer** derived from your own training corpus, going beyond the retrieved context if necessary or if no context was available.
- Address **nuances, limitations, or broader considerations** not covered by the retrieved context.  
- Distinguish this as **model-derived knowledge**, separate from retrieved evidence.  
- DO NOT include direct URLs or links in this section, as they may be outdated or incorrect.
- If references would be helpful, format them as Google search links: [Search for topic](https://www.google.com/search?q=search+terms+here) or Google Scholar links: [Academic research on topic](https://scholar.google.com/scholar?q=search+terms+here)

   **Example Format (if applicable):**  
   - "Based on broader evidence, [strategy] has been effective in [specific patient populations] but remains controversial due to [X factor]."

### **3. Application & Decision-Making** 
(Adapt title and content if not a clinical/medical question. E.g., "Practical Application & Implications" or "Data Analysis Insights & Actions")
- Discuss **feasibility and specific considerations** (e.g., for medical: comorbidities, socioeconomic factors, adherence; for data analysis: data quality, model limitations, interpretation of results; for general questions: practical implications, real-world use).
- Outline **ethical issues, risks, or implementation hurdles** in real-world settings.  
- Provide **actionable recommendations**, addressing any serious implications, monitoring strategies or workflows.  
- Anticipate 2–3 likely **follow-up questions** the user might ask and address them concisely.  
- DO NOT include direct URLs or links in this section, as they may be outdated or incorrect.
- If references would be helpful, format them as Google search links: [Search for topic](https://www.google.com/search?q=search+terms+here) or Google Scholar links: [Academic research on topic](https://scholar.google.com/scholar?q=search+terms+here)

   **Example Format (if applicable):**  
   - "Although [X intervention] is guideline-recommended, barriers such as [cost, access] should be considered. Alternatives include [Y strategy], with telehealth as a potential solution."

### **4. Future Directions & Emerging Research**  
(Adapt title and content if not a research-oriented question. E.g., "Further Exploration & Related Topics" or "Next Steps & Potential Enhancements")
- Identify **gaps in current knowledge** and areas requiring further study/exploration.  
- Highlight **ongoing or upcoming developments** (e.g., trials, technologies, trends) that could impact the topic.  
- Mention **emerging therapies, diagnostics, or evolving guidelines/approaches.**  
- DO NOT include direct URLs or links in this section, as they may be outdated or incorrect.
- If references would be helpful, format them as Google search links: [Search for topic](https://www.google.com/search?q=search+terms+here) or Google Scholar links: [Academic research on topic](https://scholar.google.com/scholar?q=search+terms+here)

   **Example Format (if applicable):**  
   - "Ongoing phase III trials of [novel therapy] may redefine treatment guidelines. Additionally, studies on [biomarker X] could improve risk stratification in [condition]."

---
## **Formatting and Clarity Guidelines**
- Use precise, technical language. If the question is medical, language appropriate for healthcare professionals.
- **Do not include unnecessary disclaimers.** Clearly indicate evidence strength when needed.
- Ensure a **clear, structured response** without redundancy.
- Format with **Markdown** to display well in a Streamlit app.
- **IMPORTANT: Begin your response *exactly* with the following line (using four hash symbols):**
   `#### 1. Best Answer from Retrieved Context`
"""
        final_summary = "Could not generate summary."
        try:
            with st.spinner("🤖 Generating final summary..."):
                summary_response = llm.invoke(summary_prompt_text)
            final_summary = (
                summary_response.content
                if hasattr(summary_response, "content")
                else str(summary_response)
            )
        except Exception as e_summ:
            final_summary = f"Error generating summary: {e_summ}"

        # Store all results in session state using Pydantic model
        # Ensure current_history (list of dicts) is correctly converted to List[AnalysisIteration]
        history_objects = [
            AnalysisIteration(**item)
            for item in current_history
            if isinstance(item, dict)
        ]

        # If no images were captured in final_images but we have them in all_analysis_images,
        # use those instead to ensure images are displayed
        images_to_store = final_images
        if not images_to_store and st.session_state.all_analysis_images:
            print(
                f"No images in final_images, but found {len(st.session_state.all_analysis_images)} in all_analysis_images"
            )
            images_to_store = st.session_state.all_analysis_images

        analysis_results = AnalysisResults(
            question=user_question,
            summary=final_summary,
            code=final_code,
            output=final_output,
            images=images_to_store,  # This now contains base64 strings
            error=final_error,
            history=history_objects,
            mappings=st.session_state.get("categorical_mappings", {}),
            # pubmed_search_terms, pubmed_articles, is_medical_question will be set if applicable
            # by other parts of the logic or can be added here if available at this point.
        )
        # Add PubMed related info if available from session state (set during medical question processing)
        if st.session_state.get("is_medical_question"):
            analysis_results.is_medical_question = True
            analysis_results.pubmed_search_terms = st.session_state.get(
                "pubmed_search_terms"
            )
            analysis_results.pubmed_articles = st.session_state.get("articles", [])

        st.session_state.analysis_results = analysis_results.model_dump()
        # Update the main working dfs in session state with the final state from analysis
        st.session_state.gpt_working_dfs = current_dfs_for_analysis

    # --- Display Analysis Results ---
    results_data = st.session_state.get("analysis_results", None)  # This is now a dictionary
    if results_data and results_data.get(
        "question"
    ):  # Check if results_data and question exist
        with main_col1:
            st.markdown("---")
            st.subheader("📊 Analysis Results")
            st.markdown(f"**Asked:** *{results_data['question']}*")
            st.markdown("**Summary from AI:**")
            expected_summary_start = "#### 1. Best Answer from Retrieved Context"
            summary_text = results_data.get("summary", "").strip()
            if not summary_text or summary_text.lower() in [
                "from the ai's own knowledge",
                "from the ai’s own knowledge",
                "no answer generated.",
                "no output.",
                "no summary.",
            ]:
                st.warning(
                    "⚠️ The AI was unable to generate a main answer for your question. This may be due to insufficient context, a code generation issue, or a model error. Please try rephrasing your question or check the Intermediate Iteration Details for more information."
                )
            else:
                section1, remainder = extract_section1(summary_text)
                if section1:
                    st.info(section1)
                if remainder:
                    st.warning(remainder)
            if results_data.get("error"):
                st.error(
                    f"**Error during execution:**\n```\n{results_data['error']}\n```"
                )
            if results_data["images"]:
                st.markdown("**Generated Plot(s):**")
                for i, img_path in enumerate(results_data["images"]):
                    if not img_path:  # Skip if path is empty
                        st.warning(f"Skipping empty image data for plot {i + 1}.")
                        continue

                    # Check if it's a string path or still a base64 string (for backward compatibility)
                    if isinstance(img_path, str) and os.path.exists(img_path):
                        # It's a file path - use Streamlit's native image display
                        st.image(
                            img_path,
                            caption=f"plot_iter{i + 1}.png",
                            use_container_width=True,
                        )
                        print(f"Successfully displayed image from path: {img_path}")
                    elif isinstance(img_path, str) and img_path.startswith(
                        ("data:", "http")
                    ):
                        # It's a data URL or web URL
                        st.image(
                            img_path,
                            caption=f"plot_iter{i + 1}.png",
                            use_container_width=True,
                        )
                        print("Successfully displayed image from URL")
                    elif isinstance(img_path, str) and len(img_path) > 100:
                        # It might be a base64 string (for backward compatibility)
                        try:
                            import base64

                            img_bytes = base64.b64decode(img_path)
                            st.image(
                                img_bytes,
                                caption=f"plot_iter{i + 1}.png",
                                use_container_width=True,
                            )
                            print("Successfully displayed image from base64 data")
                        except Exception as e:
                            st.warning(f"Could not display image {i + 1}: {e}")
                    else:
                        st.warning(f"Unknown image format for plot {i + 1}")

        with main_col2:
            # --- Display Web Search Results in Expanders ---
            if results_data["output"]:
                search_sections = []
                lines = results_data["output"].splitlines()
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    if (
                        line.startswith("Search results for:")
                        and "No results found." not in line
                    ):
                        section_lines = [lines[i]]
                        j = i + 1
                        while j < len(lines) and not lines[j].strip().startswith(
                            "Retrieved documents for query:"
                        ):
                            section_lines.append(lines[j])
                            j += 1
                        search_sections.append("\n".join(section_lines))
                        i = j - 1
                    i += 1
                if search_sections:
                    with st.expander("🔎 Web Search Results", expanded=False):
                        for idx, section in enumerate(search_sections):
                            if idx > 0:
                                st.markdown("---")
                            st.markdown(section)
            
            # Display all generated images in a dedicated expander
            with st.expander("🖼️ All Generated Images (All Iterations)", expanded=False):
                if st.session_state.get("all_analysis_images"):
                    st.success(
                        f"Found {len(st.session_state.all_analysis_images)} images"
                    )

                    # Create a grid layout for images
                    cols = st.columns(2)  # Display images in 2 columns

                    for i, img_item in enumerate(st.session_state.all_analysis_images):
                        if not img_item:
                            cols[i % 2].warning(f"Skipping empty image {i + 1}")
                            continue

                        col = cols[i % 2]  # Alternate between columns

                        # Check if it's a file path or base64 string
                        if isinstance(img_item, str) and os.path.exists(img_item):
                            # It's a file path
                            col.image(
                                img_item,
                                caption=f"Image {i + 1}",
                                use_container_width=True,
                            )
                            file_size = os.path.getsize(img_item)
                            col.caption(f"Size: {file_size:,} bytes")
                        elif isinstance(img_item, str) and len(img_item) > 100:
                            # It might be a base64 string (for backward compatibility)
                            try:
                                import base64

                                img_bytes = base64.b64decode(img_item)
                                col.image(
                                    img_bytes,
                                    caption=f"Image {i + 1}",
                                    use_container_width=True,
                                )
                                col.caption(f"Size: {len(img_bytes):,} bytes")
                            except Exception as e:
                                col.warning(f"Image {i + 1} error: {str(e)}")
                        else:
                            col.warning(f"Unknown image format for image {i + 1}")
                else:
                    st.info("No images were generated during this analysis.")
                    

            # --- Display PubMed Results if Available ---
            if st.session_state.get("pubmed_search_terms") and st.session_state.get("articles"):
                with st.expander("📚 PubMed Search Results", expanded=False):
                    pubmed_link = (
                        "https://pubmed.ncbi.nlm.nih.gov/?term="
                        + st.session_state.get("pubmed_search_terms", "")
                    )
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.subheader("PubMed Literature Search")
                        st.write(
                            f"Found {len(st.session_state.articles)} relevant articles"
                        )
                    with col2:
                        st.page_link(pubmed_link, label="View in PubMed", icon="📚")
                    with st.popover("PubMed Search Strategy"):
                        st.code(st.session_state.pubmed_search_terms, language="text")
                        st.write(
                            "This search strategy was automatically generated based on your question."
                        )
                    if st.session_state.older_pubmed_articles_alert:
                        st.warning(
                            "Note: The relevant PubMed references identified are primarily older (extending up to 15 years back). "
                            "Please use the PubMed search link to ensure no more recent articles exist."
                        )
                    for i, article in enumerate(st.session_state.articles):
                        with st.container():
                            st.markdown(
                                f"### {i + 1}. [{article['title']}]({article['link']})"
                            )
                            st.markdown(f"**Publication Year:** {article['year']}")
                            st.markdown(
                                f"<details><summary>View Abstract</summary><p>{article['abstract'] if article['abstract'] else 'No abstract available'}</p></details>",
                                unsafe_allow_html=True,
                            )
                            if i < len(st.session_state.articles) - 1:
                                st.divider()
                    st.info(f"""
                    **Search Summary:**
                    - Query executed on: {datetime.now().strftime("%Y-%m-%d")}
                    - Years searched: {st.session_state.get("years_back", 4)} years back
                    - Articles found: {len(st.session_state.articles)}
                    - Search strategy: {"Cutting-edge research" if st.session_state.get("cutting_edge", False) else "Consensus/review articles"}
                    """)

            # --- Display All Retrieved Chunks in a Dedicated Expander ---
            all_chunks = []
            if st.session_state.retrieved_chunks:
                all_chunks.extend(st.session_state.retrieved_chunks)
            if results_data["history"]:
                for item_dict in results_data["history"]:
                    if "chunks" in item_dict and item_dict["chunks"]:
                        for chunk_text_item in item_dict["chunks"]:
                            if chunk_text_item not in all_chunks:
                                all_chunks.append(chunk_text_item)
                    if (
                        "output" in item_dict
                        and "Retrieved documents for query:" in item_dict["output"]
                    ):
                        output_text = item_dict["output"]
                        lines = output_text.splitlines()
                        i = 0
                        while i < len(lines):
                            line = lines[i].strip()
                            if line.startswith("Retrieved documents for query:"):
                                section_start = i
                                j = i + 1
                                while j < len(lines) and not lines[
                                    j
                                ].strip().startswith("Search results for:"):
                                    j += 1
                                retrieval_section = "\n".join(lines[section_start:j])
                                if retrieval_section not in all_chunks:
                                    all_chunks.append(retrieval_section)
                                i = j - 1
                            i += 1
            if all_chunks:
                with st.expander("📚 Retrieved Document Chunks", expanded=False):
                    for idx, section in enumerate(all_chunks):
                        if idx > 0:
                            st.markdown(
                                "<hr style='border: 2px solid #f0f2f6; margin: 30px 0;'>",
                                unsafe_allow_html=True,
                            )
                        formatted_lines = []
                        in_combined_snippets = False
                        for line in section.splitlines():
                            if line.strip().startswith(
                                "Retrieved documents from UPLOADED FILES for query:"
                            ) or line.strip().startswith(
                                "Retrieved documents for query:"
                            ):
                                formatted_lines.append("### " + line)
                            elif line.strip().startswith("DOCUMENTS FOUND:"):
                                formatted_lines.append("### " + line)
                            elif line.strip().startswith("Combined Snippets:"):
                                in_combined_snippets = True
                                formatted_lines.append("\n## Combined Snippets")
                            elif in_combined_snippets and line.strip().startswith(
                                "   Sources:"
                            ):
                                formatted_lines.append("### Sources:")
                            elif in_combined_snippets and line.strip().startswith(
                                "   - "
                            ):
                                formatted_lines.append(line)
                            elif in_combined_snippets and line.strip().startswith(
                                "   Content:"
                            ):
                                formatted_lines.append("### Content:")
                            elif in_combined_snippets and line.strip().startswith(
                                "   From "
                            ):
                                formatted_lines.append("> " + line.strip())
                            elif line.strip().startswith(
                                tuple(str(i) + "." for i in range(1, 11))
                            ):
                                in_combined_snippets = False
                                formatted_lines.append("\n**" + line + "**")
                            elif line.strip().startswith("   Source:"):
                                formatted_lines.append("* " + line.strip())
                            elif line.strip().startswith("   Title:"):
                                formatted_lines.append("* " + line.strip())
                            elif line.strip().startswith("   Content:"):
                                content = line.replace("   Content:", "").strip()
                                formatted_lines.append("* **Content:**")
                                formatted_lines.append("> " + content)
                            elif line.strip() == "":
                                if not in_combined_snippets:
                                    formatted_lines.append("\n---\n")
                            else:
                                formatted_lines.append(line)
                        st.markdown("\n".join(formatted_lines))
            elif "vector store" in results_data["output"].lower():
                with st.expander("📚 Retrieved Document Chunks", expanded=True):
                    st.info(
                        "Vector store was created but no specific chunks were retrieved. This might be due to the search not finding relevant documents or an error in the retrieval process."
                    )
            elif "Content:" in results_data["output"] and (
                "Source:" in results_data["output"]
                or "Title:" in results_data["output"]
            ):
                with st.expander("📚 Retrieved Document Chunks", expanded=True):
                    output_text = results_data["output"]
                    retrieval_sections = []
                    lines = output_text.splitlines()
                    i = 0
                    while i < len(lines):
                        line = lines[i].strip()
                        if line.startswith("Retrieved documents") or line.startswith(
                            "DOCUMENTS FOUND:"
                        ):
                            section_start = i
                            j = i + 1
                            while j < len(lines) and not (
                                lines[j].strip().startswith("Search results for:")
                                or lines[j].strip().startswith("Retrieved documents")
                            ):
                                j += 1
                            section_text = "\n".join(lines[section_start:j])
                            retrieval_sections.append(section_text)
                            i = j - 1
                        i += 1
                    if retrieval_sections:
                        for idx, section in enumerate(retrieval_sections):
                            if idx > 0:
                                st.markdown(
                                    "<hr style='border: 2px solid #f0f2f6; margin: 30px 0;'>",
                                    unsafe_allow_html=True,
                                )
                            formatted_lines = []
                            for line in section.splitlines():
                                if line.strip().startswith(
                                    "Retrieved documents"
                                ) or line.strip().startswith("DOCUMENTS FOUND:"):
                                    formatted_lines.append("### " + line)
                                elif line.strip().startswith(
                                    tuple(str(i) + "." for i in range(1, 11))
                                ):
                                    formatted_lines.append("\n**" + line + "**")
                                elif line.strip().startswith("   Source:"):
                                    formatted_lines.append("* " + line.strip())
                                elif line.strip().startswith("   Title:"):
                                    formatted_lines.append("* " + line.strip())
                                elif line.strip().startswith("   Content:"):
                                    content = line.replace("   Content:", "").strip()
                                    formatted_lines.append("* **Content:**")
                                    formatted_lines.append("> " + content)
                                elif line.strip() == "":
                                    formatted_lines.append("\n---\n")
                                else:
                                    formatted_lines.append(line)
                            st.markdown("\n".join(formatted_lines))
                    else:
                        st.info(
                            "Document chunks may be present in the output but couldn't be properly formatted for display."
                        )

            with st.expander("Show Final Code & Output Details", expanded=False):
                st.markdown("**Final Executed Python Code:**")
                st.code(results_data["code"], language="python")
                st.markdown("**Full Textual Output:**")
                st.code(results_data["output"], language="text")
                if results_data.get("mappings"):
                    st.markdown("**Categorical Variable Mappings Applied:**")
                    for col, mapping_info in results_data["mappings"].items():
                        st.text(f"  Column '{col}': {mapping_info}")

            with st.expander("Show Iteration History", expanded=False):
                if results_data["history"]:
                    # Create tabs for each iteration instead of nested expanders
                    iteration_tabs = st.tabs([f"Iteration {item['iteration']}" for item in results_data["history"]])
                    
                    for i, tab in enumerate(iteration_tabs):
                        item = results_data["history"][i]
                        with tab:
                            st.markdown(f"--- **Iteration {item['iteration']}** ---")
                            st.markdown("*Code:*")
                            st.code(item["code"], language="python")
                            st.markdown("*Output:*")
                            st.code(item["output"], language="text")
                            if item["error"]:
                                st.error(f"*Error:*\n```\n{item['error']}\n```")
                            output_text = item["output"]
                            if "Retrieved documents for query:" in output_text:
                                # Use a collapsible container with a visible header instead of an expander
                                st.markdown(
                                    f"<details><summary>📚 <b>Retrieved Chunks (Iteration {item['iteration']})</b></summary>",
                                    unsafe_allow_html=True,
                                )
                                retrieval_lines = []
                                in_retrieval = False
                                in_combined_snippets = False
                                for line in output_text.splitlines():
                                    if line.strip().startswith(
                                        "Retrieved documents for query:"
                                    ):
                                        in_retrieval = True
                                        retrieval_lines.append("### " + line)
                                    elif in_retrieval:
                                        if line.strip().startswith(
                                            "Combined Snippets:"
                                        ):
                                            in_combined_snippets = True
                                            retrieval_lines.append(
                                                "\n## Combined Snippets"
                                            )
                                        elif (
                                            in_combined_snippets
                                            and line.strip().startswith("   Sources:")
                                        ):
                                            retrieval_lines.append("### Sources:")
                                        elif (
                                            in_combined_snippets
                                            and line.strip().startswith("   - ")
                                        ):
                                            retrieval_lines.append(line)
                                        elif (
                                            in_combined_snippets
                                            and line.strip().startswith("   Content:")
                                        ):
                                            retrieval_lines.append("### Content:")
                                        elif (
                                            in_combined_snippets
                                            and line.strip().startswith("   From ")
                                        ):
                                            retrieval_lines.append("> " + line.strip())
                                        elif line.strip().startswith(
                                            tuple(str(i) + "." for i in range(1, 11))
                                        ):
                                            in_combined_snippets = False
                                            retrieval_lines.append("\n**" + line + "**")
                                        elif line.strip().startswith("   Source:"):
                                            retrieval_lines.append("* " + line.strip())
                                        elif line.strip().startswith("   Title:"):
                                            retrieval_lines.append("* " + line.strip())
                                        elif line.strip().startswith("   Content:"):
                                            content = line.replace(
                                                "   Content:", ""
                                            ).strip()
                                            retrieval_lines.append("* **Content:**")
                                            retrieval_lines.append("> " + content)
                                        elif line.strip() == "":
                                            if not in_combined_snippets:
                                                retrieval_lines.append("\n---\n")
                                        else:
                                            retrieval_lines.append(line)
                                st.markdown("\n".join(retrieval_lines))
                                st.markdown("</details>", unsafe_allow_html=True)
                            if item["images"]:
                                st.markdown("*Image(s) from this iteration:*")
                                for i_img, img_item in enumerate(item["images"]):
                                    if not img_item:
                                        st.warning(
                                            f"Skipping empty image data for plot {i_img + 1} in iteration history."
                                        )
                                        continue

                                    # Check if it's a file path or base64 string
                                    if isinstance(img_item, str) and os.path.exists(
                                        img_item
                                    ):
                                        # It's a file path
                                        st.image(
                                            img_item,
                                            caption=f"Iter {item['iteration']}: plot_iter{i_img + 1}.png",
                                            width=300,
                                        )
                                    elif (
                                        isinstance(img_item, str)
                                        and len(img_item) > 100
                                    ):
                                        # It might be a base64 string (for backward compatibility)
                                        try:
                                            import base64

                                            img_bytes = base64.b64decode(img_item)
                                            st.image(
                                                img_bytes,
                                                caption=f"Iter {item['iteration']}: plot_iter{i_img + 1}.png",
                                                width=300,
                                            )
                                        except Exception as e:
                                            st.warning(
                                                f"Could not display image from iteration {item['iteration']}: {e}"
                                            )
                                    else:
                                        st.warning(
                                            f"Unknown image format for image in iteration {item['iteration']}"
                                        )
                else:
                    st.info(
                        "No iteration history recorded (likely only one iteration or an early error)."
                    )

            # --- RAGAS Validation and Follow-up Questions Section ---
            if "ragas_model" not in st.session_state:
                st.session_state.ragas_model = "gpt-4o-mini"

            if "validate_section1" not in st.session_state:
                st.session_state.validate_section1 = False

            if st.button(
                f"Validate Response Section 1 against Sources (using {st.session_state.ragas_model})",
                help="The Faithfulness Score = Number of claims supported by the sources / Total number of claims in the response",
                key="hallucination_check_button",
            ):
                st.session_state.validate_section1 = True

            # --- Helper to collect all retrieved chunks from session and history ---
            # (Already defined above for summary; use the same function here)

            if st.session_state.validate_section1:
                st.session_state.validate_section1 = False
                # Import necessary libraries for RAGAS
                try:
                    from ragas import SingleTurnSample
                    from ragas.metrics import Faithfulness
                    from ragas.metrics import AspectCritic
                    from ragas.metrics import RubricsScore
                    from ragas.llms import LangchainLLMWrapper
                    from ragas.embeddings import LangchainEmbeddingsWrapper
                    from langchain_openai import ChatOpenAI
                    from langchain_openai import OpenAIEmbeddings
                except ImportError:
                    st.error(
                        "RAGAS library not installed. Please install it with: pip install ragas"
                    )
                    st.stop()

                # Extract Section 1 from the response
                if "current_section1" not in st.session_state:
                    st.session_state.current_section1 = ""
                if "validated_section1" not in st.session_state:
                    st.session_state.validated_section1 = ""
                if "validation_results" not in st.session_state:
                    st.session_state.validation_results = None

                # Standardize Section 1 extraction using the helper function
                section1, _ = extract_section1(results_data["summary"])

                # --- Gather retrieved document chunks for RAGAS context ---
                # Always use the full set of retrieved document chunks for both standard and custom RAGAS assessment

                # Use robust helper to collect all retrieved chunks
                retrieved_chunks = get_all_retrieved_chunks(results_data)

                # Compose the context for RAGAS as a single string (for reference) and as a list (for faithfulness)
                citations_text = (
                    "\n\n---\n\n".join(retrieved_chunks) if retrieved_chunks else ""
                )
                # st.write(f'**Retrieved Context for RAGAS Validation:**')
                # st.code(citations_text, language="text")
                retrieved_contexts_list = (
                    [chunk for chunk in retrieved_chunks] if retrieved_chunks else []
                )
                # st.write("**Retrieved Contexts List for Faithfulness Evaluation:**")
                # st.code(retrieved_contexts_list, language="python")
                

                # Initialize variables to default values
                current_rubric_score = 0
                current_faithfulness_score = 0.0

                # Check if we've already validated this exact section1 content and context
                cache_key = (section1, tuple(retrieved_chunks), st.session_state.get("ragas_model", "gpt-4o-mini"))
                if (
                    section1 == st.session_state.validated_section1
                    and st.session_state.validation_results is not None
                    and st.session_state.validation_results.get("cache_key")
                    == cache_key
                ):
                    st.info(
                        "Using cached validation results since content and context haven't changed."
                    )
                    # Display the cached validation results
                    current_rubric_score = st.session_state.validation_results[
                        "rubric_score"
                    ]
                    current_faithfulness_score = st.session_state.validation_results[
                        "faithfulness_score"
                    ]
                    # Skip to displaying results
                else:
                    # New content/context to validate or first validation
                    st.session_state.validated_section1 = section1

                    # Prepare the sample for RAGAS using the retrieved document chunks as context
                    sample = SingleTurnSample(
                        response=section1,
                        reference=citations_text,
                        query=results_data["question"],
                    )

                    sample_faithfulness = SingleTurnSample(
                        user_input=results_data["question"],
                        response=section1,
                        retrieved_contexts=retrieved_contexts_list,
                    )

                    # Define rubrics for evaluation
                    rubrics = {
                        "score1_description": "There is no hallucination in the response. The response is fully supported by the reference.",
                        "score2_description": "Factual statements are supported by the reference but the response is not fully accurate and lacks important details.",
                        "score3_description": "There are some factual statements that are not present in the reference.",
                        "score4_description": "The response contains some factual errors and lacks important details based on the reference.",
                        "score5_description": "The model adds new information and statements that contradict the reference.",
                    }

                    # Initialize the evaluator
                    ragas_model_name = st.session_state.get(
                        "ragas_model", "gpt-4o-mini"
                    )  # Renamed to avoid conflict
                    evaluator_llm = LangchainLLMWrapper(
                        ChatOpenAI(model=ragas_model_name)
                    )
                    evaluator_embeddings = LangchainEmbeddingsWrapper(
                        OpenAIEmbeddings()
                    )  # Assumes OPENAI_API_KEY is set
                    scorer = RubricsScore(rubrics=rubrics, llm=evaluator_llm)
                    scorer_faithfulness = Faithfulness(llm=evaluator_llm)

                    # Define async function to evaluate metrics
                    async def evaluate_ragas_metrics():
                        # Call internal _single_turn_ascore to get full result objects
                        # Pass empty list for callbacks as it's expected by the internal methods
                        rubric_result_obj = await scorer._single_turn_ascore(
                            sample, callbacks=[]
                        )
                        faithfulness_result_obj = (
                            await scorer_faithfulness._single_turn_ascore(
                                sample_faithfulness, callbacks=[]
                            )
                        )
                        return rubric_result_obj, faithfulness_result_obj

                    # Run the evaluation
                    with st.spinner("Evaluating response against sources..."):
                        direct_rubric_score, direct_faithfulness_score = asyncio.run(
                            evaluate_ragas_metrics()
                        )

                        current_rubric_score = int(direct_rubric_score)
                        current_faithfulness_score = float(direct_faithfulness_score)

                        # Cache the results
                        st.session_state.validation_results = {
                            "rubric_score": current_rubric_score,
                            "faithfulness_score": current_faithfulness_score,
                            "cache_key": cache_key,
                        }

                # Define rubrics here to ensure it's always available
                rubrics = {
                    "score1_description": "There is no hallucination in the response. The response is fully supported by the reference.",
                    "score2_description": "Factual statements are supported by the reference but the response is not fully accurate and lacks important details.",
                    "score3_description": "There are some factual statements that are not present in the reference.",
                    "score4_description": "The response contains some factual errors and lacks important details based on the reference.",
                    "score5_description": "The model adds new information and statements that contradict the reference.",
                }

                # --- DEBUG: Show what is being sent to RAGAS ---
                # st.markdown("#### [DEBUG] Retrieved Chunks/Context Sent to RAGAS")
                # st.code(citations_text, language="text")
                # --- END DEBUG ---

                # Display evaluation results
                st.markdown("### RAGAS Library Evaluation Results")
                if current_rubric_score == 1:
                    st.success(
                        "Section 1 is supported by the sources (RAGAS Rubric Score: 1)."
                    )
                elif current_rubric_score == 2:
                    st.error(
                        "Caution: Factual statements supported, but Section 1 may lack accuracy/details (RAGAS Rubric Score: 2). Confirm with references."
                    )
                elif current_rubric_score == 3:
                    st.warning(
                        "Caution: Some factual statements in Section 1 may not be fully supported (RAGAS Rubric Score: 3). Confirm with references."
                    )
                elif current_rubric_score == 4:
                    st.warning(
                        "Warning: Section 1 may contain factual errors/lack details (RAGAS Rubric Score: 4). Confirm with references."
                    )
                elif current_rubric_score == 5:
                    st.error(
                        "Warning!!! Section 1 may add new information contradicting sources (RAGAS Rubric Score: 5). Confirm with references."
                    )
                else:
                    st.error(
                        f"Error: Unable to evaluate the response based on rubrics (RAGAS Rubric Score: {current_rubric_score})."
                    )

                if current_faithfulness_score > 0.9:
                    st.success(
                        f"**RAGAS Faithfulness Score:** {current_faithfulness_score:.3f} (High confidence in factual consistency with sources)."
                    )
                else:
                    st.warning(
                        f"**RAGAS Faithfulness Score:** {current_faithfulness_score:.3f}. Review carefully, some assertions might not be fully backed by provided sources."
                    )

                # Detailed RAGAS results expander
                with st.expander("View RAGAS Evaluation Details"):
                    st.subheader("Rubric Score Details")
                    st.markdown(
                        f"**Overall RAGAS Rubric Score:** {current_rubric_score}"
                    )
                    # Derive reason from the rubrics dictionary
                    rubric_reason_key = f"score{current_rubric_score}_description"
                    rubric_reason = rubrics.get(
                        rubric_reason_key, "Specific reason not found for this score."
                    )
                    st.markdown("**Reasoning (from predefined rubrics):**")
                    st.markdown(rubric_reason)

                    st.divider()

                    st.subheader("Faithfulness Score Details")
                    st.markdown(
                        f"**Overall RAGAS Faithfulness Score:** {current_faithfulness_score:.3f}"
                    )
                    st.markdown(
                        "The Faithfulness Score measures the factual consistency of the generated answer against the provided context. A higher score indicates better alignment."
                    )

                    # Add more detailed explanation of how faithfulness is calculated
                    st.markdown("### How Faithfulness is Calculated")
                    st.markdown("""
                    1. **Statement Generation**: The response is broken down into individual factual statements.
                    2. **Statement Verification**: Each statement is checked against the retrieved context.
                    3. **Verdict Assignment**: Each statement receives a verdict (1 if supported by context, 0 if not).
                    4. **Score Calculation**: Final score = Number of supported statements / Total number of statements.

                    A score closer to 1.0 means most statements in the response are directly supported by the provided sources.
                    A lower score indicates the response contains statements not found in or contradicted by the sources.
                    """)

                    # Add statement breakdown and verdicts
                    st.markdown("### Custom Statement Breakdown")
                    st.info(
                        "This is our own implementation of statement verification, separate from the RAGAS library's internal process."
                    )

                    # Check if we already have cached statement analysis
                    if (
                        section1 == st.session_state.validated_section1
                        and st.session_state.validation_results is not None
                        and "statement_analysis" in st.session_state.validation_results
                    ):
                        st.info(
                            "Using cached statement analysis since content hasn't changed."
                        )

                        # Display the cached statement analysis
                        statement_data = st.session_state.validation_results[
                            "statement_analysis"
                        ]
                        calculated_score = st.session_state.validation_results[
                            "calculated_score"
                        ]
                        supported = st.session_state.validation_results["supported"]
                        total = st.session_state.validation_results["total"]

                        # Display as a DataFrame
                        import pandas as pd

                        df = pd.DataFrame(statement_data)
                        st.dataframe(df, use_container_width=True)

                        # Display the calculated score
                        st.markdown(
                            f"**Calculated score:** {calculated_score:.3f} ({supported} supported statements out of {total} total)"
                        )

                        # Compare with RAGAS score
                        st.markdown(
                            f"**RAGAS faithfulness score:** {current_faithfulness_score:.3f}"
                        )
                        st.markdown(f"**Our calculated score:** {calculated_score:.3f}")
                        if abs(calculated_score - current_faithfulness_score) > 0.2:
                            st.info(
                                "Note: There's a significant difference between our calculated score and the RAGAS score. This could be due to differences in statement extraction or evaluation methods."
                            )
                        else:
                            st.success(
                                "Our calculated score is similar to the RAGAS score, which provides additional confidence in the evaluation."
                            )

                        # Cache the statement analysis results
                        if "validation_results" not in st.session_state:
                            st.session_state.validation_results = {}

                        st.session_state.validation_results.update(
                            {
                                "statement_analysis": statement_data,
                                "calculated_score": calculated_score,
                                "supported": supported,
                                "total": total,
                            }
                        )
                    # If not cached, generate new statement analysis
                    else:
                        # Generate example statements with verdicts for demonstration
                        # In a real implementation, these would come from the RAGAS evaluation
                        with st.spinner("Analyzing statements..."):
                            statements_prompt = [
                                {
                                    "role": "system",
                                    "content": "You are an AI assistant that breaks down text into individual factual statements. For the given text, extract 5-8 key factual claims as a JSON list of strings. Each statement should be self-contained and represent a single factual claim. Return a JSON object with a 'statements' array. If you can't find any statements, include at least 3 general claims from the text.",
                                },
                                {
                                    "role": "user",
                                    "content": f"Extract factual statements from this text:\n\n{section1}",
                                },
                            ]

                        try:
                            client = (
                                OpenAI()
                            )  # Assumes OPENAI_API_KEY is set in environment/secrets
                            ragas_model_name_for_statements = st.session_state.get(
                                "ragas_model", "gpt-4o-mini"
                            )
                            statements_response = client.chat.completions.create(
                                model=ragas_model_name_for_statements,
                                messages=statements_prompt,
                                response_format={"type": "json_object"},
                                temperature=0.1,
                            )
                            statements_json = statements_response.choices[
                                0
                            ].message.content

                            try:
                                parsed_json = json.loads(statements_json)
                                statements = parsed_json.get("statements", [])

                                # Fallback if no statements were extracted
                                if not statements:
                                    st.warning(
                                        "No specific factual statements were identified. Using general statements from the text."
                                    )
                                    # Create some basic statements from the text
                                    statements = [
                                        f"The text discusses {results_data['question']}",
                                        "The response provides information based on available sources",
                                        "The text contains factual content related to the query",
                                    ]
                            except json.JSONDecodeError:
                                st.warning(
                                    "Could not parse JSON response. Using general statements."
                                )
                                statements = [
                                    f"The text discusses {results_data['question']}",
                                    "The response provides information based on available sources",
                                    "The text contains factual content related to the query",
                                ]

                            # For each statement, determine if it's supported by the context
                            verdicts = []
                            # Use the full set of retrieved document chunks as the context for each statement
                            full_chunks_context = (
                                "\n\n---\n\n".join(retrieved_chunks)
                                if retrieved_chunks
                                else ""
                            )
                            for statement in statements:
                                try:
                                    verdict_prompt = [
                                        {
                                            "role": "system",
                                            "content": "You are an AI assistant that determines if a statement is supported by the provided context. Return a JSON object with keys 'verdict' (1 if supported, 0 if not) and 'reason' (brief explanation). A statement is supported if the information it contains can be directly found in or reasonably inferred from the context.",
                                        },
                                        {
                                            "role": "user",
                                            "content": f"Question: {results_data['question']}\n\nStatement to verify: {statement}\n\nContext from sources:\n{full_chunks_context}\n\nIs this statement supported by the context? Return only a JSON object.",
                                        },
                                    ]
                                    # client = OpenAI() # Already initialized
                                    verdict_response = client.chat.completions.create(
                                        model=ragas_model_name_for_statements,
                                        messages=verdict_prompt,
                                        response_format={"type": "json_object"},
                                        temperature=0.1,
                                    )
                                    verdict_json = verdict_response.choices[
                                        0
                                    ].message.content
                                    try:
                                        verdict_data = json.loads(verdict_json)
                                        verdict = verdict_data.get("verdict", 0)
                                        # Ensure verdict is either 0 or 1
                                        if not isinstance(
                                            verdict, int
                                        ) or verdict not in [0, 1]:
                                            if isinstance(
                                                verdict, str
                                            ) and verdict.lower() in [
                                                "true",
                                                "yes",
                                                "supported",
                                            ]:
                                                verdict = 1
                                            elif isinstance(
                                                verdict, str
                                            ) and verdict.lower() in [
                                                "false",
                                                "no",
                                                "not supported",
                                            ]:
                                                verdict = 0
                                            else:
                                                verdict = 0
                                        verdicts.append(
                                            {
                                                "statement": statement,
                                                "verdict": verdict,
                                                "reason": verdict_data.get(
                                                    "reason", "No reason provided"
                                                ),
                                            }
                                        )
                                    except json.JSONDecodeError:
                                        st.warning(
                                            f"Could not parse verdict JSON for statement: {statement[:50]}..."
                                        )
                                        verdicts.append(
                                            {
                                                "statement": statement,
                                                "verdict": 0,
                                                "reason": "Error processing verdict",
                                            }
                                        )
                                except Exception as e:
                                    st.warning(f"Error evaluating statement: {str(e)}")
                                    verdicts.append(
                                        {
                                            "statement": statement,
                                            "verdict": 0,
                                            "reason": f"Error during evaluation: {str(e)}",
                                        }
                                    )

                            # Display statements and verdicts in a flat list (no nested expanders)
                            if verdicts:
                                st.markdown("The following statements were evaluated:")

                                # Create a table for all statements
                                statement_data = []
                                for i, v_item in enumerate(
                                    verdicts, 1
                                ):  # Renamed v to v_item
                                    verdict_icon = (
                                        "✅" if v_item["verdict"] == 1 else "❌"
                                    )
                                    statement_data.append(
                                        {
                                            "Statement": f"{verdict_icon} {v_item['statement']}",
                                            "Verdict": "Supported"
                                            if v_item["verdict"] == 1
                                            else "Not supported",
                                            "Reason": v_item["reason"],
                                        }
                                    )

                                # Display as a DataFrame
                                import pandas as pd

                                df = pd.DataFrame(statement_data)
                                st.dataframe(df, use_container_width=True)

                                # Calculate and display the faithfulness score based on these verdicts
                                supported = sum(
                                    1 for v_item in verdicts if v_item["verdict"] == 1
                                )
                                total = len(verdicts)
                                calculated_score = supported / total if total > 0 else 0
                                st.markdown(
                                    f"**Calculated score:** {calculated_score:.3f} ({supported} supported statements out of {total} total)"
                                )

                                # Compare with RAGAS score
                                st.markdown(
                                    f"**RAGAS faithfulness score:** {current_faithfulness_score:.3f}"
                                )
                                st.markdown(
                                    f"**Our calculated score:** {calculated_score:.3f}"
                                )
                                if (
                                    abs(calculated_score - current_faithfulness_score)
                                    > 0.2
                                ):
                                    st.info(
                                        "Note: There's a significant difference between our calculated score and the RAGAS score. This could be due to differences in statement extraction or evaluation methods."
                                    )
                                else:
                                    st.success(
                                        "Our calculated score is similar to the RAGAS score, which provides additional confidence in the evaluation."
                                    )

                                # Cache the statement analysis results
                                if "validation_results" not in st.session_state:
                                    st.session_state.validation_results = {}

                                st.session_state.validation_results.update(
                                    {
                                        "statement_analysis": statement_data,
                                        "calculated_score": calculated_score,
                                        "supported": supported,
                                        "total": total,
                                    }
                                )
                            else:
                                st.error(
                                    "No statements were successfully evaluated. Please try again."
                                )

                        except Exception as e:
                            st.error(f"Error analyzing statements: {str(e)}")
                            st.markdown(
                                "Unable to display statement breakdown. Please try again later."
                            )

                    # Add explanation about limitations
                    st.markdown("### Limitations")
                    st.markdown("""
                    - The metric evaluates factual consistency, not completeness or relevance.
                    - Complex or nuanced statements may be difficult to verify automatically.
                    - The evaluation depends on the quality of the statement breakdown process.
                    - Statements that are common knowledge but not in the sources may be marked as unfaithful.
                    - The statement extraction and verification process shown here is a demonstration and may differ from RAGAS's internal implementation.
                    """)

            if DOCX_AVAILABLE:
                if st.button(
                    "📄 Generate Word Document from Analysis", key="docx_button"
                ):
                    with st.spinner("Creating Word document..."):
                        docx_file_path = generate_gpt_analysis_docx(
                            file_name="ai_iterative_analysis",
                            question=results_data["question"],
                            research_summary=results_data["summary"],
                            code=results_data["code"],
                            output=results_data["output"],
                            image_paths=results_data[
                                "images"
                            ],  # Pass base64 encoded images
                            categorical_mappings=results_data.get("mappings"),
                        )
                    if docx_file_path and os.path.exists(docx_file_path):
                        with open(docx_file_path, "rb") as fp:
                            st.download_button(
                                label="⬇️ Download DOCX Report",
                                data=fp,
                                file_name="ai_analysis_report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        try:
                            os.remove(docx_file_path)
                        except Exception as e_del:
                            st.warning(f"Could not delete temporary DOCX file: {e_del}")
                    else:
                        st.error("Failed to generate DOCX file.")
            else:
                st.info(
                    "DOCX generation is disabled as the `markdown_to_docx` module is not available."
                )

            # --- Follow-up Questions UI and Logic ---
            st.markdown("---")  # Visual separator
            initial_followup = st.checkbox(
                "💬 Ask Follow-Up Questions for This Analysis",
                key="initial_followup_checkbox",
            )

            if initial_followup:
                # Display existing follow-up messages
                for msg in st.session_state.followup_chat_thread:
                    if msg["role"] != "system":  # Don't display system prompts
                        with st.chat_message(
                            msg["role"], avatar=role_emojis.get(msg["role"])
                        ):
                            st.markdown(msg["content"])
                if st.session_state.get("google_search_followup_results", ""):
                    with st.expander("🔍 Follow-up Google Search Results", expanded=False):
                        st.markdown(st.session_state.google_search_followup_results)

                if followup_question_text := st.chat_input(
                    "Your follow-up question:", key="followup_chat_input"
                ):
                    # Prepare context for the follow-up system prompt
                    original_question = results_data.get("question", "Not available")
                    initial_summary = results_data.get(
                        "summary", "Initial summary not available."
                    )

                    # Gather all retrieved chunks for evidence
                    all_evidence_chunks = []
                    if st.session_state.retrieved_chunks:
                        all_evidence_chunks.extend(st.session_state.retrieved_chunks)
                    if results_data.get("history"):
                        for item_dict_hist in results_data[
                            "history"
                        ]:  # Renamed item to avoid conflict
                            if isinstance(item_dict_hist, dict) and item_dict_hist.get(
                                "chunks"
                            ):
                                for chunk_text_item_hist in item_dict_hist[
                                    "chunks"
                                ]:  # Renamed to avoid conflict
                                    if chunk_text_item_hist not in all_evidence_chunks:
                                        all_evidence_chunks.append(chunk_text_item_hist)
                    retrieved_chunks_text_for_followup = (
                        "\n\n---\n".join(all_evidence_chunks)
                        if all_evidence_chunks
                        else "No specific document chunks were retrieved for the initial answer."
                    )

                    # Initialize follow-up thread with system prompt if it's the first message
                    if not st.session_state.followup_chat_thread:
                        formatted_system_prompt = (
                            followup_system_prompt_template.format(
                                prior_question=original_question,
                                evidence=retrieved_chunks_text_for_followup,
                                prior_answer=initial_summary,
                            )
                        )
                        st.session_state.followup_chat_thread.append(
                            {"role": "system", "content": formatted_system_prompt}
                        )

                    # Add user's raw follow-up question to history for display
                    st.session_state.followup_chat_thread.append(
                        {"role": "user", "content": followup_question_text}
                    )

                    # Perform Google Search for the follow-up question
                    with st.spinner("Searching Google for additional context..."):
                        from rag_utils import google_search  # Ensure it's imported
                        from models import (
                            SearchResult,
                        )  # Ensure SearchResult is available for type hint

                        def format_google_search_results_for_llm(
                            search_results: List[SearchResult],
                        ) -> str:
                            if not search_results:
                                return "No relevant search results found from Google for the follow-up question."
                            output_lines = [
                                "Fresh Google Search Results for Follow-up Question:"
                            ]
                            for i, res_item in enumerate(
                                search_results
                            ):  # Renamed res to res_item
                                output_lines.append(f"Result {i + 1}:")
                                output_lines.append(f"  Title: {res_item.title}")
                                output_lines.append(f"  Link: {res_item.link}")
                                output_lines.append(f"  Snippet: {res_item.snippet}")
                            return "\n".join(output_lines)

                        # Create a combined query for Google search
                        combined_search_query = (
                            f"{original_question} {followup_question_text}"
                        )
                        try:
                            google_results_list = google_search(
                                combined_search_query,
                                num_results=3,
                                scrape_content=False,
                            )  # Don't scrape for follow-up to keep it fast
                            formatted_google_results = (
                                format_google_search_results_for_llm(
                                    google_results_list
                                )
                            )
                            st.session_state.google_search_followup_results = (
                                formatted_google_results
                            )
                        except Exception as e_gs:
                            st.warning(f"Google search for follow-up failed: {e_gs}")
                            st.session_state.google_search_followup_results = "Google search for follow-up failed or returned no results."

                    # Display Google search results for transparency
                    with st.expander(
                        "🔍 Google Search Results for Follow-up", expanded=False
                    ):
                        st.markdown(st.session_state.google_search_followup_results)

                    # Prepare message for LLM, augmenting the user's question with search results
                    augmented_followup_for_llm = (
                        f"{followup_question_text}\n\n"
                        f"Consider these recent Google search results when formulating your answer:\n"
                        f"{st.session_state.google_search_followup_results}"
                    )

                    # Create the list of messages to send to the LLM
                    # System prompt + all prior user/assistant messages + current augmented user message
                    messages_for_llm = list(
                        st.session_state.followup_chat_thread[:-1]
                    )  # All but the last raw user message
                    messages_for_llm.append(
                        {"role": "user", "content": augmented_followup_for_llm}
                    )

                    # Call Azure LLM
                    with st.spinner("AI is thinking about your follow-up..."):
                        try:
                            llm_instance = get_llm_instance()
                            if llm_instance:
                                response_obj = llm_instance.invoke(messages_for_llm)
                                assistant_response_text = (
                                    response_obj.content
                                    if hasattr(response_obj, "content")
                                    else str(response_obj)
                                )
                            else:
                                assistant_response_text = (
                                    "Error: LLM instance not available."
                                )
                        except Exception as e_llm_followup:
                            assistant_response_text = (
                                f"Error during LLM call for follow-up: {e_llm_followup}"
                            )

                    st.session_state.followup_chat_thread.append(
                        {"role": "assistant", "content": assistant_response_text}
                    )
                    st.rerun()  # Rerun to display the new messages

                # Add download button for the follow-up thread if there's content
                if any(
                    msg["role"] != "system"
                    for msg in st.session_state.followup_chat_thread
                ):
                    original_question_for_docx = results_data.get(
                        "question", "Original question not available."
                    )
                    initial_summary_for_docx = results_data.get(
                        "summary", "Initial summary not available."
                    )

                    # Prepare DOCX data using the new helper function
                    # Pass the global role_emojis dict
                    docx_bytes_io = chat_thread_to_docx(
                        st.session_state.followup_chat_thread,
                        original_question_for_docx,
                        initial_summary_for_docx,
                        role_emojis,
                    )

                    st.download_button(
                        label="⬇️ Download Follow-up Thread (DOCX)",
                        data=docx_bytes_io,
                        file_name=f"followup_chat_{st.session_state.get('current_analysis_timestamp', 'current')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_followup_docx_button",
                    )

            # --- Ask 3 AI Expert Personas ---
            st.markdown("---")
            if st.button(
                "🤔 Ask 3 AI Expert Personas for Opinions", key="ask_experts_button"
            ):
                with st.spinner("Generating expert personas and their opinions..."):
                    llm = get_llm_instance()
                    if llm is None:
                        st.error("LLM not available for expert consultation.")
                    else:
                        # 1. Generate Expert Personas and Rephrased Questions
                        find_experts_messages = [
                            {
                                "role": "system",
                                "content": system_prompt_expert_questions,
                            },
                            {"role": "user", "content": results_data["question"]},
                        ]
                        try:
                            # Forcing JSON-like output through prompt, actual JSON mode might need model/API adjustment
                            # For Azure, the model should be prompted to return JSON.
                            # If the model supports it, `llm.invoke(messages, response_format={"type": "json_object"})`
                            # could be used, but this depends on Langchain and Azure API version.
                            # For now, relying on strong prompting for JSON.
                            experts, domains, expert_questions_list = [], [], []
                            experts_json_response = llm.invoke(find_experts_messages)
                            experts_json_output = (
                                experts_json_response.content
                                if hasattr(experts_json_response, "content")
                                else str(experts_json_response)
                            )

                            try:
                                experts, domains, expert_questions_list = (
                                    extract_expert_info(experts_json_output)
                                )
                            except Exception as e_extract:
                                st.error(f"Failed to parse expert info: {e_extract}")
                                experts, domains, expert_questions_list = [], [], []

                            st.session_state.experts = experts
                            st.session_state.expert_domains = domains
                            st.session_state.expert_questions = expert_questions_list
                            st.session_state.expert_answers = []

                            if not experts or len(experts) < 3:
                                st.error(
                                    "Could not generate 3 distinct expert personas. Please try again."
                                )
                            else:
                                # 2. Gather Context for Experts
                                # Consolidate all retrieved chunks text
                                all_retrieved_chunks_for_experts = []
                                if st.session_state.retrieved_chunks:
                                    all_retrieved_chunks_for_experts.extend(
                                        st.session_state.retrieved_chunks
                                    )
                                if results_data["history"]:
                                    for item_dict_expert_hist in results_data[
                                        "history"
                                    ]:
                                        if isinstance(
                                            item_dict_expert_hist, dict
                                        ) and item_dict_expert_hist.get("chunks"):
                                            for (
                                                chunk_text_item_expert_hist
                                            ) in item_dict_expert_hist["chunks"]:
                                                if (
                                                    chunk_text_item_expert_hist
                                                    not in all_retrieved_chunks_for_experts
                                                ):
                                                    all_retrieved_chunks_for_experts.append(
                                                        chunk_text_item_expert_hist
                                                    )

                                context_for_experts = results_data["summary"]
                                if all_retrieved_chunks_for_experts:
                                    context_for_experts += (
                                        "\n\n--- Additional Retrieved Context ---\n"
                                        + "\n\n---\n".join(
                                            all_retrieved_chunks_for_experts
                                        )
                                    )

                                # 3. Get Opinion from Each Expert
                                expert_prompts_templates = [
                                    expert1_system_prompt,
                                    expert2_system_prompt,
                                    expert3_system_prompt,
                                ]
                                temp_expert_answers = []
                                for i in range(len(st.session_state.experts)):
                                    expert_name = st.session_state.experts[i]
                                    expert_domain = st.session_state.expert_domains[i]
                                    expert_question = st.session_state.expert_questions[
                                        i
                                    ]

                                    expert_system_prompt = expert_prompts_templates[
                                        i
                                    ].format(expert=expert_name, domain=expert_domain)

                                    expert_messages = [
                                        {
                                            "role": "system",
                                            "content": expert_system_prompt,
                                        },
                                        {
                                            "role": "user",
                                            "content": f"Rephrased question for you: {expert_question}\n\nHere's the initial analysis and retrieved context:\n{context_for_experts}",
                                        },
                                    ]

                                    with st.spinner(
                                        f"Waiting for {expert_name} ({expert_domain}) to respond..."
                                    ):
                                        expert_response_obj = llm.invoke(
                                            expert_messages
                                        )
                                        expert_opinion = (
                                            expert_response_obj.content
                                            if hasattr(expert_response_obj, "content")
                                            else str(expert_response_obj)
                                        )
                                        temp_expert_answers.append(expert_opinion)
                                st.session_state.expert_answers = temp_expert_answers
                                st.success("Expert opinions generated!")
                                st.rerun()

                        except Exception as e_expert_gen:
                            st.error(
                                f"Error during expert consultation: {e_expert_gen}"
                            )
                            st.session_state.experts = []
                            st.session_state.expert_answers = []

            # Display Expert Opinions if available
            if (
                st.session_state.experts
                and st.session_state.expert_answers
                and len(st.session_state.experts)
                == len(st.session_state.expert_answers)
            ):
                st.markdown("---")
                st.subheader("🧐 AI Expert Persona Opinions")
                for i in range(len(st.session_state.experts)):
                    expert_name = st.session_state.experts[i]
                    expert_domain = st.session_state.expert_domains[i]
                    expert_opinion = st.session_state.expert_answers[i]
                    with st.expander(
                        f"Perspective from {expert_name} (Expert in {expert_domain})",
                        expanded=False,
                    ):
                        st.markdown(expert_opinion)

                        # Prepare citations for Word doc - using original analysis citations
                        # Attempt to get structured citations if available from the main analysis results
                        citations_for_expert_docx = results_data.get("citations")

                        # Fallback: If structured citations are not directly available in results_data["citations"],
                        # try to construct a list of dictionaries from retrieved_chunks.
                        # This is a simplified approach. Ideally, citations would be more structured.
                        if not citations_for_expert_docx and st.session_state.get(
                            "retrieved_chunks"
                        ):
                            citations_for_expert_docx = []
                            for chunk_text in st.session_state.retrieved_chunks:
                                # Basic parsing: look for a URL or a title-like structure in the chunk
                                # This is very rudimentary and might need refinement based on chunk format
                                url_match = re.search(
                                    r"Source: (https?://[^\s]+)", chunk_text
                                )
                                title_match = re.search(r"Title: ([^\n]+)", chunk_text)

                                citation_entry = {"context": chunk_text}
                                if url_match:
                                    citation_entry["url"] = url_match.group(1)
                                if title_match:
                                    citation_entry["title"] = title_match.group(1)
                                else:  # Fallback title
                                    citation_entry["title"] = (
                                        "Retrieved Context Snippet"
                                    )
                                citations_for_expert_docx.append(citation_entry)
                        elif not citations_for_expert_docx:
                            citations_for_expert_docx = []

                        docx_buffer_expert = markdown_to_word(
                            expert_opinion,
                            original_question_text=st.session_state.expert_questions[
                                i
                            ],  # Use rephrased question
                            citations_list=citations_for_expert_docx,  # Pass the list of citation dicts
                        )
                        st.download_button(
                            label=f"⬇️ Download {expert_name}'s Opinion (DOCX)",
                            data=docx_buffer_expert,
                            file_name=f"expert_opinion_{expert_name.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_expert_{i}_docx",
                        )


# --- App Entry Point ---
if __name__ == "__main__":
    if (
        not os.environ["AZURE_OPENAI_API_KEY"]
        or not os.environ["AZURE_OPENAI_ENDPOINT"]
        or not AZURE_DEPLOYMENT
    ):
        st.error(
            "Azure OpenAI service credentials (API Key, Base URL, Deployment Name) are not configured. Please check secrets/environment variables."
        )
    elif check_password():  # Check password first
        # Check for Tavily API key
        # Check if RAG search is available
        search_available, search_error = get_rag_search_available()
        if not search_available:
            st.warning(
                f"RAG search functionality will be disabled. {search_error}. Set the GOOGLE_API_KEY and GOOGLE_CSE_ID in environment variables or secrets to enable web search."
            )
        else:
            st.sidebar.success("🔍 RAG search is enabled and ready to use")

            # Display current RAG settings in sidebar
            with st.sidebar.expander("Current RAG Settings", expanded=False):
                st.write(
                    f"- Search results: {st.session_state.get('rag_num_results', 5)}"
                )
                st.write(
                    f"- Chunk size: {st.session_state.get('rag_chunk_size', 1000)}"
                )
                st.write(
                    f"- Chunk overlap: {st.session_state.get('rag_chunk_overlap', 200)}"
                )
                st.write(
                    f"- Full content scraping: {'Enabled' if st.session_state.get('rag_scrape_content', True) else 'Disabled'}"
                )
                st.write(
                    f"- Documents retrieved: {st.session_state.get('rag_k_docs', 5)}"
                )
        main_application()
    else:
        st.info("Please enter the password in the sidebar to use the application.")
        # Optionally, you can display a less functional page or just the login prompt.
