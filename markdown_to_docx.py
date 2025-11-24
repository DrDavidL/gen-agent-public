#!/usr/bin/env python3
import os
import tempfile
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import markdown
from bs4 import BeautifulSoup, NavigableString
from docx.oxml.shared import OxmlElement, qn
import base64  # Import base64 for decoding


def write_out_html(file_name, text_html, encoding="utf8"):
    """Write HTML content to a file"""
    try:
        with open(file_name, "w", encoding=encoding) as output_fd:
            output_fd.write(text_html)
    except Exception as e:
        print(f"Could not write HTML file {file_name}: {e}")


def do_table_of_contents(document):
    """Add a table of contents to the document"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:t")
    fld_char3.text = "Right-click to update field."
    fld_char2.append(fld_char3)
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    # Add a style for the hyperlink if desired (e.g., underline, color)
    # Example:
    # rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
    # rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink') # Assumes 'Hyperlink' style is defined
    # rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_code_block_style(doc):
    """Create a style for code blocks"""
    styles = doc.styles
    style_name = "Code Block"
    if style_name not in styles:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = "Courier New"
        font.size = Pt(9)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.right_indent = Inches(0.5)
    return style_name


def enhance_html_for_docx(html_content):
    """Enhance HTML content for better DOCX conversion"""
    soup = BeautifulSoup(html_content, "html.parser")

    # Improve code blocks
    for pre in soup.find_all("pre"):
        code = pre.find("code")
        if code:
            # Add a div with special class for code blocks
            code_div = soup.new_tag("div")
            code_div["class"] = "code-block"
            code_div["style"] = (
                "background-color: #f5f5f5; padding: 10px; font-family: Courier New; font-size: 9pt; margin: 10px 0;"
            )
            code.wrap(code_div)

    # Improve tables
    for table in soup.find_all("table"):
        table["border"] = "1"
        table["style"] = "border-collapse: collapse; width: 100%;"

        # Add thead if not present
        if not table.find("thead") and table.find("tr"):
            first_row = table.find("tr")
            thead = soup.new_tag("thead")
            table.insert(0, thead)
            thead.append(first_row)

            # Convert td to th in the header row
            for td in first_row.find_all("td"):
                th = soup.new_tag("th")
                th.string = td.string
                th["style"] = (
                    "background-color: #4472C4; color: white; font-weight: bold; text-align: center; padding: 5px;"
                )
                td.replace_with(th)

        # Style all cells
        for td in table.find_all("td"):
            td["style"] = "padding: 5px; border: 1px solid #DDDDDD;"

            # Right-align numeric cells
            if (
                td.string
                and td.string.strip().replace(".", "", 1).replace("-", "", 1).isdigit()
            ):
                td["style"] += " text-align: right;"

    # Improve images
    for img in soup.find_all("img"):
        if "width" not in img.attrs:
            img["width"] = "80%"
        img["style"] = "display: block; margin: 10px auto;"

    return str(soup)


def generate_gpt_analysis_docx(
    file_name,
    question,
    research_summary,
    code,
    output,
    image_paths=None,  # This will now be a list of base64 strings
    categorical_mappings=None,
):
    """
    Generate a DOCX file for GPT analysis, with sections for question, summary, code, output, images, and categorical mappings.

    This function properly handles bullet points and list formatting in the Word document.
    """
    """
    Generate a DOCX file for GPT analysis, with sections for question, summary, code, output, images, and categorical mappings.
    """
    doc = docx.Document()

    # Title
    doc.add_heading("GPT Analysis Report", 0)

    # Question
    doc.add_heading("Original Question", level=1)
    doc.add_paragraph(question)

    def add_html_to_doc(doc_obj, html_content_str):
        """Add HTML content to a docx document using BeautifulSoup parsing."""
        # Use BeautifulSoup parsing as the primary method
        soup = BeautifulSoup(html_content_str, "html.parser")

        # Ensure code block style exists
        code_style_name = create_code_block_style(doc_obj)

        def _process_node_recursive(
            node,
            current_paragraph,
            doc_obj,
            is_bold=False,
            is_italic=False,
            in_list_style=None,
        ):
            """
            Recursively processes HTML nodes and adds formatted content to the document.
            Creates new paragraphs/elements for block tags.
            """
            if isinstance(node, NavigableString):
                text = str(node)
                if current_paragraph:
                    # Add text to the current paragraph with formatting
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
                elif (
                    text.strip()
                ):  # If text is directly under body and not just whitespace
                    # Create a new paragraph for this text
                    p = doc_obj.add_paragraph()
                    run = p.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic

            elif node.name:  # It's a Tag
                new_bold = is_bold or (node.name in ["strong", "b"])
                new_italic = is_italic or (node.name in ["em", "i"])

                # Handle block-level elements by creating new DOCX elements
                if node.name.startswith("h") and len(node.name) == 2:
                    try:
                        level = int(node.name[1])
                        # Add heading and set current_paragraph to None for subsequent content
                        doc_obj.add_heading(node.get_text(strip=True), level=level)
                        current_paragraph = None
                    except ValueError:
                        pass  # Not h1-h6, treat as inline or process children
                elif node.name == "p":
                    if current_paragraph is None:
                        # If no current paragraph, create a new one for this <p>
                        p = doc_obj.add_paragraph()
                        # Process children into this new paragraph
                        for child in node.contents:
                            _process_node_recursive(
                                child, p, doc_obj, new_bold, new_italic
                            )
                        current_paragraph = (
                            p  # Set current paragraph to the one just created
                        )
                    else:
                        # If there is a current paragraph (e.g., inside an <li>),
                        # process children into the current paragraph without adding a newline.
                        # This prevents extra space for <p> tags within list items.
                        for child in node.contents:
                            _process_node_recursive(
                                child, current_paragraph, doc_obj, new_bold, new_italic
                            )
                        # current_paragraph remains the same
                elif node.name == "pre":
                    # Create a code block paragraph
                    code_text = node.get_text()
                    p = doc_obj.add_paragraph(code_text, style=code_style_name)
                    current_paragraph = p  # Set current paragraph
                elif node.name == "code":  # Inline code
                    if current_paragraph is None:
                        # If no current paragraph, create one for inline code
                        current_paragraph = doc_obj.add_paragraph()
                    run = current_paragraph.add_run(node.get_text())
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)  # Use a standard size
                    run.bold = new_bold  # Apply outer bold/italic
                    run.italic = new_italic
                    # Don't recurse into children of <code> as get_text() handles it
                elif node.name == "br":
                    if current_paragraph:
                        current_paragraph.add_run(
                            "\n"
                        )  # Add newline to current paragraph
                    else:
                        # If no current paragraph, create one and add newline
                        doc_obj.add_paragraph().add_run("\n")
                elif node.name == "hr":
                    # Add a simple horizontal rule representation
                    doc_obj.add_paragraph("─" * 50)
                    current_paragraph = None  # HR is a block
                elif node.name == "a" and node.has_attr("href"):
                    if current_paragraph is None:
                        current_paragraph = doc_obj.add_paragraph()
                    # Add hyperlink - this creates a new run internally
                    add_hyperlink(
                        current_paragraph, node["href"], node.get_text(strip=True)
                    )
                    # Note: Nested formatting within link text is not handled by add_hyperlink
                elif node.name == "img" and node.has_attr("src"):
                    try:
                        # Add image as a block element
                        doc_obj.add_picture(
                            node["src"], width=Inches(5)
                        )  # Use a default width
                        current_paragraph = None  # Image is a block
                    except Exception as img_e:
                        # Add placeholder if image fails
                        p_img_error = doc_obj.add_paragraph()
                        p_img_error.add_run(
                            f"[Image could not be loaded: {node['src']}. Error: {img_e}]"
                        )
                        current_paragraph = p_img_error  # Set current paragraph to the error message one
                elif node.name in ["ul", "ol"]:
                    # Handle lists
                    list_style = "List Bullet" if node.name == "ul" else "List Number"
                    for li in node.find_all("li", recursive=False):
                        # Only create a paragraph for non-empty list items
                        if li.get_text(strip=True):
                            # Each li is a new paragraph with list style
                            p_li = doc_obj.add_paragraph(style=list_style)

                            # Add a small amount of text to ensure bullet appears properly aligned
                            # This is a workaround for the bullet alignment issue
                            if not p_li.runs:
                                p_li.add_run("")

                            # Check if the first child is a text node and add it directly to ensure proper bullet alignment
                            children = list(li.contents)
                            if children and isinstance(children[0], NavigableString):
                                # Add the first text node directly to the paragraph
                                p_li.add_run(str(children[0]))
                                # Process remaining children
                                for child_li in children[1:]:
                                    _process_node_recursive(
                                        child_li,
                                        p_li,
                                        doc_obj,
                                        new_bold,
                                        new_italic,
                                        list_style,
                                    )
                            else:
                                # Process all children if first child is not a text node
                                for child_li in children:
                                    _process_node_recursive(
                                        child_li,
                                        p_li,
                                        doc_obj,
                                        new_bold,
                                        new_italic,
                                        list_style,
                                    )
                        # If li is empty, do nothing, don't create a paragraph
                    current_paragraph = (
                        None  # After list, subsequent content should be new para
                    )
                elif node.name == "table":
                    # Handle tables - simplified fallback
                    # This is still complex. Let's just add the text content for now.
                    # A proper table conversion would be needed for full fidelity.
                    doc_obj.add_paragraph(
                        f"[Table content placeholder]\n{node.get_text()}"
                    )
                    current_paragraph = None  # Table is a block
                else:
                    # For other tags (div, span, etc.), just process their children
                    # Need to ensure there's a current paragraph context for inline children
                    if current_paragraph is None:
                        # If no current paragraph, create one for inline content within this tag
                        current_paragraph = doc_obj.add_paragraph()
                    for child_node in node.contents:
                        # Pass the current paragraph and updated formatting states
                        _process_node_recursive(
                            child_node,
                            current_paragraph,
                            doc_obj,
                            new_bold,
                            new_italic,
                            in_list_style,
                        )

                # After processing a block element, subsequent content should start a new paragraph
                # This is handled by setting current_paragraph = None for block elements.
                # For inline elements or containers (like div/span), current_paragraph remains the same.

        # Start recursive processing from the body's contents
        # Pass None as the initial current_paragraph, so block elements create new ones.
        initial_paragraph = None
        if soup.body:
            for element in soup.body.contents:
                _process_node_recursive(element, initial_paragraph, doc_obj)
        else:  # If no body tag (e.g. very simple fragment), process children of soup directly
            for element in soup.contents:
                _process_node_recursive(element, initial_paragraph, doc_obj)

    # Research Summary
    if research_summary:
        doc.add_heading("Research Summary", level=1)
        # Convert markdown to HTML, then add to docx
        # Removed 'nl2br' and 'sane_lists' extensions to prevent extra newlines and list formatting
        html = markdown.markdown(
            research_summary,
            extensions=["fenced_code", "tables", "markdown.extensions.extra"],
        )
        add_html_to_doc(doc, html)  # Uses the helper above

    # Code
    if code:
        doc.add_heading("Code used for analysis", level=1)
        code_block_style = create_code_block_style(doc)
        for line in code.strip().split("\n"):
            p = doc.add_paragraph(line, style=code_block_style)

    # Helper: parse markdown-style bold/italic in a string and add to a paragraph
    # This is the function being improved.
    def add_markdown_text(paragraph, text):
        """Adds markdown-formatted text to a paragraph, robustly handling inline styles."""
        try:
            # Convert markdown to HTML
            # Using 'extra' for features like fenced code, footnotes, etc.
            # 'sane_lists' for better list behavior if lists were handled here.
            # 'nl2br' converts newlines to <br>
            html_content = markdown.markdown(
                text,
                extensions=[
                    "nl2br",
                    "fenced_code",
                    "tables",
                    "sane_lists",
                    "markdown.extensions.extra",
                ],
            )
            soup = BeautifulSoup(html_content, "html.parser")

            def _process_html_node(
                node, current_paragraph, is_bold=False, is_italic=False
            ):
                """
                Recursively processes HTML nodes and adds formatted runs to the paragraph.
                """
                if isinstance(node, NavigableString):
                    # Add text with current formatting state
                    run = current_paragraph.add_run(str(node))
                    run.bold = is_bold
                    run.italic = is_italic
                elif node.name:  # It's a Tag
                    # Update formatting state based on the tag
                    new_bold = is_bold or (node.name in ["strong", "b"])
                    new_italic = is_italic or (node.name in ["em", "i"])

                    if node.name == "br":
                        current_paragraph.add_run("\n")
                    elif node.name == "a" and node.has_attr("href"):
                        # Hyperlink handling: python-docx creates a new run for hyperlinks.
                        # For simplicity, nested formatting within link text is not preserved by this call.
                        # To support it, add_hyperlink would need to be more complex or we'd parse children
                        # and build the hyperlink run by run (very complex with python-docx).
                        link_text = node.get_text()  # Get all text from link
                        add_hyperlink(current_paragraph, node["href"], link_text)
                    elif node.name == "code":
                        # Inline code
                        run = current_paragraph.add_run(node.get_text())
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                        # Apply bold/italic if the <code> tag is nested within <strong> or <em>
                        run.bold = is_bold
                        run.italic = is_italic
                    # elif node.name in ['ul', 'ol', 'table', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'pre', 'blockquote']:
                    # Block-level elements: This function is designed to format text within
                    # a *single* given paragraph. Handling block elements that create new
                    # paragraphs (like lists, tables, headers) here would require access
                    # to the main 'doc' object and would change the function's contract.
                    # The original code's attempts were flawed.
                    # For now, we process their children inline if they appear.
                    # A more robust solution for full markdown documents would use html2docx
                    # or a more comprehensive parser at a higher level.
                    # If these appear from 'nl2br' or simple markdown, process children:
                    # For example, if 'text' is "para1\n\npara2", markdown might give <p>para1</p><p>para2</p>.
                    # This function will render "para1para2" in the current paragraph.
                    # If nl2br is used, "line1\nline2" becomes "line1<br>line2", handled above.
                    # For now, just recurse for children of any other tag.
                    # This means <p> tags from markdown will have their content processed inline.
                    # If `text` is truly complex markdown, this function might not be the right tool.
                    # It's best for inline formatting within a line/paragraph.
                    # Fallthrough to process children:
                    # for child in node.contents:
                    #    _process_html_node(child, current_paragraph, new_bold, new_italic)
                    else:
                        # Recursively process child nodes with updated formatting state
                        for child in node.contents:
                            _process_html_node(
                                child, current_paragraph, new_bold, new_italic
                            )

            # Process all top-level nodes from the parsed HTML fragment
            # BeautifulSoup wraps fragments in <html><body>...</body></html>.
            # We process the contents of <body>.
            if soup.body:
                for element in soup.body.contents:
                    _process_html_node(element, paragraph)
            else:  # Fallback if no body tag (e.g. very simple fragment)
                for element in soup.contents:
                    _process_html_node(element, paragraph)

        except Exception as e:
            # Fallback: add raw text with an error message
            error_run = paragraph.add_run(
                f"[Error processing markdown: {e}] Raw text below:\n"
            )
            error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for error
            paragraph.add_run(text)

    # Output (formatted as a code block)
    if output:
        doc.add_heading("Analysis Output", level=1)
        code_block_style = create_code_block_style(doc)  # Ensure style is available

        # Add the entire output as paragraphs with the code block style
        # Split the output by lines and add each line as a paragraph
        for line in output.strip().split("\n"):
            # Add an empty paragraph for blank lines to preserve spacing
            if not line.strip():
                doc.add_paragraph()
            else:
                doc.add_paragraph(line, style=code_block_style)

    # Categorical mappings
    if categorical_mappings:
        doc.add_heading("Categorical Variable Encodings", level=1)
        for col, mapping in categorical_mappings.items():
            doc.add_heading(f"Column '{col}'", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Original Value"
            table.rows[0].cells[1].text = "Encoded Value"
            for orig, enc in mapping.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(orig)
                row_cells[1].text = str(enc)

    # Images
    if image_paths:  # image_paths now contains base64 strings
        doc.add_heading("Generated Plots", level=1)
        for i, img_b64 in enumerate(image_paths):
            try:
                # Decode base64 string to bytes
                img_bytes = base64.b64decode(img_b64)

                # Create a temporary file to save the image bytes
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=".png"
                ) as temp_img_file:
                    temp_img_file.write(img_bytes)
                    temp_img_path = temp_img_file.name

                # Add the image from the temporary file to the document
                doc.add_picture(temp_img_path, width=Inches(5.5))

                # Clean up the temporary file immediately
                os.remove(temp_img_path)
            except Exception as e:
                doc.add_paragraph(f"[Image {i + 1} could not be loaded. Error: {e}]")

    # Save the document
    docx_file_path = file_name + ".docx"
    doc.save(docx_file_path)
    return docx_file_path


def fallback_markdown_to_docx(project_name, markdown_content):
    """Fallback method using the original implementation"""
    try:
        # Create a new document
        doc = docx.Document()

        # Convert markdown to HTML
        import markdown as md

        html_content = md.markdown(
            markdown_content,
            extensions=[
                "markdown.extensions.tables",
                "markdown.extensions.fenced_code",
                "markdown.extensions.codehilite",  # For syntax highlighting in <pre><code>
                "markdown.extensions.nl2br",  # Newlines to <br>
                "markdown.extensions.extra",  # Includes many useful extensions
            ],
        )

        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")

        # Create code block style
        code_style_name = create_code_block_style(doc)  # Get the style name

        # Process each element (simplified recursive approach)
        def _process_node_fallback(
            node,
            current_paragraph,
            doc_obj,
            is_bold=False,
            is_italic=False,
            in_list_style=None,
        ):
            if isinstance(node, NavigableString):
                if current_paragraph:  # Ensure we have a paragraph to add to
                    run = current_paragraph.add_run(str(node))
                    run.bold = is_bold
                    run.italic = is_italic
                # If no current_paragraph, this text might be lost or needs a new para.
                # This can happen if text is directly under soup.body without a <p> tag.
                elif str(node).strip():  # Only add if there's non-whitespace text
                    p = doc_obj.add_paragraph()
                    run = p.add_run(str(node))
                    run.bold = is_bold
                    run.italic = is_italic

            elif node.name:  # It's a Tag
                new_bold = is_bold or (node.name in ["strong", "b"])
                new_italic = is_italic or (node.name in ["em", "i"])

                new_paragraph_created = (
                    False  # Flag to track if this node created a new paragraph
                )

                if node.name.startswith("h") and len(node.name) == 2:
                    try:
                        level = int(node.name[1])
                        doc_obj.add_heading(node.get_text(strip=True), level=level)
                        current_paragraph = None  # After a heading, subsequent content should form new paragraphs
                        new_paragraph_created = True
                    except ValueError:  # Not h1-h6
                        pass  # Process children in current paragraph
                elif node.name == "p":
                    current_paragraph = doc_obj.add_paragraph()
                    new_paragraph_created = True
                elif node.name == "pre":
                    code_text = node.get_text()  # Usually pre contains code
                    current_paragraph = doc_obj.add_paragraph(
                        code_text, style=code_style_name
                    )
                    new_paragraph_created = True
                elif node.name == "code":  # Inline code
                    if current_paragraph is None:
                        current_paragraph = doc_obj.add_paragraph()
                    run = current_paragraph.add_run(node.get_text())
                    run.font.name = "Courier New"
                    run.bold = new_bold  # Apply outer bold/italic
                    run.italic = new_italic
                    return  # No further processing for children of <code>
                elif node.name == "br":
                    if current_paragraph:
                        current_paragraph.add_run("\n")
                    else:
                        doc_obj.add_paragraph().add_run(
                            "\n"
                        )  # New para if no current one
                elif node.name == "hr":
                    doc_obj.add_paragraph("─" * 50)  # Simple HR
                    current_paragraph = None
                    new_paragraph_created = True
                elif node.name == "a" and node.has_attr("href"):
                    if current_paragraph is None:
                        current_paragraph = doc_obj.add_paragraph()
                    add_hyperlink(
                        current_paragraph, node["href"], node.get_text(strip=True)
                    )
                    # Link text itself won't be further formatted by this simple call
                    return  # No further processing for children of <a>
                elif node.name == "img" and node.has_attr("src"):
                    try:
                        # Images are block elements in docx typically
                        doc_obj.add_picture(node["src"], width=Inches(5))
                        current_paragraph = None  # Image acts as a block
                        new_paragraph_created = True
                    except Exception as img_e:
                        p_img_error = doc_obj.add_paragraph()
                        p_img_error.add_run(
                            f"[Image could not be loaded: {node['src']}. Error: {img_e}]"
                        )
                        current_paragraph = p_img_error
                        new_paragraph_created = True
                elif node.name == "ul" or node.name == "ol":
                    list_style = "List Bullet" if node.name == "ul" else "List Number"
                    for li in node.find_all("li", recursive=False):
                        # Each li is a new paragraph
                        p_li = doc_obj.add_paragraph(style=list_style)

                        # Add a small amount of text to ensure bullet appears properly aligned
                        # This is a workaround for the bullet alignment issue
                        if not p_li.runs:
                            p_li.add_run("")

                        # Check if the first child is a text node and add it directly to ensure proper bullet alignment
                        children = list(li.contents)
                        if children and isinstance(children[0], NavigableString):
                            # Add the first text node directly to the paragraph
                            p_li.add_run(str(children[0]))
                            # Process remaining children
                            for child_li in children[1:]:
                                _process_node_fallback(
                                    child_li,
                                    p_li,
                                    doc_obj,
                                    new_bold,
                                    new_italic,
                                    list_style,
                                )
                        else:
                            # Process all children if first child is not a text node
                            for child_li in children:
                                _process_node_fallback(
                                    child_li,
                                    p_li,
                                    doc_obj,
                                    new_bold,
                                    new_italic,
                                    list_style,
                                )
                    current_paragraph = (
                        None  # After list, new content should be new para
                    )
                    new_paragraph_created = True
                    return  # Children of ul/ol (i.e. li) handled, so return
                elif node.name == "table":
                    # Simplified table handling for fallback
                    html_rows = node.find_all("tr")
                    if html_rows:
                        try:
                            first_row_cells = html_rows[0].find_all(["td", "th"])
                            if not first_row_cells:
                                raise ValueError("No cells in first row")
                            docx_tbl = doc_obj.add_table(
                                rows=0, cols=len(first_row_cells)
                            )
                            docx_tbl.style = "Table Grid"
                            for html_row_idx, html_row_elem in enumerate(html_rows):
                                cells_in_row = html_row_elem.find_all(["td", "th"])
                                if len(cells_in_row) == len(
                                    first_row_cells
                                ):  # Consistent columns
                                    row_cells = docx_tbl.add_row().cells
                                    for cell_idx, html_cell_elem in enumerate(
                                        cells_in_row
                                    ):
                                        cell_para = row_cells[cell_idx].paragraphs[0]
                                        # Process children of cell into cell_para
                                        for cell_child in html_cell_elem.contents:
                                            _process_node_fallback(
                                                cell_child,
                                                cell_para,
                                                doc_obj,
                                                new_bold,
                                                new_italic,
                                            )
                                        # Bold header
                                        if html_cell_elem.name == "th" or (
                                            html_row_idx == 0 and not node.find("thead")
                                        ):
                                            for run_in_cell in cell_para.runs:
                                                run_in_cell.bold = True
                            current_paragraph = None  # Table is a block
                            new_paragraph_created = True
                        except Exception as table_e:
                            p_table_error = doc_obj.add_paragraph()
                            p_table_error.add_run(f"[Error creating table: {table_e}]")
                            current_paragraph = p_table_error
                            new_paragraph_created = True
                    return  # Children of table handled

                # Generic recursion for other tags or if no new paragraph was made by this tag
                if not new_paragraph_created and current_paragraph is None:
                    # If we are in a state where current_paragraph is None (e.g. after a heading)
                    # and this element doesn't create its own paragraph (like <p> or <h1>),
                    # we need to create one for its inline content.
                    # However, skip if the node is just whitespace.
                    if isinstance(node, NavigableString) and not str(node).strip():
                        pass  # Skip whitespace nodes if no paragraph context
                    elif node.name not in [
                        "html",
                        "body",
                    ]:  # Avoid creating paras for html/body tags
                        current_paragraph = doc_obj.add_paragraph()

                # Process children with updated states
                for child_node in node.contents:
                    _process_node_fallback(
                        child_node,
                        current_paragraph,
                        doc_obj,
                        new_bold,
                        new_italic,
                        in_list_style,
                    )

        # Initial call to process the parsed HTML structure
        # Start with no current paragraph; block elements will create them.
        initial_paragraph = None
        if soup.body:
            for element in soup.body.contents:
                _process_node_fallback(element, initial_paragraph, doc)
        else:  # If no body, process children of soup directly
            for element in soup.contents:
                _process_node_fallback(element, initial_paragraph, doc)

        # Save the document
        docx_path = f"{project_name}_fallback.docx"
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        print(f"Fallback markdown_to_docx conversion failed: {e}")
        # As a last resort, try to use html2docx directly if everything else fails
        try:
            from html2docx import html2docx as h2d_converter

            buf = h2d_converter(
                markdown_content, project_name
            )  # html2docx expects html string
            docx_path_h2d = f"{project_name}_html2docx_fallback.docx"
            with open(docx_path_h2d, "wb") as f:
                f.write(buf.getvalue())
            return docx_path_h2d
        except Exception as h2d_e:
            print(f"html2docx direct fallback also failed: {h2d_e}")
            return None


if __name__ == "__main__":
    # Example usage
    project_name = "example_project"
    markdown_content_example = """
# Main Title: ***Bold & Italic***

This is a paragraph with **bold text**, *italic text*, and ***bold italic text***.
Here's `inline code`. And a [link to Google](https://www.google.com).

Adjacent: *italic***bold***italic* **bold**.

Nested: **bold *bold-italic* bold** and *italic **italic-bold***.

## Lists
- Unordered item 1
  - Nested unordered
- Unordered item 2 `code in list`

1. Ordered item 1
   1. Nested ordered
2. Ordered item 2 **bold in list**

> This is a blockquote.
> With another line.

---

## Code Block
```python
def hello_world():
    # This is a comment
    print("Hello, World! From Python.")
    print("***Just asterisks, not formatting***")
```

## Table Example

| Header 1 | Header 2 | Header 3         |
|----------|:--------:|------------------|
| *CellA1* |  CellA2  | `CellA3 code`    |
| CellB1   | **CellB2** | ***CellB3 BI***  |

Another paragraph after everything.
    """
    output_file_main = generate_gpt_analysis_docx(
        project_name,
        "Example question: How to parse complex markdown?",
        "## Summary of Findings\n\n- Parsing **markdown** is *tricky*.\n- **Bold, _italic_, and `code`** are common.\n- Nested: ***bold with _italic_ inside***.",
        "def complex_code_example(param1, param2):\n    # This is some Python code\n    if param1 > param2:\n        print(f'{param1} is greater')\n    else:\n        print(f'{param2} is greater or equal')\n    return abs(param1 - param2)",
        "### Sample Output\n\nThis is the output from the analysis. It can include:\n\n- **Bolded results**\n- *Emphasized points*\n- `monospace_code_snippets`\n- And even ***bolded italics***.\n\nLine breaks should be preserved.\n\nAnother line here.\n\n```\nPreformatted block\n  should also work\n    with indents.\n```\n\n| Col1 | Col2 |\n|------|------|\n| Val1 | Val2 |",
        image_paths=None,  # Provide path to a test image if available
        categorical_mappings={
            "gender": {"Male": 0, "Female": 1},
            "outcome": {"Survived": 1, "Died": 0},
        },
    )
    if output_file_main:
        print(f"Main Docx file created: {output_file_main}")
    else:
        print("Main Docx file creation failed.")

    # Test fallback with the same complex markdown
    # output_file_fallback = fallback_markdown_to_docx(project_name, markdown_content_example)
    # if output_file_fallback:
    #     print(f"Fallback Docx file created: {output_file_fallback}")
    # else:
    #     print("Fallback Docx file creation failed.")
