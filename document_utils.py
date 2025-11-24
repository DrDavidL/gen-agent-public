"""
Utility functions for processing document files (PDF, DOCX)
"""

import os
import re
import tempfile
from typing import List, Tuple
import streamlit as st

# Import models

# For PDF processing
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# For DOCX processing
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf(file_obj) -> Tuple[str, bool]:
    """
    Extract text from a PDF file. If no text is extracted, use Azure gpt-4o-AutoAnalyzer to analyze as image.

    Args:
        file_obj: File-like object containing PDF data

    Returns:
        Tuple of (extracted text, success flag)
    """
    if not PDF_AVAILABLE:
        return "PDF extraction not available. Please install PyPDF2.", False

    try:
        text_content = []
        pdf_reader = PyPDF2.PdfReader(file_obj)

        # Get number of pages
        num_pages = len(pdf_reader.pages)

        # Extract text from each page
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text_content.append(page.extract_text())

        # Join all pages with page markers
        full_text = "\n\n--- Page Break ---\n\n".join(text_content)

        # Clean up text
        full_text = re.sub(r"\s+", " ", full_text)
        full_text = re.sub(r"\n\s*\n", "\n\n", full_text)

        # If no text extracted, try Azure gpt-4o-AutoAnalyzer as image
        if not full_text.strip():
            try:
                import tempfile
                from pdf2image import convert_from_bytes
                import requests
                import os

                # Read PDF bytes
                pdf_bytes = file_obj.read()

                # Convert first page of PDF to image using pdf2image
                images = convert_from_bytes(pdf_bytes, first_page=1, last_page=1)
                if not images:
                    return "PDF has no pages to analyze.", False
                img = images[0]
                with tempfile.NamedTemporaryFile(
                    suffix=".png", delete=False
                ) as img_file:
                    img.save(img_file, format="PNG")
                    img_path = img_file.name

                # Read image bytes
                with open(img_path, "rb") as img_file_rb:
                    image_bytes = img_file_rb.read()

                # Azure OpenAI Vision API call
                AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY") or (
                    hasattr(st, "secrets") and st.secrets.get("AZURE_OPENAI_API_KEY")
                )
                AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT") or (
                    hasattr(st, "secrets") and st.secrets.get("AZURE_OPENAI_ENDPOINT")
                )
                DEPLOYMENT_NAME = (
                    os.environ.get("AZURE_VISION_DEPLOYMENT")
                    or (
                        hasattr(st, "secrets")
                        and st.secrets.get("AZURE_VISION_DEPLOYMENT")
                    )
                    or "gpt-4o-AutoAnalyzer"
                )
                API_VERSION = (
                    os.environ.get("API_VERSION")
                    or (hasattr(st, "secrets") and st.secrets.get("API_VERSION"))
                    or "2024-05-01-preview"
                )

                if (
                    not AZURE_OPENAI_API_KEY
                    or not AZURE_OPENAI_ENDPOINT
                    or not DEPLOYMENT_NAME
                    or not API_VERSION
                ):
                    return (
                        "Azure OpenAI credentials, vision deployment, or API version not set for vision analysis.",
                        False,
                    )

                url = f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{DEPLOYMENT_NAME}/chat/completions?api-version={API_VERSION}"

                headers = {
                    "api-key": AZURE_OPENAI_API_KEY,
                    "Content-Type": "application/json",
                }

                import base64

                image_base64 = base64.b64encode(image_bytes).decode("utf-8")
                vision_prompt = [
                    {
                        "role": "system",
                        "content": "You are an expert document analyzer. Extract all readable text and summarize the content of the provided PDF page image. If tables or figures are present, describe them as well.",
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{image_base64}"
                                },
                            }
                        ],
                    },
                ]

                data = {"messages": vision_prompt, "max_tokens": 2048}

                response = requests.post(url, headers=headers, json=data, timeout=60)
                if response.status_code == 200:
                    result = response.json()
                    # Extract the text from the response
                    if "choices" in result and result["choices"]:
                        vision_text = result["choices"][0]["message"]["content"]
                        # Clean up temp files
                        try:
                            os.remove(img_path)
                        except Exception:
                            pass
                        return vision_text, True
                    else:
                        return "Azure Vision API did not return any content.", False
                else:
                    return (
                        f"Azure Vision API error: {response.status_code} {response.text}",
                        False,
                    )

            except Exception as vision_e:
                return (
                    f"Error extracting text from PDF using vision model: {vision_e}",
                    False,
                )
        return full_text, True
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}", False


def extract_text_from_docx(file_obj) -> Tuple[str, bool]:
    """
    Extract text from a DOCX file.

    Args:
        file_obj: File-like object containing DOCX data

    Returns:
        Tuple of (extracted text, success flag)
    """
    if not DOCX_AVAILABLE:
        return "DOCX extraction not available. Please install python-docx.", False

    try:
        # Save to a temporary file since python-docx needs a file path
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_file.write(file_obj.getvalue())
            temp_path = temp_file.name

        # Open the document
        doc = docx.Document(temp_path)

        # Extract text from paragraphs
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        # Clean up the temporary file
        try:
            os.unlink(temp_path)
        except:
            pass

        # Join all content
        full_text = "\n\n".join(text_content)

        return full_text, True
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}", False


def process_document_file(uploaded_file) -> Tuple[str, bool]:
    """
    Process an uploaded document file (PDF or DOCX) and extract its text.
    If PDF text extraction fails, will attempt Azure gpt-4o-AutoAnalyzer vision analysis.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Tuple of (extracted text, success flag)
    """
    file_type = uploaded_file.name.split(".")[-1].lower()

    if file_type == "pdf":
        if not PDF_AVAILABLE:
            return (
                "PDF processing requires PyPDF2. Please install it with 'pip install PyPDF2'.",
                False,
            )
        return extract_text_from_pdf(uploaded_file)

    elif file_type in ["docx", "doc"]:
        if not DOCX_AVAILABLE:
            return (
                "DOCX processing requires python-docx. Please install it with 'pip install python-docx'.",
                False,
            )
        return extract_text_from_docx(uploaded_file)

    else:
        return f"Unsupported document type: {file_type}", False


@st.cache_data(ttl=3600, show_spinner=False)
def create_document_chunks(
    text: str, chunk_size: int = 1000, chunk_overlap: int = 200
) -> List[str]:
    """
    Split document text into chunks for vector store processing.
    Uses Streamlit caching to avoid reprocessing the same text.

    Args:
        text: The document text to split
        chunk_size: Size of each chunk
        chunk_overlap: Overlap between chunks

    Returns:
        List of text chunks
    """
    if not text:
        return []

    # Clean the text - use a single regex operation for better performance
    text = re.sub(
        r"(\s+|\n\s*\n)", lambda m: " " if m.group(1).strip() == "" else "\n\n", text
    )

    # Try to split on meaningful boundaries
    # First try to split by double newlines (paragraphs)
    paragraphs = re.split(r"\n\s*\n", text)

    # If we have very few paragraphs, try splitting by single newlines
    if len(paragraphs) < 3:
        paragraphs = re.split(r"\n", text)

    # If we still have very few chunks, split by sentences
    if len(paragraphs) < 3:
        # Split by sentences (simple heuristic)
        paragraphs = re.split(r"(?<=[.!?])\s+", text)

    # Create smaller chunks for better retrieval
    chunks = []
    current_chunk = ""

    # Add document metadata to each chunk for better context
    metadata_prefix = "Document Content: "

    for para in paragraphs:
        # Skip empty paragraphs
        if not para.strip():
            continue

        # If adding this paragraph would exceed chunk size, save current chunk and start a new one
        if len(current_chunk) + len(para) > chunk_size and current_chunk:
            # Add metadata prefix to the chunk
            chunks.append(metadata_prefix + current_chunk)

            # Start new chunk with overlap from the end of the previous chunk
            if len(current_chunk) > chunk_overlap:
                # Find the last complete sentence within the overlap
                overlap_text = current_chunk[-chunk_overlap:]
                sentence_break = (
                    overlap_text.rfind(". ") + 2
                )  # +2 to include the period and space

                if sentence_break > 0:
                    current_chunk = current_chunk[-(chunk_overlap - sentence_break) :]
                else:
                    current_chunk = current_chunk[-chunk_overlap:]
            else:
                current_chunk = current_chunk  # Keep the entire short chunk

        # Add paragraph to current chunk
        if current_chunk:
            current_chunk += " " + para
        else:
            current_chunk = para

    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(metadata_prefix + current_chunk)

    # If we have very few or no chunks, create overlapping chunks manually
    if len(chunks) < 2 and len(text) > chunk_size:
        chunks = []
        for i in range(0, len(text), chunk_size - chunk_overlap):
            chunk = text[i : i + chunk_size]
            if chunk:  # Skip empty chunks
                chunks.append(chunk)

    # Print some debug info
    print(f"Created {len(chunks)} chunks from document text of length {len(text)}")
    if chunks:
        print(
            f"Average chunk size: {sum(len(c) for c in chunks) / len(chunks):.1f} characters"
        )
        # Print a sample of the first chunk to help with debugging
        if len(chunks) > 0:
            first_chunk = chunks[0]
            print(f"First chunk sample (first 100 chars): {first_chunk[:100]}...")

    return chunks
